import streamlit as st
import json
import re
import requests
from datetime import datetime
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import funcoes_texto as ft
import agentes_nocivos as agnocivo


# Função para criar texto_base
def texto_base(doc, fundamento_questao):
    if fundamento_questao == 1:
        for i, paragrafo in enumerate(fundamento_base):
            parag = doc.add_paragraph(paragrafo)
            parag.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY 
            # if i in [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 25, 26, 27, 28, 31, 33, 59]:
            #     parag.paragraph_format.first_line_indent = Cm(2) 
            # elif i in [13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24, 29, 30, 32, 34, 35, 36, 37, 38 ,39, 40 ,41, 42, 43, 44, 45, 46, 47 ,48 ,49, 50 ,51, 52, 53, 54, 55, 56, 57, 58]:
            #     parag.paragraph_format.left_indent = Cm(2)
            parag.paragraph_format.first_line_indent = Cm(2)

    if fundamento_questao == 2:
        for linha in fundamento_custom.split("\n"):
            if linha.split():
                parag = doc.add_paragraph(linha.strip())
                parag.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                parag.paragraph_format.first_line_indent = Cm(2)

# Função auxiliar para formatar dias em anos/meses/dias
def formatar_dias_em_anos_meses_dias(dias: int) -> str:
    anos = dias // 360
    resto = dias % 360
    meses = resto // 30
    dias_restantes = resto % 30
    partes = []
    if anos:
        partes.append(f"{anos} ano(s)")
    if meses:
        partes.append(f"{meses} mês(es)")
    if dias_restantes or not partes:
        partes.append(f"{dias_restantes} dia(s)")
    return ", ".join(partes)
# Função para transformar data
def formatar_data_iso_para_br(data_iso: str) -> str:
    try:
        return datetime.fromisoformat(data_iso.replace("Z", "")).strftime("%d/%m/%Y")
    except:
        return data_iso
# Função para transformar data
def iso_to_br(date_str):
    try:
        return datetime.fromisoformat(date_str.replace("Z", "")).strftime("%d/%m/%Y")
    except:
        return ""
# Função para transformar data
def corrigir_data_em_string(texto: str) -> str:
    padrao = r"(\d{1,2})/(\d{1,2})/(\d{4})"
    def converter_para_br(match):
        mm, dd, aaaa = match.groups()
        try:
            data = datetime.strptime(f"{mm}/{dd}/{aaaa}", "%m/%d/%Y")
            return data.strftime("%d/%m/%Y")
        except ValueError:
            return match.group(0)  # retorna o original se der erro
    return re.sub(padrao, converter_para_br, texto)
# Função para tratamento do arquivo .calc de entrada
def parse_calc_format(dados_orig):
    from datetime import datetime
    FORMA_CONTAGEM_ENUM = [
        "NAO_COMPUTAR", "COMUM", "ESPECIAL_25", "ESPECIAL_20", "ESPECIAL_15", "RURAL", "MAGISTERIO"
    ]
    NATUREZA_VINCULO_ENUM = [
        "NAO_INFORMADO", "BENEFICIO_INCAPACIDADE", "EMPREGADO", "EMPREGADO_DOMESTICO",
        "CONTRIBUINTE_INDIVIDUAL", "TRABALHADOR_AVULSO", "SEGURADO_ESPECIAL", "SEGURADO_FACULTATIVO"
    ]
    GRAU_DEFICIENCIA_ENUM = ["SEM_DEFICIENCIA", "LEVE", "MODERADO", "GRAVE"]
    def iso_to_br(iso):
        try:
            return datetime.fromisoformat(iso.replace("Z", "")).strftime("%d/%m/%Y")
        except:
            return ""
    novo_payload = {}
    # Processo
    novo_payload["processo"] = dados_orig.get("dados", {}).get("processo", {})
    # Segurado
    segurado = dados_orig.get("dados", {}).get("segurado", {})
    sexo_map = {0: "M", 1: "F"}
    tipo_map = {0: "COMUM", 1: "PCD", 2: "RURICOLA"}
    novo_payload["dadosSegurado"] = {
        "nascimento": iso_to_br(segurado.get("nascimento", "")),
        "sexo": sexo_map.get(segurado.get("sexo"), ""),
        "tipo": tipo_map.get(segurado.get("tipo"), "")
    }
    # Períodos de contribuição
    periodos_raw = dados_orig.get("dados", {}).get("periodos", {})
    periodos_convertidos = []
    for p in periodos_raw.values():
        periodo = {
            "dataInicial": iso_to_br(p.get("dataInicial", "")),
            "dataFinal": iso_to_br(p.get("dataFinal", "")),
            "descricao": p.get("descricao", ""),
            "formaContagem": FORMA_CONTAGEM_ENUM[p.get("formaContagem", 1)],
            "grauDeficiencia": GRAU_DEFICIENCIA_ENUM[p.get("grauDeficiencia", 0)],
            "contarCarencia": p.get("contarCarencia", True),
            "naturezaVinculo": NATUREZA_VINCULO_ENUM[p.get("naturezaVinculo", 0)],
            "observacao": p.get("observacao", "")
        }
        periodos_convertidos.append(periodo)
    novo_payload["periodosContribuicao"] = periodos_convertidos
    # Períodos PCD
    periodos_pcd_raw = dados_orig.get("dados", {}).get("periodosPcd", {})
    periodos_pcd_convertidos = []
    for p in periodos_pcd_raw.values():
        pcd = {
            "dataInicial": iso_to_br(p.get("dataInicial", "")),
            "dataFinal": iso_to_br(p.get("dataFinal", "")),
            "observacao": p.get("observacao", ""),
            "grauDeficiencia": GRAU_DEFICIENCIA_ENUM[p.get("grauDeficiencia", 0)]
        }
        periodos_pcd_convertidos.append(pcd)
    novo_payload["periodosPcd"] = periodos_pcd_convertidos
    # Salários e benefícios
    novo_payload["salarios"] = dados_orig.get("salarios", {})
    novo_payload["beneficios"] = dados_orig.get("beneficios", {})
    # Parâmetros Gerais e Tempo
    dados = dados_orig.get("dados", {})
    param_gerais = dados.get("parametrosGerais", {})
    param_tempo = dados.get("parametrosTempo", {})
    param_rmi_desv = dados.get("parametrosRmiDesvinculada", {})
    marcoCarencia_map = ["IMPLEMENTO_IDADE", "DER"]
    requisito_map = ["INICIO_CONTRIBUICAO_COM_DEFICIENCIA", "INICIO_DEFICIENCIA"]
    novo_payload["opcoesContagem"] = {
        "especie": param_gerais.get("especie"),
        "marcoCarencia": marcoCarencia_map[param_tempo.get("marcoCarencia", 0)],
        "requisitoB41Pcd": requisito_map[param_tempo.get("requisitoB41Pcd", 0)],
        "carenciaAposEc103": param_tempo.get("carenciaAposEc103", True),
        "aplicarFatorAposEc103": param_rmi_desv.get("aplicarFatorAposEc103", False),
        "analisarCompetenciasAbaixoDoMinimo": param_tempo.get("excluirCompetenciasAbaixoDoMinimoAposEc103", True)
    }
    novo_payload["dataApuracao"] = iso_to_br(param_gerais.get("der", ""))
    return novo_payload
# Função para criar json para envio à API
def montar_payload_final(novo_payload):
    payload_final = {
        "dataApuracao": novo_payload.get("dataApuracao"),
        "processo": f"{processo_formatado}",
        "dadosSegurado": novo_payload.get("dadosSegurado", {}),
        "opcoesContagem": novo_payload.get("opcoesContagem", {}),
        "periodosContribuicao": [],
    }
    # Garantir ordem e campos esperados em periodosContribuicao
    for p in novo_payload.get("periodosContribuicao", []):
        periodo = {
            "dataInicial": p.get("dataInicial", ""),
            "dataFinal": p.get("dataFinal", ""),
            "descricao": p.get("descricao", ""),
            "formaContagem": p.get("formaContagem", ""),
            "contarCarencia": p.get("contarCarencia", True),
            "prioritario": p.get("prioritario", False),
            "computavelDesde": p.get("dataInicial", ""),  # mesmo valor da dataInicial
#            "fatorConversaoModificado": float(p.get("fatorConversaoModificado", 1.0) or 1.0),
            "naturezaVinculo": p.get("naturezaVinculo", ""),
            "observacao": p.get("observacao", "")
        }
        payload_final["periodosContribuicao"].append(periodo)
    # Adiciona periodosPcd se houver
    if "periodosPcd" in payload:
        payload_final["periodosPcd"] = payload["periodosPcd"]
    return payload_final
#=================================================================================================================
#=================================================================================================================

API_URL = "https://fabrica-de-calculos.dev.trf3.jus.br/api/v1/tempo/calcular"

fundamento_questao = st.radio(
"**Relatório e fundamentação jurídica**\n\n"
"Toda sentença gerada possui dispensa de relatório e uma fundamentação jurídica básica, que vai até onde se inicia a análise do caso concreto.\n\n"  
"Você pode usar esta fundamentação, ou fornecer sua **própria fundamentação**, com ou sem relatório (fazendo menção à dispensa, se for o caso.)\n\n"  
"**Como você deseja prosseguir?**",
[1, 2],
format_func=lambda x: "Vou usar a fundamentação padrão deste aplicativo." if x == 1 else "Desejo fornecer a minha fundamentação."
)
if fundamento_questao == 1:
     fundamento_base = [
    f"""Vistos.""",
    f"""Trata-se de pedido envolvendo aposentadoria.""",
    f"""Dispensado o relatório nos termos do art. 38 da Lei n. 9.099/95.""",
    f"""DECIDO.""",
    f"""O feito comporta julgamento imediato.""",
    f"""O Juízo é competente, pois o valor da causa é inferior a 60 salários-mínimos. Não há que se falar em renúncia ao valor excedente à esta alçada.""",
    f"""O tema 1030 do STJ prevê que: “Ao autor que deseje litigar no âmbito de Juizado Especial Federal Cível, é lícito renunciar, de modo expresso e para fins de atribuição de valor à causa, ao montante que exceda os 60 (sessenta) salários-mínimos previstos no art. 3º, caput, da Lei 10.259/2001, aí incluídas, sendo o caso, até doze prestações vincendas, nos termos do art. 3º, § 2º, da referida lei, c/c o art. 292, §§ 1º e 2º, do CPC/2015.”""",
    f"""Logo, a renúncia só tem sentido nos casos em que o valor da causa, apurado sem qualquer renúncia, pudesse ser superior à alçada, pois nesta hipótese a renúncia daria ensejo à parte litigar sob a competência do Juizado. Não é o caso dos autos, onde o valor da causa apontado já está abaixo da alçada, e a parte ré não aponta erro na sua apuração.""",
    f"""Importante frisar que a eventual aplicação do Tema 629 do STJ, embora possa resultar em extinção, sem resolução do mérito, depende da apreciação de provas. Por isso, por pragmatismo, será analisada ao tempo do mérito, se o caso implicar seu reconhecimento.""",
    f"""Foi cumprido o estipulado no Tema 350 do STF. Partes legítimas e bem representadas.""",
    f"""Passo ao mérito.""",
    f"""Prejudicialmente, analiso a prescrição. Não há prescrição de fundo de direito, mas apenas das parcelas vencidas, na forma da súmula 85 do STJ. Estão prescritas eventuais parcelas anteriores ao quinquênio que precede a propositura da ação.""",
    f"""APOSENTADORIA PROGRAMADA NA EC 103/19""",
    f"""Com a promulgação da Emenda Constitucional nº 103, de 13 de novembro de 2019, o sistema previdenciário brasileiro passou por uma profunda reformulação, com especial impacto sobre os requisitos e as modalidades de concessão das aposentadorias no âmbito do Regime Geral de Previdência Social (RGPS). As antigas espécies denominadas aposentadoria por idade e aposentadoria por tempo de contribuição foram reagrupadas sob o título de aposentadoria programada, constituindo a nova espinha dorsal da previdência social pública.""",
    f"""A aposentadoria programada é devida aos segurados filiados ao RGPS a partir da vigência da EC nº 103/2019. São exigidos, cumulativamente, os seguintes requisitos: (i) 62 anos de idade para a mulher e 65 anos para o homem; (ii) 15 anos de tempo de contribuição para a mulher e 20 anos para o homem; e (iii) 180 contribuições mensais de carência.""",
    f"""Para os segurados que já se encontravam filiados ao RGPS até a data da promulgação da referida emenda constitucional, foram instituídas regras de transição, com vistas a mitigar os impactos da nova ordem jurídica sobre situações consolidadas. As regras de transição têm aplicação imediata aos pedidos protocolados a partir de 13 de novembro de 2019, assegurado o direito adquirido àqueles que já preenchiam os requisitos para aposentadoria pelas normas anteriores.""",
    f"""As principais regras de transição são as seguintes:""",
    f"""1. APOSENTADORIA POR IDADE – REGRA DE TRANSIÇÃO (ART. 18 DA EC 103/2019):""",
    f"""Aplica-se à mulher que tenha, no mínimo, 15 anos de tempo de contribuição. A idade mínima, originalmente de 60 anos, é acrescida de seis meses a cada ano, iniciando-se em 2020, até atingir 62 anos em 2023. Para o homem, a regra permanece inalterada: 65 anos de idade e 15 anos de tempo de contribuição.""",
    f"""2. APOSENTADORIA POR TEMPO DE CONTRIBUIÇÃO – REGRAS DE TRANSIÇÃO:""",
    f"""As aposentadorias por tempo de contribuição, ainda que extintas para os novos segurados, podem ser concedidas aos filiados até 13 de novembro de 2019 mediante o cumprimento de uma das seguintes regras de transição:""",
    f"""a) SISTEMA DE PONTOS (ART. 15 DA EC 103/2019):""",
    f"""Nesta regra, exige-se o cumprimento simultâneo de:""",
    f"""30 anos de tempo de contribuição para mulheres e 35 para homens;""",
    f"""180 meses de carência;""",
    f"""Soma da idade com o tempo de contribuição, atingindo pontuação mínima progressiva: em 2019, 86 pontos (mulher) e 96 pontos (homem); acrescendo-se 1 ponto por ano, até alcançar, em 2033, 100 pontos (mulher) e 105 pontos (homem).""",
    f"""b) IDADE MÍNIMA PROGRESSIVA (ART. 16 DA EC 103/2019):""",
    f"""Exige-se:""",
    f"""30 anos de tempo de contribuição para a mulher e 35 para o homem;""",
    f"""180 meses de carência;""",
    f"""Idade mínima inicial de 56 anos (mulher) e 61 anos (homem), com acréscimo de 6 meses por ano, até atingir, em 2031, a idade de 62 anos (mulher) e 65 anos (homem).""",
    f"""c) PEDÁGIO DE 50% (ART. 17 DA EC 103/2019):""",
    f"""Aplicável apenas a quem, na data da EC 103/2019, já havia cumprido 28 anos (mulher) ou 33 anos (homem) de contribuição. Exige-se:""",
    f"""Cumprimento do tempo mínimo de 30/35 anos;""",
    f"""Pagamento de um pedágio de 50% do tempo que faltava para atingir o tempo mínimo em 13/11/2019;""",
    f"""180 contribuições mensais de carência.""",
    f"""d) PEDÁGIO DE 100% COM IDADE MÍNIMA (ART. 20 DA EC 103/2019):""",
    f"""Exige-se:""",
    f"""Idade mínima de 57 anos (mulher) e 60 anos (homem);""",
    f"""30/35 anos de contribuição acrescidos de um pedágio de 100% do tempo que faltava para atingir esse tempo mínimo em 13/11/2019;""",
    f"""180 meses de carência.""",
    f"""3. APOSENTADORIA PROGRAMADA ESPECIAL (ART. 19 DA EC 103/2019):""",
    f"""Voltada a trabalhadores que atuam expostos a agentes nocivos, a nova sistemática exige, além dos 180 meses de carência, os seguintes critérios:""",
    f"""Idade mínima de 55, 58 ou 60 anos, conforme a atividade especial exija 15, 20 ou 25 anos de exposição efetiva;""",
    f"""Tempo de contribuição correspondente à atividade insalubre.""",
    f"""Na regra de transição (art. 21), exige-se pontuação mínima (soma entre idade e tempo de contribuição), aliada ao tempo de exposição: 66 pontos (15 anos), 76 pontos (20 anos) ou 86 pontos (25 anos).""",
    f"""4. APOSENTADORIA PROGRAMADA DO PROFESSOR (ART. 201, §8º, DA CF):""",
    f"""Aplicável a quem exerce exclusivamente funções de magistério na educação básica. A regra permanente exige:""",
    f"""57 anos de idade (mulher) e 60 anos (homem);""",
    f"""25 anos de tempo de contribuição (mulher) e 30 anos (homem);""",
    f"""180 meses de carência.""",
    f"""Regras de transição:""",
    f"""Pontuação: 81 pontos (mulher) e 91 (homem) em 2019, com acréscimo de 1 ponto ao ano até 92/100.""",
    f"""Idade mínima: Inicialmente 51 anos (mulher) e 56 anos (homem), com acréscimo de 6 meses por ano, até alcançar 57/60.""",
    f"""Pedágio de 100%: Para quem estava próximo do tempo mínimo em 13/11/2019. Exige-se a idade mínima de 52 anos (mulher) e 55 (homem), o tempo mínimo de 25/30 anos de contribuição acrescido de um pedágio de 100% do tempo faltante, além da carência.""",
    f"""DIREITO ADQUIRIDO ANTES DA EC 103/2019""",
    f"""1. APOSENTADORIA POR TEMPO DE CONTRIBUIÇÃO:""",
    f"""As modificações introduzidas pela EC 103/2019 alteraram profundamente o sistema de acesso às aposentadorias do RGPS. As regras de transição refletem um esforço de preservação dos direitos em formação, assegurando uma adaptação gradativa às novas exigências constitucionais.""",
    f"""Observe-se, como dito, que é assegurada a concessão do benefício, conforme às regras anteriores à EC 103/2019, aos segurados que tenham cumprido os requisitos para concessão antes da entrada em vigor da emenda. A par desta constatação, observa-se o seguinte regramento em relação à aposentadoria por tempo de contribuição.""",
    f"""(i) segurado que ingressou no RGPS antes da vigência da Lei 8.213/91 sem preenchimento de todos os requisitos para a aposentadoria à época da EC 20/98 (16/12/98), deve comprovar 35 anos de tempo de serviço/contribuição, se homem, e 30 anos, se mulher, além de número de contribuições (carência) correspondente ao ano de implemento das condições previsto na tabela constante do artigo 142 da LBPS;""",
    f"""(ii) segurado que ingressou no RGPS a partir da vigência da Lei 8.213/91 (25/07/91) e antes da EC 20/98 (16/12/98), sem preenchimento de todas as condições à época da vigência da EC 20/98, deve comprovar 35 anos de tempo de serviço/contribuição, se homem, e 30 anos, se mulher, sendo ainda permitida a consideração do tempo de serviço como tempo de contribuição, em conformidade com o artigo 4º da EC 20/98. A carência exigida para o benefício é de 180 contribuições (art. 25, inciso II, LBPS);""",
    f"""(iii) segurado que ingressou no RGPS a partir da vigência da EC Nº 20/98 (16/12/98), deve comprovar tempo de efetiva contribuição correspondente a 35 anos, se homem, e 30 anos, se mulher, aplicando-se o disposto no artigo 55 da Lei 8.213/91 e art. 60 do RPS, que descrevem hipóteses consideradas como tempo de contribuição, até que lei específica discipline a matéria em consonância com o comando do artigo 4º da EC 20/98. A carência para o benefício é de 180 contribuições (art. 25, inciso II, da Lei Nº 8.213/91).""",
    f"""2. APOSENTADORIA POR IDADE URBANA, RURAL E HÍBRIDA:""",
    f"""No que se refere à aposentadoria por idade urbana, com base no art. 201, §7º, CF/88 e artigos 48 a 50 da Lei 8.213/91, verificava-se que sua concessão exigia a idade mínima de 65 anos para o homem e 60 para a mulher, e o cumprimento da carência de 180 meses de contribuição, para segurados inscritos no RGPS após 24/07/1991, observando-se a tabela progressiva do art. 142, Lei n. 8.213/91, para os inscritos anteriormente.""",
    f"""Para o trabalhador rural, a concessão do benefício por idade rural exige a comprovação da atividade rural em período imediatamente anterior, ainda que de forma descontínua, ao requerimento administrativo (advento da idade) e o cumprimento do prazo de carência de 180 meses, ou aquele previsto no art. 142 da Lei n. 8.213/91, tal qual se dava para os segurados urbanos. Os requisitos de idade são reduzidos para 60 anos, o homem, e 55 anos, a mulher.""",
    f"""A aposentadoria por idade rural não foi alterada pela EC 103/19, sendo estes os requisitos ainda vigentes. Importante, para esta aposentadoria, as disposições da Lei n. 11.718/2008, que revogou o art. 143 da Lei n. 8.213/91 passou a fazer distinção entre trabalhador rural empregado, trabalhador rural contribuinte individual, no que se refere à necessidade de contribuição a partir de 2010 e 2020, para contagem de carência, em contraste com sua inexigibilidade para os segurados especial, nos termos do art. 39 da Lei n. 8.213/91.""",
    f"""Para o trabalhador com vínculos rurais e urbanos, a idade é 65 anos para homens e 60 anos para mulher. A contagem da carência leva em conta os vínculos urbanos e rurais, e, em relação a estes últimos, seguindo as mesmas regras de prova e desnecessidade (ou não) de recolhimento de contribuição, previstas para a aposentadoria rural (Lei n. 11.718/2008).""",
    f"""3. APOSENTADORIA DA PESSOA COM DEFICIÊNCIA (LC 142/2013):""",
    f"""A aposentadoria da pessoa com deficiência, regida pela Lei Complementar n. 142/2013, também não foi atingida pela EC 103/19. Trata-se de modalidade devida à pessoa com deficiência, assim entendida aquela com impedimento de longo prazo (físico, mental, intelectual ou sensorial), que, em interação com barreiras, limita sua participação na sociedade. Subdivide-se em aposentadoria por idade e por tempo de contribuição.""",
    f"""A aposentadoria por idade da pessoa com deficiência exige idade de 60 anos para  homem e 55 anos para a mulher, com tempo de contribuição mínimo de 15 anos como deficiente em qualquer grau. Na aposentadoria por tempo de contribuição da pessoa com deficiência, o tempo exigido varia conforme o grau da deficiência, sem necessidade de idade mínima:""",
    f"""Grave: 25 anos (homem) / 20 anos (mulher)""",
    f"""Moderada: 29 anos (homem) / 24 anos (mulher)""",
    f"""Leve: 33 anos (homem) / 28 anos (mulher)""",
    f"""A carência mínima de 180 meses de contribuição, considerada desde que o segurado seja portador da deficiência, e a comprovação de que a deficiência já existia na data do requerimento ou na data em que se completaram os requisitos e persistia durante todo o período contributivo.""",
    f"""A aposentadoria da pessoa com deficiência possui regras para conversão do tempo contribuição comum e com deficiência, para efeitos de aposentadoria por tempo de contribuição, vedada a conversão para efeito de aposentadoria especial.""",
    f"""4. APOSENTADORIA ESPECIAL:""",
    f"""No que se refere à aposentadoria especial, em si, originariamente prevista na Lei nº 3.807/60, e posteriormente prevista na Lei nº 8.213/91, especialmente nos arts. 57 e 58, representa subespécie da aposentadoria por tempo de serviço (ou de contribuição), que leva em conta a realização de atividades em condições penosas, insalubres ou perigosas, potencialmente causadoras de danos à saúde ou à integridade física do trabalhador.""",
    f"""As sucessivas modificações legislativas ocorridas em relação à aposentadoria especial exigem, como premissa necessária à interpretação de seus preceitos, a fixação do entendimento de que a norma aplicável ao trabalho exercido em condições especiais é a norma vigente ao tempo em que tais atividades foram realizadas. Essa orientação tornou-se a regra do atual § 1º no artigo 70 do Regulamento da Previdência Social (Decreto n. 3.048/99). A esse respeito, o Superior Tribunal de Justiça assentou no REsp 1.151.363/MG, processado na forma do artigo 543-C do Código de Processo Civil de 1973: “observa-se o regramento da época do trabalho para a prova da exposição aos agentes agressivos à saúde: se pelo mero enquadramento da atividade nos anexos dos Regulamentos da Previdência, se mediante as anotações de formulários do INSS ou, ainda, pela existência de laudo assinado por médico do trabalho”.""",
    f"""Neste ponto, até a vigência da Lei n.º 9.032/95, para comprovação do tempo especial, bastaria a apresentação do formulário SB-40, DISES SE 5235 ou DSS 8030, preenchido pela empresa, empregador ou preposto, comprovando o enquadramento do segurado numa das atividades elencadas nas listas dos Decretos n.º 53.831/64 e 83.080/79.""",
    f"""Após a Lei n.º 9.032/95, até a publicação da medida provisória n.º 1.523, de 13 de outubro de 1996, basta apresentação dos mesmos formulários, que devem fazer menção ao agente nocivo, já que, nesta época, não mais vigia a sistemática de enquadramento em atividade profissional considerada especial, sendo necessária a comprovação de exposição do segurado aos agentes nocivos também previstos nos Decretos n.º 53.831/64 e 83.080/79.""",
    f"""Como os referidos formulários são preenchidos pelo empregador mediante declaração de responsabilidade criminal pela veracidade das informações, a este Juízo parece claro que eventuais suspeitas sobre as informações contidas no documento devem ser dirimidas pelo INSS, a tempo e modo oportuno, a fim de retirar a presunção de veracidade do documento. Com a edição do Decreto nº 4.032/2001, que determinou a redação do artigo 338, § 2º do Decreto nº 3.048/99 houve expressa previsão de fiscalização a cargo do INSS.""",
    f"""Portanto, nestes períodos não se pode exigir laudo para comprovação da exposição do segurado a agentes nocivos, pois a exigência de laudo somente teve lugar após a edição da medida provisória nº 1.523, de 13 de outubro de 1996. É anotação comum da doutrina, no entanto, que para o agente “ruído”, por imperiosa necessidade de medição, a apresentação do laudo é indispensável, qualquer que seja o período trabalhado.""",
    f"""Após 13 de outubro de 1996, por força da citada medida provisória, definitivamente convertida na Lei n.º 9.528/97, que alterou a redação do artigo 58 da Lei n.º 8.213/91, exige-se formulário emitido pela empresa ou seu preposto, com base em laudo técnico de condições ambientais do trabalho expedido por médico do trabalho ou engenheiro de segurança do trabalho, atestando a exposição aos agentes nocivos previstos nos Decretos n.º 53.831/64 e 83.080/79, e, partir de 05 de março de 1997, com base no Decreto 2.172/97, até edição do Decreto 3.048/99, que passa a embasar os enquadramentos posteriores.""",
    f"""O perfil profissiográfico mencionado pelo § 4º acrescentado ao artigo 58 da Lei nº 8.213/91 por força da medida provisória nº 1.523, de 13 de outubro de 1996, definitivamente convertida na Lei n.º 9.528/97 somente teve seu conceito introduzido pelo Decreto n.º 4.032, de 26 de novembro de 2001, a partir de quando se tornou o documento probatório da efetiva exposição dos segurados aos agentes nocivos.""",
    f"""a) PARÂMETROS TÉCNICOS E JURISPRUDENCIAIS – EPI E RUÍDO""",
    f"""Quanto ao agente nocivo ruído, consolidou-se o entendimento de que se especial a atividade sujeita ao agente ruído superior a 80 dB(A) até 05.3.1997; superior a 90 dB(A) de 06.3.1997 a 18.11.2003; superior a 85 dB(A) a partir de 19.11.2003.""",
    f"""Quanto à suposta alegação de falta de custeio para o pagamento da aposentadoria especial, é evidente que o sistema concedeu contribuições específicas para o custeio dessas aposentadorias, essencialmente o Seguro de Acidentes do Trabalho (SAT), exigido na forma do art. 22, II, da Lei nº 8.212/91, do acréscimo de que cuida o art. 57, §§ 6º e 7º, da Lei nº 8.213/91, bem como da possibilidade de redução prevista no art. 10 da Lei nº 10.666/2003.""",
    f"""Nesses termos, sem embargo da possibilidade de que a União, por meio da Secretaria da Receita Federal do Brasil, institua e cobre tais contribuições adicionais, não há como recusar o direito à aposentadoria especial a quem preencheu todos os requisitos legais.""",
    f"""Recorde-se que o Supremo Tribunal Federal, no julgamento do ARE 664.335, com repercussão geral reconhecida, fixou duas teses quanto à utilização de equipamentos de proteção individuais (EPI’s):""",
    f"""1. “O direito à aposentadoria especial pressupõe a efetiva exposição do trabalhador a agente nocivo a sua saúde, de modo que se o Equipamento de Proteção Individual (EPI) for realmente capaz de neutralizar a nocividade, não haverá respaldo à concessão constitucional de aposentadoria especial”.""",
    f"""2. “Na hipótese de exposição do trabalhador a ruído acima dos limites legais de tolerância, a declaração do empregador no âmbito do Perfil Profissiográfico Previdenciário (PPP), no sentido da eficácia do Equipamento de Proteção Individual (EPI), não descaracteriza o tempo de serviço especial para a aposentadoria”.""",
    f"""Importante destacar que para o reconhecimento de tempo especial, em relação a serviço prestado antes de 29.04.95, data da publicação da Lei n. 9.032/95, não se exige o requisito da permanência, embora seja exigível a demonstração da habitualidade na exposição a agente nocivo à saúde. A premissa reflete o entendimento da TNU (PEDILEF 200451510619827, Juíza Federal Jaqueline Michels Bilhalva, TNU - Turma Nacional de Uniformização, DJ 20/10/2008).""",
    f"""Conforme ficou decidido pela Turma Nacional de Uniformização (Pedido de Uniformização de Interpretação de Lei Federal 0501419-87.2015.4.05.8312, Juíza Federal Gisele Chaves Sampaio Alcântara, DOU 18/05/2017 pág. 99/220): “A permanência e a habitualidade da exposição a agentes nocivos à saúde são requisitos exigidos para as atividades exercidas a partir de 29/04/1995, quando entrou em vigor a Lei n. 9.032/95” – grifamos.""",
    f"""Assim, a presença do agente nocivo nas condições de trabalho, por si só, não caracteriza a atividade como especial para fins previdenciários. Além da sua presença é imprescindível que a exposição tenha ocorrido de modo habitual e permanente e que não tenha sido utilizado Equipamentos de Proteção Coletiva ou Individual realmente eficazes.""",
    f"""Em caso de não haver no PPP menção expressa à habitualidade e permanência, tal fato, por si só, não obsta o reconhecimento da especialidade. Como se sabe, o formulário é preenchido pelo empregador, motivo pelo qual o segurado não pode ser prejudicado em virtude de irregularidade formal.""",
    f"""Assim, ressalto que se as atividades descritas na profissiografia revelarem que o fator de risco se mostra inerente e indissociável às tarefas do segurado, deve-se considerá-la como permanente.""",
    f"""O PPP deverá ser assinado por representante legal da empresa, com poderes específicos outorgados por procuração, contendo a indicação dos responsáveis técnicos legalmente habilitados, por período, pelos registros ambientais e resultados de monitoração biológica.""",
    f"""A não apresentação de procuração do representante legal ou o contrato social da empresa, a meu ver, não autorizam a conclusão de que o PPP seria inidôneo. Diferente seria o caso, se se tratasse de PPP sem o responsável técnico legalmente habilitado, visto que nesse caso, é ele o engenheiro ou médico do trabalho que fará a análise do agente nocivo no ambiente laboral. Sem ele, de fato o PPP é irregular. Mas a extemporaneidade do formulário ou a ausência de procuração do representante legal que o assinou, por si só, não invalida o PPP.""",
    f"""A partir da Lei 9.032/95, editada em 28/04/1995, a ausência de responsável técnico no PPP não constitui mera irregularidade formal, visto que é o referido profissional (médico ou engenheiro do trabalho) que aferirá a presença ou não do agente nocivo no ambiente de trabalho e irá se responsabilizar pela veracidade e eficácia das suas informações. Sem o referido profissional, não há como se reconhecer a especialidade por agente nocivo.""",
    f"""De todo modo, saliente-se que a ausência de indicação de responsável técnico no PPP poderá ser suprida pela juntada do Laudo Técnico de condições ambientais do trabalho expedido por médico do trabalho ou engenheiro de segurança do trabalho, que deu fundamento às anotações dos fatores de risco.""",
    f"""Quanto à extemporaneidade do laudo, a TNU consolidou a controvérsia por meio da Súmula nº 68: “O laudo pericial não contemporâneo ao período trabalhado é apto à comprovação da atividade especial do segurado”.""",
    f"""Do mesmo modo, o fato do responsável técnico ter sido contratado em período posterior ao que o segurado exerceu suas atividades laborais na empresa, também não invalida o referido laudo.""",
    f"""Por fim, é possível a conversão do tempo especial em tempo comum, para efeito de concessão de aposentadoria por tempo de contribuição. A conversão, no entanto, tem por termo final a EC 103/19, a partir de quando se tornou vedado qualquer acréscimo de tempo fictício.""",
    f"""Feitas estas premissas, passo a apreciar o caso concreto."""
     ]

else:
    fundamento_custom = st.text_area(
        "Redija, ou copie e cole, a fundamentação que deseja inserir na sentença. \nO texto deve englobar tudo, desde o 'vistos em sentença' até um parágrafo assim redigido: 'Feitas estas considerações, passo a analisar o caso concreto', ou expressão equivalente."
        )


st.subheader("Fábrica de Cálculos - TRF3")
st.write("Acesse a Fábrica de Cálculos (TRF3), em https://www.trf3.jus.br/cecalc/tc/ e siga os passos para realizar o cálculo de tempo de contribuição. Após, salve seu cálculo, em seu computador, em arquivo com qualquer nome, com extensão .calc. Faça o upload o arquivo que você salvou, no campo abaixo")

uploaded_file = st.file_uploader("Anexe aqui o arquivo .calc que você salvou em seu computador", type=["calc"], key="uploader")

if uploaded_file:
    try:
        payload = json.load(uploaded_file)
    except json.JSONDecodeError:
        st.error("Erro ao ler o arquivo .calc.")
        st.stop()

    payload = parse_calc_format(payload)
    payload_final = montar_payload_final(payload)
    #st.json(payload_final, expanded=False)

    # Envio e tratamento da resposta
    if payload_final:
        st.session_state["payload_final"] = payload_final
        with st.spinner("Enviando requisição para a API..."):
            response = requests.post(API_URL, json=payload_final)

        if response.status_code == 200:
            st.success("Resposta recebida com sucesso!")
            resposta = response.json()
            st.session_state["resposta"] = resposta
            st.json(resposta, expanded=False)

            # ESCOLHA DOS VINCULOS QUE SERÃO ANALISADOS
            # Inicializa estruturas de estado
            if "opcoes_periodos" not in st.session_state:
                st.session_state["opcoes_periodos"] = []
            if "mapa_periodos" not in st.session_state:
                st.session_state["mapa_periodos"] = {}

            st.subheader("🎯 Escolha dos vínculos que serão analisados")
            with st.expander("📖 Instruções ao usuário"):
                st.markdown(f"""**Toda a lógica do Aplicativo é justificar as escolhas que você fez na hora que elaborou o cálculo na "Fábrica de Cálculos"**.
                            \nOs vínculos são apresentados **conforme foram inseridos na "Fábrica de Cálculos"**.
                            \nPor este motivo, há casos em que o período foi recortado. O recorte ocorre porque o vínculo refere-se a um período maior, mas somente foi reconhecida parte da pretensão, ou, então, porque a "Fábrica de Cálculos" o recorta automaticamente, nos marcos da Lei 9.876/99, EC 103/19, por exemplo.
                            \nNestes casos, selecione apenas a parte do vínculo recortada, que vai corresponder ao período que você vai reconhecer. Você será chamado a especificar, posteriormente, que se trata de parte de um período controvertido maior, e que somente esta parte será deferida, e o restante do pedido não, de modo que a redação da sentença sairá ajustada.
                            \nSe preferir, nos casos em que há vários recortes do mesmo período, você pode inseri-lo manualmente. O valor inserido manualmente será o que período apreciado. Se você preferir analisar parte de um período dentro de um período maior, você será chamado a especificar o período maior.
                            \nO mesmo ocorre em relação ao reconhecimento de deficiência. Somente será exibido o período que foi deferido, embora a controvérsia possa abranger período maior. Você será chamado a especificar.
                            \nNeste lógica, você perceberá que os pedidos de reconhecimento de tempo (urbano ou rural, especial ou comum) que serão julgados **totalmente improcedentes**, não estão listados.
                            \nIsso ocorre porque você não os inseriu no "Fábrica de Cálculos" (corretamente, porque serão improcedentes no total). Contudo, estes períodos devem compor a fundamentação da sentença.
                            \nComo os períodos não constam, sequer em parte, no cálculo advindo da "Fábrica de Cálculos", eles devem ser inseridos manualmente abaixo.
                            \nIgualmente, se a parte alegou ter trabalhado como DEFICIENTE, e isso não será reconhecido pela sentença em nenhum grau, o período respectivo também deve ser inserido manualmente.""")

            # Adiciona períodos da API nas estruturas de estado
            for p in resposta.get("dadosPeriodos", []):
                rotulo = f"{formatar_data_iso_para_br(p.get('dataInicial'))} a {formatar_data_iso_para_br(p.get('dataFinal'))} - {p.get('descricao')}"
                if rotulo not in st.session_state.opcoes_periodos:
                    st.session_state.opcoes_periodos.append(rotulo)
                    st.session_state.mapa_periodos[rotulo] = {
                        "data_inicio": formatar_data_iso_para_br(p.get("dataInicial", "")),
                        "data_fim": formatar_data_iso_para_br(p.get("dataFinal", "")),
                        "vinculo": p.get("descricao", ""),
                        "data_inicio_maior": "",
                        "data_fim_maior": "",
                        "origem_do_dado": "controvertido_sem_deficiencia",
                        "grau_deficiencia": "",
                        "o_que_parte_pede": "",
                        "houve_prova_material": "",
                        "documento_prova_material": "",
                        "conclusao_prova_material": "",
                        "houve_prova_testemunhal": "",
                        "depoimento": "",
                        "conclusao_depoimento": "",
                        "tipo_tempo": "",
                        "resultado": "",
                        "dispositivo": ""
                    }
            for p in payload_final.get("periodosPcd", []):
                rotulo = f"{p.get('dataInicial', '')} a {p.get('dataFinal', '')} - Deficiência {p.get('grauDeficiencia', '')}"
                if rotulo not in st.session_state.opcoes_periodos:
                    st.session_state.opcoes_periodos.append(rotulo)
                    st.session_state.mapa_periodos[rotulo] = {
                        "data_inicio": p.get("dataInicial", ""),
                        "data_fim": p.get("dataFinal", ""),
                        "vinculo": f"Deficiência {p.get('grauDeficiencia')}",
                        "data_inicio_maior": "",
                        "data_fim_maior": "",
                        "origem_do_dado": "controvertido_deficiencia",
                        "grau_deficiencia": p.get("grauDeficiencia"),
                        "o_que_parte_pede": "",
                        "houve_prova_material": "",
                        "documento_prova_material": "",
                        "conclusao_prova_material": "",
                        "houve_prova_testemunhal": "",
                        "depoimento": "",
                        "conclusao_depoimento": "",
                        "conclusao_especial_ou_comum": "",
                        "tipo_tempo": "",
                        "resultado": "",
                        "texto_final_periodos": "",
                        "dispositivo": ""
                    }

            # Inicializa listas no session_state, que resultarão na lista "periodos_para_sentenca"
            if "periodos_manuais" not in st.session_state:
                st.session_state["periodos_manuais"] = []
            if "periodos_da_api" not in st.session_state:
                st.session_state["periodos_da_api"] = []

            abas = st.tabs(["🗂 Escolha dos Períodos existentes na Fábrica de Cálculos", "✏️ Inserção manual de período"])

            with abas[0]:
                # Multiselect para seleção dos períodos analisados
                escolhidos = st.multiselect("Escolha o(s) período(s) que deve(m) ser analisado(s) nesta sentença:",
                    options=st.session_state.opcoes_periodos,
                    key="periodos_controvertidos_selecionados")

                st.session_state["periodos_da_api"] = [st.session_state.mapa_periodos[r] for r in st.session_state["periodos_controvertidos_selecionados"]]

            with abas[1]:
                st.write("Haverá necessidade de inserção manual de algum período?")
                contador = [0]
                id_contador = len(contador) + 1
                if id_contador <= 100:
                    tipo = st.radio("Trata-se de tempo de serviço ou de período de deficiência", [1, 2],
                                    format_func=lambda x: "Tempo de serviço" if x == 1 else "Deficiência", key=f"tipo_{id_contador}")
                    if tipo == 1:
                        origem_dado = "controvertido_sem_deficiencia"
                        # vinculo_inserido = st.text_input("Descrição do vínculo (empregador / contribuinte / etc.)", key=f"vinculo_inserido_{id_contador}")
                        data_inicio_inserido = st.text_input("Data inicial (dd/mm/aaaa)", key=f"data_inicio_inserido_{id_contador}")
                        data_fim_inserido = st.text_input("Data final (dd/mm/aaaa)", key=f"data_fim_inserido_{id_contador}")
                    else:
                        origem_dado = "controvertido_deficiencia"
                        vinculo_inserido = "Deficiência"
                        data_inicio_inserido = st.text_input("Data inicial da deficiência (dd/mm/aaaa)", key=f"data_inicio_inserido_{id_contador}")
                        data_fim_inserido = st.text_input("Data final da deficiência (dd/mm/aaaa)", key=f"data_fim_inserido_{id_contador}")

                    if data_inicio_inserido and data_fim_inserido:
                        if st.button("Salvar período"):
                            novo_periodo = {
                                "data_inicio": data_inicio_inserido,
                                "data_fim": data_fim_inserido,
                                "vinculo": vinculo_inserido,
                                "data_inicio_maior": "",
                                "data_fim_maior": "",
                                "origem_do_dado": origem_dado,
                                "grau_deficiencia": "",
                                "o_que_parte_pede": "",
                                "houve_prova_material": "",
                                "documento_prova_material": "",
                                "conclusao_prova_material": "",
                                "houve_prova_testemunhal": "",
                                "depoimento": "",
                                "conclusao_depoimento": "",
                                "conclusao_especial_ou_comum": "",
                                "tipo_tempo": "",
                                "resultado": "",
                                "texto_final_periodos": "",
                                "dispositivo": ""
                            }
                            st.session_state.periodos_manuais.append(novo_periodo)
                            st.success("Período manual inserido com sucesso.")
                            contador.append(len(contador)+1)

                        # Permite inserir outro
                        if st.button("Inserir outro período"):
                            st.rerun()
            
            # Junta os períodos da API + manuais para análise/sentença
            periodos_para_sentenca = st.session_state["periodos_da_api"] + st.session_state["periodos_manuais"]

            # LÓGICA DE APRECIAÇÃO DE CADA PERÍODO
            if periodos_para_sentenca:
                abasx = st.tabs(["🔍 Análise de cada período", "☢️ Exemplos de redações para Agentes Nocivos"])
                with abasx[0]:
                    st.subheader("🔍 Análise de cada período")
                    st.write("Cada período deverá ser analisado separadamente, dentro de cada pasta abaixo.")
                    for i, p in enumerate(periodos_para_sentenca):
                        with st.expander(f"De {p['data_inicio']} a {p['data_fim']}"):
                            if p["origem_do_dado"] == "controvertido_sem_deficiencia":
                                p["vinculo"] = st.text_input("Qual a descrição do vínculo (nome do empregador ou contratante / contribuinte individual / rural / etc.)", key=f"vinculo_{i}")
                                periodo_maior = st.radio("Este período analisado é todo o período controvertido pela parte autora, ou é apenas parte de um período maior controvertido pela parte autora?",[1,2], format_func=lambda x: ("O período analisado é parte de um período controvertido maior" if x == 1 else "O período analisado é toda o período controvertido pela parte autora"), key=f"periodo_maior_{i}", index=1)
                                if periodo_maior == 1:
                                    p["data_inicio_maior"] = st.text_input("Qual a data inicial do período maior controvertido? Digite em formato dd/mm/aaaa.", key=f"data_inicio_maior_{i}")
                                    p["data_fim_maior"] = st.text_input("Qual a data final do período maior controvertido? Digite em formato dd/mm/aaaa.", key=f"data_fim_maior_{i}")
                                else: 
                                    p["data_inicio_maior"] = p["data_inicio"]
                                    p["data_fim_maior"] = p["data_fim"]
                                o_que_parte_pede = st.radio(f"O que a parte pede para este período de {p['data_inicio_maior']} ate {p['data_fim_maior']}?", [1, 2, 3, 4], format_func=lambda x: (
                                        "Reconhecimento de tempo urbano comum" if x == 1 else
                                        "Reconhecimento de tempo urbano comum cumulado com pedido de conversão do período para tempo especial" if x == 2 else
                                        "Reconhecimento de tempo rural" if x == 3 else
                                        "Apenas conversão de tempo especial para comum"
                                    ), key=f"o_que_parte_pede_{i}")
    
                                p["o_que_parte_pede"] = {
                                    1: "reconhecimento de tempo urbano comum",
                                    2: "reconhecimento de tempo urbano cumulado com pedido de conversão do período para tempo especial",
                                    3: "reconhecimento de tempo rural",
                                    4: "conversão de tempo especial para comum"
                                }[o_que_parte_pede]
                                
                                if o_que_parte_pede in [1, 2, 3]:
                                    aplicacao_sumula_75_reconhecimento_CTPS = []
                                    aplicacao_sumula_75_negativa_CTPS = []                                    
                                    if o_que_parte_pede in [1, 2]:
                                        sumula_75_TNU = st.radio(f"Responda, quanto à Súmula 75 da TNU: A Carteira de Trabalho e Previdência Social (CTPS) em relação à qual não se aponta defeito formal que lhe comprometa a fidedignidade goza de presunção relativa de veracidade, formando prova suficiente de tempo de serviço para fins previdenciários, ainda que a anotação de vínculo de emprego não conste no Cadastro Nacional de Informações Sociais (CNIS).", [1, 2, 3], format_func=lambda x: "Vou aplicar a súmula e julgar procedente o pedido em relação a este período, pois a CTPS está formalmente em ordem" if x == 1 else "Vou aplicar a súmula, mas por algum motivo, a CTPS não serve como prova definitiva e nem como início de prova material" if x == 2 else "Não vou aplicar a súmula, pois não há registro em CTPS", key=f"sumula_75_TNU_{i}")
                                        if sumula_75_TNU == 1:
                                            aplicacao_sumula_75_reconhecimento_CTPS.extend([f"A súmula 75 da TNU dispõe: {"A Carteira de Trabalho e Previdência Social (CTPS) em relação à qual não se aponta defeito formal que lhe comprometa a fidedignidade goza de presunção relativa de veracidade, formando prova suficiente de tempo de serviço para fins previdenciários, ainda que a anotação de vínculo de emprego não conste no Cadastro Nacional de Informações Sociais (CNIS)."}",
                                                                        f"Não se justifica que este período não seja computado para cálculo do tempo de contribuição da parte autora."
                                                                         ])
                                        if sumula_75_TNU == 2:
                                            motivo_CTPS = [p.strip() for p in st.text_area(f"Redija porque a CTPS não serve como prova para fins da aplicação da súmula 75 da TNU. Inicie com letra maiúscula e encerre com ponto final.", key=f"motivo_CTPS_{i}").split("\n") if p.strip()]
                                            aplicacao_sumula_75_negativa_CTPS.extend(motivo_CTPS)

                                    if aplicacao_sumula_75_reconhecimento_CTPS:
                                        p["texto_final_periodos"] = []
                                        p["texto_final_periodos"].append(f"DO PERÍODO ENTRE {p["data_inicio_maior"]} ATÉ {p["data_fim_maior"]} - {p["vinculo"]}:")
                                        p["texto_final_periodos"].append(f"Em relação ao tempo de trabalho de {p["data_inicio_maior"]} até {p["data_fim_maior"]} - {p["vinculo"]} - em que a parte autora pede {p["o_que_parte_pede"]}, foi apresentada CTPS formalmente em ordem, com referido vínculo anotado.")
                                        p["texto_final_periodos"].extend(aplicacao_sumula_75_reconhecimento_CTPS)
                                        
                                        if o_que_parte_pede == 1:
                                            p["tipo_tempo"] = "comum" 
                                            p["resultado"] = "Procedente" if periodo_maior == 2 else "Procedente em parte" 
                                            p["dispositivo"] = f"Com resolução de mérito, nos termos do art. 487, I do CPC, em relação ao pedido de {p['o_que_parte_pede']} referente ao período de {p['data_inicio_maior']} a {p['data_fim_maior']} - {p['vinculo']}, JULGO {p['resultado'].upper()} o pedido e declaro o período de {p['data_inicio']} a {p['data_fim']} como tempo comum, determinando sua averbação."                                                                                
    
                                        if o_que_parte_pede == 2:
                                            sera_reconhecido_comum_especial = st.radio(f"Dentro do período de {p['data_inicio_maior']} até {p['data_fim_maior']}, a sentença reconhecerá o período de {p['data_inicio']} ate {p['data_fim']} como tempo comum ou tempo especial?", [1, 2], format_func=lambda x: "Comum" if x == 1 else "Especial", key=f"sera_reconhecido_comum_especial_{i}")
                                            p["conclusao_especial_ou_comum"] = [p.strip() for p in st.text_area(f"Redija porque chegou nesta conclusão, com base nas provas. Inicie com letra maiúscula e encerre com ponto final.", key=f"conclusao_comum_especial{i}").split("\n") if p.strip()]
                                            p["tipo_tempo"] = "comum" if sera_reconhecido_comum_especial == 1 else "especial"
                                            p["resultado"] = "Procedente" if sera_reconhecido_comum_especial == 2 and periodo_maior == 2 else "Procedente em parte"
                                            p["dispositivo"] = f"Com resolução de mérito, nos termos do art. 487, I do CPC, em relação ao pedido de {p['o_que_parte_pede']} referente ao período de {p['data_inicio_maior']} a {p['data_fim_maior']} - {p['vinculo']}, JULGO {p['resultado'].upper()} o pedido " + (f"e declaro o período entre {p['data_inicio']} até {p['data_fim']} como tempo {p['tipo_tempo']}, determinando sua averbação." if p["tipo_tempo"] == "comum" else f"e declaro o período entre {p['data_inicio']} até {p['data_fim']}, como tempo {p['tipo_tempo']}, sujeito a conversão, determinando sua averbação." if p["tipo_tempo"] == "especial" else "")
                                            p["texto_final_periodos"].append(f"Passo a analisar a alegação de que o tempo de trabalho é tempo especial.")
                                            p["texto_final_periodos"].extend(p['conclusao_especial_ou_comum'])

                                    else:  
                                        precisa_prova = st.radio(f"A parte trouxe algum início de prova material para este período de {p['data_inicio_maior']} ate {p['data_fim_maior']}? (responda {"Sim"} se a parte autora trouxe algum documento, mesmo que não seja inicío de prova suficiente. Somente responda {"Não"} se nenhum documento foi apresentado).", [1, 2], format_func=lambda x: "Sim" if x == 1 else "Não (pedido será extinto - Tema 629 STJ)", key=f"precisa_prova_{i}")
        
                                        if precisa_prova == 2:
                                            p["houve_prova_material"] = f"Não"
                                            p["resultado"] = f"Extinto sem julgamento de mérito - Tema 629 - Falta de início de prova material"
                                            p["dispositivo"] = f"Sem resolução de mérito, nos termos do art. 485, VI do CPC, JULGO EXTINTO o pedido de {p['o_que_parte_pede']} referente ao período de {p['data_inicio_maior']} a {p['data_fim_maior']} - {p['vinculo']}."
                                            p["texto_final_periodos"] = []
                                            p["texto_final_periodos"].append(f"DO PERÍODO ENTRE {p["data_inicio_maior"]} ATÉ {p["data_fim_maior"]} - {p["vinculo"]}:")
                                            p["texto_final_periodos"].append(f"Em relação ao tempo de trabalho de {p["data_inicio_maior"]} até {p["data_fim_maior"]} - {p["vinculo"]} - em que a parte autora pede {p["o_que_parte_pede"]}, o art. 55, § 3º da Lei n. 8.213/91 exige a apresentação de início de prova material, sendo ônus da parte autora junta-lo aos autos. No caso, a parte não trouxe nenhum documento como prova material.")
                                            if aplicacao_sumula_75_negativa_CTPS:
                                                p["texto_final_periodos"].extend([f"Quanto à CTPS apresentada, não se presta para fins da súmula 75 da TNU, ou como início de prova material. {aplicacao_sumula_75_negativa_CTPS[0]}"] + aplicacao_sumula_75_negativa_CTPS[1:])
                                            p["texto_final_periodos"].append(f"Ausente início de prova material, inviável o acolhimento do pleito. Em que pese, em regra, a ausência de prova implique na improcedência do feito, a jurisprudência do STJ consolidou-se no sentido de que, nas demandas previdenciárias, a ausência de prova de tempo de contribuição deve resultar na extinção do feito sem resolução de mérito, por ausência de pressuposto processual, diante do caráter social que a lide envolve. Trata-se do Resp 1352721/SP, Rel. Ministro NAPOLEÃO NUNES MAIA FILHO, julgado em regime de recursos repetitivos (tema 629) - REsp 1352721/SP, Rel. Ministro NAPOLEÃO NUNES MAIA FILHO, CORTE ESPECIAL, julgado em 16/12/2015, DJe 28/04/2016.")
                                                        
        
                                        if precisa_prova == 1:
                                            p["houve_prova_material"] = "Sim"
                                            p["documento_prova_material"] = [p.strip() for p in st.text_area(f"Qual(is) documento(s) a parte trouxe, como início de prova material? Especifique-os, mas redija como um (ou mais) parágrafo(s) completo(s), iniciando com letra maiúscula, e encerrando com ponto final", key=f"doc_prova_{i}").split("\n") if p.strip()]
                                            inicio_prova_material_apresentado = st.radio("O início de prova material apresentado é suficiente?", [1, 2], format_func=lambda x: "Sim" if x == 1 else "Não (pedido será extinto - Tema 629 STJ)", key=f"inicio_prova_material_suficiente_{i}")
                                            if inicio_prova_material_apresentado == 2:
                                                p["conclusao_prova_material"] = [p.strip() for p in st.text_area(f"Explique o motivo pelo qual você concluiu que os documentos apresentados não são suficientes para início de prova material. Redija como um (ou mais) parágrafo(s) completo(s), iniciando com letra maiúscula, e encerrando com ponto final", key=f"conclusao_prova_material_{i}").split("\n") if p.strip()]
                                                p["resultado"] = f"Extinto sem julgamento de mérito - Tema 629 - Falta de início de prova material"
                                                p["dispositivo"] = f"Sem resolução de mérito, nos termos do art. 485, VI do CPC, JULGO EXTINTO o pedido de {p['o_que_parte_pede']} referente ao período de {p['data_inicio_maior']} a {p['data_fim_maior']} - {p['vinculo']}."
                                                p["texto_final_periodos"] = []
                                                p["texto_final_periodos"].append(f"DO PERÍODO ENTRE {p["data_inicio_maior"]} ATÉ {p["data_fim_maior"]} - {p["vinculo"]}:")
                                                p["texto_final_periodos"].append(f"Em relação ao tempo de trabalho de {p["data_inicio_maior"]} até {p["data_fim_maior"]} - {p["vinculo"]} - em que a parte autora pede {p["o_que_parte_pede"]}, o art. 55, § 3º da Lei n. 8.213/91 exige a apresentação de início de prova material, sendo ônus da parte autora junta-lo aos autos.")
                                                p["texto_final_periodos"].extend(p["documento_prova_material"])
                                                p["texto_final_periodos"].extend(p["conclusao_prova_material"])
                                                if aplicacao_sumula_75_negativa_CTPS:
                                                    p["texto_final_periodos"].extend([f"Quanto à CTPS apresentada, não se presta para fins da súmula 75 da TNU, ou como início de prova material. {aplicacao_sumula_75_negativa_CTPS[0]}"] + aplicacao_sumula_75_negativa_CTPS[1:])
                                                p["texto_final_periodos"].append(f"Ausente início de prova material, inviável o acolhimento do pleito. Em que pese, em regra, a ausência de prova implique na improcedência do feito, a jurisprudência do STJ consolidou-se no sentido de que, nas demandas previdenciárias, a ausência de prova de tempo de contribuição deve resultar na extinção do feito sem resolução de mérito, por ausência de pressuposto processual, diante do caráter social que a lide envolve. Trata-se do Resp 1352721/SP, Rel. Ministro NAPOLEÃO NUNES MAIA FILHO, julgado em regime de recursos repetitivos (tema 629) - REsp 1352721/SP, Rel. Ministro NAPOLEÃO NUNES MAIA FILHO, CORTE ESPECIAL, julgado em 16/12/2015, DJe 28/04/2016.")
                                                                        
                                       
                                            if inicio_prova_material_apresentado == 1:
                                                especificar_conclusao_prova_material = st.radio(f"A sentença informará genericamente que: 'Considero suficiente o início de prova material apresentado nos autos autos'. Deseja apresentar outros esclarecimentos além disso, ou esta redação é suficiente?", [1,2], format_func=lambda x: "Esta redação é suficiente" if x == 1 else "Desejo apresentar esclarecimentos", key=f"especificar_conclusao_prova_material_{i}")
                                                if especificar_conclusao_prova_material == 1:
                                                    p["conclusao_prova_material"] = []
                                                if especificar_conclusao_prova_material == 2:
                                                    p["conclusao_prova_material"] = [p.strip() for p in st.text_area(f"Explique porque você concluiu que os documentos apresentados são suficientes para início de prova material. Redija como um (ou mais) parágrafo(s) completo(s), iniciando com letra maiúscula, e encerrando com ponto final", key=f"conclusao_prova_material_{i}").split("\n") if p.strip()]
                                                testemunhal = st.radio(f"Houve prova testemunhal para este período?", [1, 2], format_func=lambda x: "Sim" if x == 1 else "Não", key=f"testemunhal_{i}")
                                                p["houve_prova_testemunhal"] = "Sim" if testemunhal == 1 else "Não"
        
                                                if testemunhal == 2:
                                                    p["resultado"] = "Improcedente"
                                                    p["dispositivo"] = f"Com resolução de mérito, nos termos do art. 487, I do CPC, JULGO IMPROCEDENTE o pedido de {p['o_que_parte_pede']} referente ao período de {p['data_inicio_maior']} a {p['data_fim_maior']} - {p['vinculo']}."
                                                    p["texto_final_periodos"] = []
                                                    p["texto_final_periodos"].append(f"DO PERÍODO ENTRE {p["data_inicio_maior"]} ATÉ {p["data_fim_maior"]} - {p["vinculo"]}:")
                                                    if aplicacao_sumula_75_negativa_CTPS:
                                                        p["texto_final_periodos"].extend([f"Quanto à CTPS apresentada, não se presta para fins da súmula 75 da TNU, ou como início de prova material. {aplicacao_sumula_75_negativa_CTPS[0]}"] + aplicacao_sumula_75_negativa_CTPS[1:])
                                                    p["texto_final_periodos"].append(f"Em relação ao tempo de trabalho de {p["data_inicio_maior"]} até {p["data_fim_maior"]} - {p["vinculo"]} - em que a parte autora pede {p["o_que_parte_pede"]}, o art. 55, § 3º da Lei n. 8.213/91 exige a apresentação de início de prova material, que, por si só, não é suficiente para reconhecimento da pretensão. É indispensável a oitiva da testemunhas para confirmar a existência do tempo de trabalho.")
                                                    p["texto_final_periodos"].append(f"O ônus da prova é da parte autora, nos termos da legislação processual. Embora tenha apresentado início de prova material, não foram indicadas testemunhas para comprovação do alegado, em audiência. A falta de prova implica em improcedência.")
                                                                                
        
                                                if testemunhal == 1:
                                                    p["texto_final_periodos"] = []
                                                    p["texto_final_periodos"].append(f"DO PERÍODO ENTRE {p["data_inicio_maior"]} ATÉ {p["data_fim_maior"]} - {p["vinculo"]}:")
                                                    if aplicacao_sumula_75_negativa_CTPS:
                                                        p["texto_final_periodos"].extend([f"Quanto à CTPS apresentada, não se presta para fins da súmula 75 da TNU, ou como início de prova material. {aplicacao_sumula_75_negativa_CTPS[0]}"] + aplicacao_sumula_75_negativa_CTPS[1:])
                                                    p["texto_final_periodos"].append(f"Em relação ao tempo de trabalho de {p["data_inicio_maior"]} até {p["data_fim_maior"]} - {p["vinculo"]} - em que a parte autora pede {p["o_que_parte_pede"]}, o art. 55, § 3º da Lei n. 8.213/91 exige a apresentação de início de prova material, para reconhecimento do pedido. No caso dos autos, houve início de prova material.")
                                                    p["texto_final_periodos"].extend(p['documento_prova_material'])
                                                    p["texto_final_periodos"].extend([f"Considero suficiente o início de prova material apresentado nos autos autos. {p['conclusao_prova_material'][0]}"] + p["conclusao_prova_material"][1:] if p["conclusao_prova_material"] else ["Considero suficiente o início de prova material apresentado nos autos autos."])
                                                    p["texto_final_periodos"].append(f"Para comprovação do alegado, houve oitiva de testemunha(s) em Juízo:")
                                                    p["texto_final_periodos"].extend(p['depoimento'])
                                                    p["texto_final_periodos"].extend(p['conclusao_depoimento'])
                                                    
                                                    p["depoimento"] = [p.strip() for p in st.text_area("Redija o(s) depoimento(s). Inicie com letra maiúscula, e encerre com ponto final", key=f"depoimento_{i}").split("\n") if p.strip()]
                                                    if o_que_parte_pede == 1:
                                                        sera_reconhecido = st.radio(f"Dentro do período de {p['data_inicio_maior']} até {p['data_fim_maior']}, a sentença reconhecerá o período de {p['data_inicio']} até {p['data_fim']} como tempo de serviço?", [1, 2], format_func=lambda x: "Sim" if x == 1 else "Não", key=f"sera_reconhecido_{i}")
                                                        p["conclusao_depoimento"] = [p.strip() for p in st.text_area(f"Redija porque chegou nesta nesta conclusão, com base nos depoimentos colhidos e demais provas. Inicie com letra maiúscula e encerre com ponto final.", key=f"conclusao_depoimento_{i}").split("\n") if p.strip()]
                                                        p["tipo_tempo"] = "comum" if sera_reconhecido == 1 else ""
                                                        p["resultado"] = "Procedente" if sera_reconhecido == 1 and periodo_maior == 2 else "Procedente em parte" if sera_reconhecido == 1 and periodo_maior == 1 else "Improcedente"
                                                        p["dispositivo"] = f"Com resolução de mérito, nos termos do art. 487, I do CPC, JULGO IMPROCEDENTE o pedido de {p['o_que_parte_pede']} referente ao período de {p['data_inicio_maior']} a {p['data_fim_maior']} - {p['vinculo']}." if p.get("resultado") == "Improcedente" else f"Com resolução de mérito, nos termos do art. 487, I do CPC, em relação ao pedido de {p['o_que_parte_pede']} referente ao período de {p['data_inicio_maior']} a {p['data_fim_maior']} - {p['vinculo']}, JULGO {p['resultado'].upper()} o pedido " + (f"e declaro o período de {p['data_inicio']} a {p['data_fim']} como tempo comum, determinando sua averbação." if p["tipo_tempo"] == "comum" else f"")                                                                                
        
                                                    if o_que_parte_pede == 2:
                                                        sera_reconhecido = st.radio(f"Dentro do período de {p['data_inicio_maior']} até {p['data_fim_maior']}, a sentença reconhecerá o período de {p['data_inicio']} até {['data_fim']} como tempo tempo de serviço?", [1, 2], format_func=lambda x: "Sim" if x == 1 else "Não", key=f"sera_reconhecido_{i}")                          
                                                        p["conclusao_depoimento"] = [p.strip() for p in st.text_area(f"Redija porque chegou nesta conclusão, com base nos depoimentos colhidos e demais provas. Inicie com letra maiúscula e encerre com ponto final.", key=f"conclusao_depoimento_{i}").split("\n") if p.strip()]
                                                        if sera_reconhecido == 2:
                                                            p["resultado"] = "Improcedente"
                                                            p["dispositivo"] = f"Com resolução de mérito, nos termos do art. 487, I do CPC, JULGO IMPROCEDENTE o pedido de {p['o_que_parte_pede']} referente ao período de {p['data_inicio_maior']} a {p['data_fim_maior']} - {p['vinculo']}."
                                                            p["texto_final_periodos"].append(f"A parte autora não comprovou o alegado, e, por isso, seu pedido de reconhecimento do período em questão deve ser improcedente")
                                                                                    
                                                        if sera_reconhecido == 1:
                                                            sera_reconhecido_comum_especial = st.radio(f"Dentro do período de {p['data_inicio_maior']} até {p['data_fim_maior']}, a sentença reconhecerá o período de {p['data_inicio']} ate {p['data_fim']} como tempo comum ou tempo especial?", [1, 2], format_func=lambda x: "Comum" if x == 1 else "Especial", key=f"sera_reconhecido_comum_especial_{i}")
                                                            p["conclusao_especial_ou_comum"] = [p.strip() for p in st.text_area(f"Redija porque chegou nesta conclusão, com base nas provas. Inicie com letra maiúscula e encerre com ponto final.", key=f"conclusao_comum_especial{i}").split("\n") if p.strip()]
                                                            p["tipo_tempo"] = "comum" if sera_reconhecido_comum_especial == 1 else "especial"
                                                            p["resultado"] = "Procedente" if sera_reconhecido_comum_especial == 2 and periodo_maior == 2 else "Procedente em parte"
                                                            p["dispositivo"] = f"Com resolução de mérito, nos termos do art. 487, I do CPC, em relação ao pedido de {p['o_que_parte_pede']} referente ao período de {p['data_inicio_maior']} a {p['data_fim_maior']} - {p['vinculo']}, JULGO {p['resultado'].upper()} o pedido " + (f"e declaro o período entre {p['data_inicio']} até {p['data_fim']} como tempo {p['tipo_tempo']}, determinando sua averbação." if p["tipo_tempo"] == "comum" else f"e declaro o período entre {p['data_inicio']} até {p['data_fim']}, como tempo {p['tipo_tempo']}, sujeito a conversão, determinando sua averbação." if p["tipo_tempo"] == "especial" else "")
                                                            p["texto_final_periodos"].append(f"Passo a analisar a alegação de que o tempo de trabalho é tempo especial.")
                                                            p["texto_final_periodos"].extend(p['conclusao_especial_ou_comum'])
                                                                                        
        
                                                    if o_que_parte_pede == 3:
                                                        sera_reconhecido = st.radio(f"Dentro do período de {p['data_inicio_maior']} até {p['data_fim_maior']}, a sentença reconhecerá o período de {p['data_inicio']} até {p['data_fim']} como tempo rural ou não reconhecerá o tempo de trabalho deste período?", [1, 2], format_func=lambda x: "Reconhecerá o tempo rural" if x == 1 else "Não reconhecerá",key=f"sera_reconhecido_{i}")
                                                        p["conclusao_depoimento"] = [p.strip() for p in st.text_area(f"Redija porque chegou nesta conclusão, com base nos depoimentos colhidos e demais provas. Inicie com letra maiúscula e encerre com ponto final.", key=f"conclusao_depoimento_{i}").split("\n") if p.strip()]
                                                        if sera_reconhecido == 2:
                                                            p["resultado"] = "Improcedente"
                                                            p["dispositivo"] = f"Com resolução de mérito, nos termos do art. 487, I do CPC, JULGO IMPROCEDENTE o pedido de {p['o_que_parte_pede']} referente ao tempo rural de {p['data_inicio_maior']} a {p['data_fim_maior']}."
                                                            p["texto_final_periodos"].append(f"A parte autora não comprovou o alegado, e, por isso, seu pedido de reconhecimento do período em questão deve ser improcedente.")
                                                                                    
        
                                                        if sera_reconhecido == 1:                                            
                                                            p["tipo_tempo"] = "rural"
                                                            p["resultado"] = "Procedente" if periodo_maior == 2 else "Procedente em parte"
                                                            p["dispositivo"] = f"Com resolução de mérito, nos termos do art. 487, I do CPC, em relação ao pedido de {p['o_que_parte_pede']} referente ao período de {p['data_inicio_maior']} a {p['data_fim_maior']}, JULGO {p['resultado'].upper()} o pedido e declaro o período entre {p["data_inicio"]} até {p["data_fim"]} como tempo rural, determinando sua averbação."
        
    
                                if o_que_parte_pede == 4:
                                    sera_reconhecido = st.radio(f"Dentro do período de {p['data_inicio_maior']} até {p['data_fim_maior']}, a sentença reconhecerá o período de {p['data_inicio']} até {p['data_fim']} como tempo comum ou tempo especial?", [1, 2],
                                                                format_func=lambda x: "Comum" if x == 1 else "Especial",
                                                                key=f"sera_reconhecido_{i}")
                                    p["conclusao_especial_ou_comum"] = [p.strip() for p in st.text_area(f"Redija porque chegou nesta conclusão, com base nas provas. Inicie com letra maiúscula e encerre com ponto final.", key=f"conclusao_comum_especial{i}").split("\n") if p.strip()]
                                    p["tipo_tempo"] = "comum" if sera_reconhecido == 1 else "especial"
                                    p["resultado"] = "Procedente" if sera_reconhecido == 2 and periodo_maior == 2 else "Procedente em parte" if sera_reconhecido == 2 and periodo_maior == 1 else "Improcedente"
                                    p["dispositivo"] = f"Com resolução de mérito, nos termos do art. 487, I do CPC, JULGO IMPROCEDENTE o pedido de {p['o_que_parte_pede']} referente ao período de {p['data_inicio_maior']} a {p['data_fim_maior']} - {p['vinculo']}." if p.get("resultado") == "Improcedente" else f"Com resolução de mérito, nos termos do art. 487, I do CPC, em relação ao pedido de {p['o_que_parte_pede']} referente ao período de {p['data_inicio_maior']} a {p['data_fim_maior']} - {p['vinculo']}, JULGO {p['resultado'].upper()} o pedido, e declaro o período entre {p["data_inicio"]} até {p["data_fim"]} como tempo especial, determinando sua averbação."
                                    p["texto_final_periodos"] = []
                                    p["texto_final_periodos"].append(f"DO PERÍODO ENTRE {p["data_inicio_maior"]} ATÉ {p["data_fim_maior"]} - {p["vinculo"]}:")
                                    p["texto_final_periodos"].append(f"A parte autora alega que o período laborado entre {p["data_inicio_maior"]} ATÉ {p["data_fim_maior"]} - {p["vinculo"]}, é tempo especial.")
                                    p["texto_final_periodos"].extend(p['conclusao_especial_ou_comum'])
                                                                
    
                            if p["origem_do_dado"] == "controvertido_deficiencia":
                                periodo_maior = st.radio("Este período analisado é todo o período controvertido pela parte autora, ou é apenas parte de um período maior controvertido pela parte autora?",[1,2], format_func=lambda x: ("O período analisado é parte de um período controvertido maior" if x == 1 else "O período analisado é toda o período controvertido pela parte autora"), key=f"periodo_maior_{i}")
                                if periodo_maior == 1:
                                    p["data_inicio_maior"] = st.text_input("Qual a data inicial do período maior controvertido? Digite em formato dd/mm/aaaa.", key=f"data_inicio_maior_{i}")
                                    p["data_fim_maior"] = st.text_input("Qual a data final do período maior controvertido? Digite em formato dd/mm/aaaa.", key=f"data_fim_maior_{i}")
                                else: 
                                    p["data_inicio_maior"] = p["data_inicio"]
                                    p["data_fim_maior"] = p["data_fim"]
                                deficiente_procedente_improcedente = st.radio("Será reconhecida a deficiência?", [1,2], format_func=lambda x: "Sim" if x == 1 else "Não", key=f"deficiente_procedente_improcedente_{i}", index=1)
                                p["conclusao_depoimento"] = [p.strip() for p in st.text_area(f"Redija porque chegou nesta conclusão. No caso da resposta ter sido pelo reconhecimento da deficiência, no todo ou em parte, e em qualquer grau, a explicação deve justificar a existência da deficiência, o grau reconhecido e o período reconhecido. Inicie com letra maiúscula e encerre com ponto final.", key=f"conclusao_depoimento_{i}").split("\n") if p.strip()]    
                                p["texto_final_periodos"] = []
                                p["texto_final_periodos"].append(f"DA ALEGADA DEFICIÊNCIA EM GRAU {p["o_que_parte_pede"].upper()} ENTRE {p["data_inicio_maior"]} ATÉ {p["data_fim_maior"]}:")
                                p["texto_final_periodos"].append(f"No que se refere ao período de {p["data_inicio_maior"]} até {p["data_fim_maior"]}, que parte alega ter realizado trabalho na condição de deficiente, foi realizada prova pericial para sua comprovação.")
                                p["texto_final_periodos"].extend(p["conclusao_depoimento"])
                                                                    
                                if deficiente_procedente_improcedente == 1:                            
                                    o_que_parte_pede = st.radio(f"A sentença reconhecerá a deficiência como em grau {p["grau_deficiencia"]}. Para efeito de se verificar se a procedência é total ou parcial, o grau de deficiencia que parte pede para ser reconhecido:", [1,2,3], format_func=lambda x: "LEVE" if x == 1 else "MODERADO" if x == 2 else "GRAVE", key=f"o_que_parte_pede_{i}")
                                    p["o_que_parte_pede"] = {
                                                        1: "LEVE",
                                                        2: "MODERADO",
                                                        3: "GRAVE"
                                                        }[o_que_parte_pede]
                                    if p["o_que_parte_pede"] == p["grau_deficiencia"] and periodo_maior == 2:
                                        p["resultado"] = "Procedente"
                                        p["dispositivo"] = f"Com resolução de mérito, nos termos do art. 487, I do CPC, em relação ao pedido de reconhecimento de deficiência em grau {p['o_que_parte_pede'].lower()} referente ao período de {p['data_inicio_maior']} a {p['data_fim_maior']}, JULGO PROCEDENTE o pedido e declaro a deficiência em grau {p['grau_deficiencia'].lower()} referente ao período de {p['data_inicio']} a {p['data_fim']}, determinando sua averbação para fins de aposentadoria e suas conversões de tempo."
                                        
                                    else:
                                        p["resultado"] = "Procedente em parte"
                                        p["dispositivo"] = f"Com resolução de mérito, nos termos do art. 487, I do CPC, em relação ao pedido de reconhecimento de deficiência em grau {p['o_que_parte_pede'].lower()} referente ao período de {p['data_inicio_maior']} a {p['data_fim_maior']}, JULGO PROCEDENTE EM PARTE o pedido, apenas para declarar a deficiência em grau {p['grau_deficiencia'].lower()} referente ao período de {p['data_inicio']} a {p['data_fim']}, determinando sua averbação para fins de aposentadoria e suas conversões de tempo."
                                        
                                else:
                                    p["resultado"] = "Improcedente"
                                    p["dispositivo"] = f"Com resolução de mérito, nos termos do art. 487, I do CPC, JULGO IMPROCEDENTE o pedido pedido de reconhecimento de deficiência em grau {str(p['o_que_parte_pede']).lower} referente ao período de {p['data_inicio_maior']} a {p['data_fim_maior']}."

                with abasx[1]:
                    st.write("Se desejar utilizar algum trecho, basta copiar aqui e colar no campo em que pretende usar.")
                    for agente, explicacoes in agnocivo.agentes_nocivos.items():
                        st.divider()
                        st.markdown(f"🧪 **{agente}**")
                        for item in explicacoes:
                            st.markdown(f"{item}")
                    
            # verificação de sentença de total extinção por falta de início de prova material - Tema 629
            sentenca_merito = True
            if periodos_para_sentenca:
                lista_de_resultados = []
                for p in periodos_para_sentenca:
                    resultado = p.get("resultado")
                    lista_de_resultados.append(resultado)
                    #st.json(lista_de_resultados)
                if all(r == "Extinto sem julgamento de mérito - Tema 629 - Falta de início de prova material" for r in lista_de_resultados):
                    sentenca_merito = False

            if sentenca_merito == True:

                abas2 = st.tabs(["✏️ Dados Básicos","📈 Resultado da Análise da Fábrica de Dados"])

                # Inicializa DADOS BÁSICOS que acaba criando "paragrafos_sobre_dados_basicos"
                with abas2[0]:
                    if "dados_basicos" not in st.session_state:
                        st.session_state["dados_basicos"] = {
                            "DER": payload_final.get("dataApuracao"),
                            "explicacao_reafirmacao": [],
                            "DIP": "",
                            "especie": int(payload_final.get("opcoesContagem", {}).get("especie")),
                            "tipo_de_pedido": "",
                            "NB": "",
                            "beneficio_revisado": "",
                            "beneficio_revisado_string": "",
                            "beneficio_pedido": "",
                            "beneficio_pedido_string": "",
                            "beneficio_mais_vantajoso": "",
                            "resultado": "",
                            "redacao": [],
                            "dispositivo": []
                        }

                    paragrafos_sobre_dados_basicos = st.session_state["dados_basicos"]

                    st.subheader("🔍 Dados Básicos")
                    data_atual = datetime.now()
                    DIP = data_atual.strftime("01/%m/%Y")
                    especie = int(payload_final.get("opcoesContagem", {}).get("especie"))
                    DER = f"{payload_final.get("dataApuracao")}"
                    st.write(f"Os cálculos estão utilizando como DER a data de ***`{DER}`***, advinda de {"Fábrica de Cálculos"}")
                    DER_ou_DER_reafirmada = st.radio(f"Esta data é a DER do benefício, ou a DER Reafirmada?", [1,2], format_func=lambda x: "DER" if x == 1 else "DER Reafirmada")
                    if DER_ou_DER_reafirmada == 2:
                        explicacao_reafirmacao = [p.strip() for p in st.text_area(f"Redija porque a DER precisou ser reafirmada, e porque a referida data foi a data escolhida. Inicie com letra maiúscula e encerre com ponto final.").split("\n") if p.strip()]    
                        if st.session_state["dados_basicos"]["explicacao_reafirmacao"] != explicacao_reafirmacao:
                            st.session_state["dados_basicos"]["explicacao_reafirmacao"] = explicacao_reafirmacao
                    if DIP not in st.session_state["dados_basicos"]["DIP"]:
                        st.session_state["dados_basicos"]["DIP"] = DIP
                    qual_pedido_concessao_revisao = st.radio("O pedido é de concessão de benefício ou de revisão de benefício já concedido", [1,2], format_func=lambda x: "Concessão" if x == 1 else "Revisão")
                    if qual_pedido_concessao_revisao == 2:  #revisão
                        tipo_de_pedido = "revisão"
                        NB = st.text_input("Qual o número do benefício a ser revisado?")
                        qual_beneficio_revisado = st.radio("O benefício a ser revisado é uma:", [1,2,3,4], format_func=lambda x: "41 - Aposentadoria por idade" if x == 1 else "42 - Aposentadoria por tempo de contribuição" if x == 2 else "46 - Aposentadoria especial" if x == 3 else "57 - Aposentadoria por tempo de contribuição do professor", index=1)
                        beneficio_revisado = {1: 41, 2: 42, 3: 46, 4: 57}[qual_beneficio_revisado]
                        beneficio_revisado_string = {41: "Aposentadoria por idade", 42: "Aposentadoria por tempo de contribuição", 46: "Aposentadoria especial", 57: "Aposentadoria por tempo de contribuição do professor"}[beneficio_revisado]
                        qual_resultado = st.radio("Em relação a apenas este pedido revisional, o que a parte pede será julgado:", [1,2,3], format_func=lambda x: "Totalmente procedente (o benefício será revisto nos exatos moldes requeridos na inicial)" if x == 1 else "Procedente em parte (o benefício será revisto, mas com qualquer parâmetro distinto daquilo pedido na inicial)" if x == 2 else "Totalmente improcedente (o benefício não será revisto)")
                        resultado = {1: "Procedente", 2: "Procedente em parte", 3:"Improcedente"}[qual_resultado]
                        if st.session_state["dados_basicos"]["tipo_de_pedido"] != tipo_de_pedido:
                            st.session_state["dados_basicos"]["tipo_de_pedido"] = tipo_de_pedido
                        if st.session_state["dados_basicos"]["NB"] != NB:
                            st.session_state["dados_basicos"]["NB"] = NB
                        if st.session_state["dados_basicos"]["beneficio_revisado"] != beneficio_revisado:
                            st.session_state["dados_basicos"]["beneficio_revisado"] = beneficio_revisado
                        if st.session_state["dados_basicos"]["beneficio_revisado_string"] != beneficio_revisado_string:
                            st.session_state["dados_basicos"]["beneficio_revisado_string"] = beneficio_revisado_string
                        if st.session_state["dados_basicos"]["resultado"] != resultado:
                            st.session_state["dados_basicos"]["resultado"] = resultado

                    else:     #concessão
                        tipo_de_pedido = "concessão"
                        qual_beneficio_pedido = st.radio("O pedido principal inicial foi para para concessão de qual benefício? Se houver pedidos subsidiários, indique apenas o pedido principal", [1,2,3,4], format_func=lambda x: "41 - Aposentadoria por idade" if x == 1 else "42 - Aposentadoria por tempo de contribuição" if x == 2 else "46 - Aposentadoria especial" if x == 3 else "57 - Aposentadoria por tempo de contribuição do professor", index=1)
                        beneficio_pedido = {1: 41, 2: 42, 3: 46, 4: 57}[qual_beneficio_pedido]
                        beneficio_pedido_string = {41: "Aposentadoria por idade", 42: "Aposentadoria por tempo de contribuição", 46: "Aposentadoria especial", 57: "Aposentadoria por tempo de contribuição do professor"}[beneficio_pedido]
                        qual_resultado = st.radio("Em relação a apenas este pedido de concessão de aposentadoria, o que a parte pede será julgado:", [1,2,3], format_func=lambda x: "Totalmente procedente (o benefício será concedido nos exatos moldes requeridos na inicial)" if x == 1 else "Procedente em parte (o benefício será concedido, mas com qualquer parâmetro distinto daquilo pedido na inicial)" if x == 2 else "Totalmente improcedente (o benefício não será concedido)")
                        resultado = {1: "Procedente", 2: "Procedente em parte", 3:"Improcedente"}[qual_resultado]
                        if st.session_state["dados_basicos"]["tipo_de_pedido"] != tipo_de_pedido:
                            st.session_state["dados_basicos"]["tipo_de_pedido"] = tipo_de_pedido
                        if st.session_state["dados_basicos"]["beneficio_pedido"] != beneficio_pedido:
                            st.session_state["dados_basicos"]["beneficio_pedido"] = beneficio_pedido
                        if st.session_state["dados_basicos"]["beneficio_pedido_string"] != beneficio_pedido_string:
                            st.session_state["dados_basicos"]["beneficio_pedido_string"] = beneficio_pedido_string
                        if st.session_state["dados_basicos"]["resultado"] != resultado:
                            st.session_state["dados_basicos"]["resultado"] = resultado
                    
                    cumpridos = resposta.get("beneficios", {}).get("cumpridos", [])
                    if cumpridos:
                        if len(cumpridos) > 1:
                            opcoes_beneficios = ["Não é possível definir o melhor benefício neste momento (escolha deve ficar a cargo das partes, apurada em liquidação)"]
                            for i, b in enumerate(cumpridos):
                                if b.get("dadosApuracao", {}).get("temDireito"):
                                    aposentadoria = f"{b.get('descricao')} com fundamento na {b.get('fundamento')}"
                                    opcoes_beneficios.append(aposentadoria)
                            qual_beneficio_mais_vantajoso = str(st.radio("Qual destes é o benefício mais vantajoso, para ser deferido?", opcoes_beneficios, key=f"beneficio_mais_vantajoso{i}"))
                            if qual_beneficio_mais_vantajoso == opcoes_beneficios[0]:                            
                                lista_melhor_beneficio = "aposentadoria mais vantajoso entre "
                                for id, paragrafo in enumerate(cumpridos, start=1):
                                    lista_melhor_beneficio += f"- {id}: {paragrafo.get('descricao')} com fundamento na {paragrafo.get('fundamento')} "
                                beneficio_mais_vantajoso = f"{lista_melhor_beneficio}- a ser definido pelo INSS quando da implantação"

                            else:
                                beneficio_mais_vantajoso = f"{qual_beneficio_mais_vantajoso}, que se mostra o benefício mais vantajoso" 
                        else:
                            for i, b in enumerate(cumpridos):
                                if b.get("dadosApuracao", {}).get("temDireito"):
                                    aposentadoria = f"{b.get('descricao')} com fundamento na {b.get('fundamento')}"
                                    beneficio_mais_vantajoso = aposentadoria
                    else:
                        beneficio_mais_vantajoso = ""

                    if st.session_state["dados_basicos"]["beneficio_mais_vantajoso"] != beneficio_mais_vantajoso:
                        st.session_state["dados_basicos"]["beneficio_mais_vantajoso"] = beneficio_mais_vantajoso
                    
                    beneficio_info = (f"{beneficio_revisado_string.lower()} (NB {NB})"if tipo_de_pedido == "revisão" else f"{beneficio_pedido_string.lower()}")
                    if resultado == "Improcedente": 
                        dispositivo = [
                                f"Com resolução de mérito, nos termos do art. 487, I do CPC, JULGO IMPROCEDENTE o pedido de {tipo_de_pedido} de {beneficio_info}."
                                ]
                    else: 
                        dispositivo = [
                                f"Com resolução de mérito, nos termos do art. 487, I do CPC, em relação ao pedido de {tipo_de_pedido} de {beneficio_info} JULGO {resultado.upper()} para determinar a {tipo_de_pedido} e consequente implantação do benefício de {beneficio_mais_vantajoso}, na forma apurada nesta sentença. Fixo a DIB em {DER}. RMI e RMA a serem calculadas pelo INSS. NB a ser definido pelo INSS",
                                f"Condeno o INSS ao pagamento dos atrasados devidos desde a DIB fixada, até a DIP em {DIP}, atualizados desde cada competência devida e com juros desde a propositura da demanda, pelos índices e percentuais do Manual de Cálculos da Justiça Federal, a ser apurado em cumprimento invertido de sentença.",
                                f"Fica autorizado o desconto de eventuais valores recebidos a título de benefícios inacumuláveis.",
                                f"Condeno o INSS ao ressarcimento de eventuais honorários periciais antecipados pela Justiça Federal nesta lide (art. 82, § 2º, do CPC).",
                                f"Considerando que o momento da prolação de sentença é oportuno para distribuir o ônus do tempo do processo, com vistas a salvaguardar a eficácia do princípio constitucional da razoável duração do processo e ao mesmo tempo privilegiar o direito provável em detrimento do improvável, demonstrada a verossimilhança das alegações da parte autora e diante do nítido caráter alimentar da verba pleiteada, nos termos do art. 294 e 300, do CPC ANTECIPA A TUTELA JURISDICIONAL para determinar ao INSS que providencie a implantação da {beneficio_mais_vantajoso} na forma concedida, com data de início de pagamento em {DIP} (DIP).",
                                f"O INSS deverá providenciar a implantação do benefício previdenciário ora concedido no prazo legal, sendo a contagem em dias úteis, sendo que constitui ônus das partes informar ao Juízo sobre a efetiva implantação do benefício ou eventual descumprimento do prazo pelo INSS/APSADJ."
                                ]
                            
                    if st.session_state["dados_basicos"]["dispositivo"] != dispositivo:
                        st.session_state["dados_basicos"]["dispositivo"] = dispositivo

                    redacao = [f"A parte autora pede a {tipo_de_pedido} do benefício de {beneficio_info}."]
                    if resultado == "Procedente em parte":
                        redacao.append(f"Não houve cumprimento dos requisitos, da exata forma que pedido na inicial. A apuração de tempo, no entanto, demonstrou que a parte autora tem direito ao benefício de {beneficio_mais_vantajoso}, com DER em {DER}.")
                        if DER_ou_DER_reafirmada == 2:
                            redacao.extend(explicacao_reafirmacao)
                    if resultado == "Procedente":
                        redacao.append(f"Pela apuração, conclui-se que a parte autora tem direito ao benefício de {beneficio_mais_vantajoso}, com DER em {DER}.")
                        if DER_ou_DER_reafirmada == 2:
                            redacao.extend(explicacao_reafirmacao)
                    if resultado == "Improcedente":
                        redacao.append(f"Pela apuração, verifica-se que a parte autora não cumpriu o requisito para concessão de aposentadoria, e seu pedido deve ser julgado improcedente.") 
                    redacao.append(f"São os seguintes, os resultados da apuração:")
                    
                    if st.session_state["dados_basicos"]["redacao"] != redacao:
                        st.session_state["dados_basicos"]["redacao"] = redacao
                        

                # Exibe os resultados advindos da API, e cria "paragrafos_sobre_tempo" como dicionário
                with abas2[1]:
                    if "analise_paragrafos" not in st.session_state:
                        st.session_state["analise_paragrafos"] = {
                            "demonstrativo_cada_vinculo": [],
                            "demonstrativo_tempo_total": [],
                            "demonstrativo_tempo_total_magisterio": [],
                            "demonstrativo_tempo_total_PCD": [],
                            "analise_API_beneficios_cumpridos": [],
                            "analise_API_beneficios_nao_cumpridos": []
                        }

                    paragrafos_sobre_tempo = st.session_state["analise_paragrafos"]
                    # 1. DEMONSTRATIVO DE CADA VÍNCULO 
                    if "dadosPeriodos" in resposta:
                        st.markdown("📊 Demonstrativo de Tempo de cada vínculo")
                        for i, item in enumerate(resposta["dadosPeriodos"]):
                            if isinstance(item, dict):
                                data_inicial_corrigida = formatar_data_iso_para_br(item.get("dataInicial"))
                                data_final_corrigida = formatar_data_iso_para_br(item.get("dataFinal"))
                                valor_forma_contagem = item.get("formaContagem")
                                if valor_forma_contagem == "COMUM":
                                    forma_contagem_corrigida = "tempo comum"
                                if valor_forma_contagem == "RURAL":
                                    forma_contagem_corrigida = "tempo rural"
                                if valor_forma_contagem in ["ESPECIAL_25", "ESPECIAL_20", "ESPECIAL_15"]:
                                    forma_contagem_corrigida = "tempo especial"
                                if valor_forma_contagem == "MAGISTERIO":
                                    forma_contagem_corrigida = "tempo de magistério"
                                valor_grau_deficiencia = item.get("grauDeficiencia")
                                if valor_grau_deficiencia == "SEM_DEFICIENCIA":
                                    grau_deficiencia_corrigido = "sem constatação de deficiência"
                                if valor_grau_deficiencia in ["LEVE", "MODERADO", "GRAVE"]:
                                    grau_deficiencia_corrigido = f"com constatação de deficiência em grau {item.get("grauDeficiencia").lower()}"
                                if item.get("tempoSimples") == item.get("tempoConvertido"):
                                    tempo_corrigido = f"tempo de contribuição de {formatar_dias_em_anos_meses_dias(item.get("tempoSimples"))}"
                                else:
                                    tempo_corrigido = f"tempo simples apurado de {formatar_dias_em_anos_meses_dias(item.get("tempoSimples"))}, que, convertido, equivale a {formatar_dias_em_anos_meses_dias(item.get("tempoConvertido"))}"
                                carencia_corrigida = f"{item.get("carencia")} meses"
                                idade_corigida = f"{item.get("idade")} anos de idade"
                                texto_demonstrativo_cada_vinculo = f"Em relação ao período de {data_inicial_corrigida} até {data_final_corrigida}, computado como {forma_contagem_corrigida} e {grau_deficiencia_corrigido}, foi apurado um {tempo_corrigido}, com carência de {carencia_corrigida} e idade da parte autora de {idade_corigida}."
                                if texto_demonstrativo_cada_vinculo not in st.session_state["analise_paragrafos"]["demonstrativo_cada_vinculo"]:
                                    st.session_state["analise_paragrafos"]["demonstrativo_cada_vinculo"].append(texto_demonstrativo_cada_vinculo)                              
                                st.markdown(texto_demonstrativo_cada_vinculo)

                    # 2. DEMONSTRATIVO DE TEMPO
                    if "demonstrativo" in resposta:
                        st.markdown("📊 Demonstrativo de Tempo")
                        for item in resposta["demonstrativo"]:
                            if isinstance(item, dict):
                                if item.get("rotulo") in ["Até 11/10/2023", "Até 12/16/1998 (EC 20)", "Até 11/28/1999 (Lei 9.876)", "Até 11/13/2019 (Lei 9.876)"]:
                                    rotulo_corrigido = corrigir_data_em_string(item.get("rotulo"))
                                    if item.get("simples") == item.get("convertido"):
                                        tempo_demonstrativo_corrigido = f"tempo de contribuição apurado de {formatar_dias_em_anos_meses_dias(item.get("simples"))}, não havendo conversões no período"
                                    else:
                                        tempo_demonstrativo_corrigido = f"tempo de contribuição total de {formatar_dias_em_anos_meses_dias(item.get("convertido"))}, já computadas as conversões de tempo reconhecidas no período" 
                                    carencia_demonstrativo_corrigida = f"{item.get("carencia")} meses"
                                    idade_demonstrativo_corrigida = f"{item.get("idade")} anos de idade"
                                    texto_demonstrativo_tempo_total = f"{rotulo_corrigido} foi apurado o {tempo_demonstrativo_corrigido}, com carência de {carencia_demonstrativo_corrigida} e idade da parte autora de {idade_demonstrativo_corrigida}."
                                    if texto_demonstrativo_tempo_total not in st.session_state["analise_paragrafos"]["demonstrativo_tempo_total"]:
                                        st.session_state["analise_paragrafos"]["demonstrativo_tempo_total"].append(texto_demonstrativo_tempo_total)
                                    st.markdown(texto_demonstrativo_tempo_total)                                    
                                if item.get("rotulo") == "Magistério":
                                    texto_demonstrativo_tempo_total_magisterio = f"O tempo apurado apenas de magistério foi de {formatar_dias_em_anos_meses_dias(item.get("simples"))}."
                                    if texto_demonstrativo_tempo_total_magisterio not in st.session_state["analise_paragrafos"]["demonstrativo_tempo_total_magisterio"]:
                                        st.session_state["analise_paragrafos"]["demonstrativo_tempo_total_magisterio"].append(texto_demonstrativo_tempo_total_magisterio)
                                    st.markdown(texto_demonstrativo_tempo_total_magisterio)
                                if item.get("rotulo") in ["Leve", "Moderada", "Grave"]:                                                           
                                    texto_tempo_PCD_simples = f"O tempo apurado de trabalho com Deficiência {item.get("rotulo")} foi de {formatar_dias_em_anos_meses_dias(item.get("simples"))}."
                                    if texto_tempo_PCD_simples not in st.session_state["analise_paragrafos"]["demonstrativo_tempo_total_PCD"]:
                                        st.session_state["analise_paragrafos"]["demonstrativo_tempo_total_PCD"].append(texto_tempo_PCD_simples)
                                    st.markdown(texto_tempo_PCD_simples)
                                if len(st.session_state["analise_paragrafos"]["demonstrativo_tempo_total_PCD"]) > 1:
                                    if item.get("rotulo") in ["Convertido para preponderante (Deficiência Grave)", "Convertido para preponderante (Deficiência Moderada)", "Convertido para preponderante (Deficiência Leve)"]:
                                        texto_tempo_PCD_convertido = f"{item.get("rotulo")}, o tempo total de trabalho com deficiência equivale a {formatar_dias_em_anos_meses_dias(item.get("convertido"))}."
                                        if texto_tempo_PCD_convertido not in st.session_state["analise_paragrafos"]["demonstrativo_tempo_total_PCD"]:
                                            st.session_state["analise_paragrafos"]["demonstrativo_tempo_total_PCD"].append(texto_tempo_PCD_convertido)
                                        st.markdown(texto_tempo_PCD_convertido)

                    if "totais" in resposta:
                        st.markdown("📊 Demonstrativo de Tempo na DER")
                        item = resposta["totais"]  
                        rotulo_DER_corrigido = corrigir_data_em_string(item.get("rotulo"))
                        if item.get("simples") == item.get("convertido"):
                            tempo_DER_demonstrativo_corrigido = (f"tempo de contribuição apurado de {formatar_dias_em_anos_meses_dias(item.get('simples'))}, não havendo conversões no período")
                        else:
                            tempo_DER_demonstrativo_corrigido = (f"tempo de contribuição total de {formatar_dias_em_anos_meses_dias(item.get('convertido'))}, já computadas as conversões de tempo reconhecidas no período")
                        carencia_DER_demonstrativo_corrigida = f"{item.get('carencia')} meses"
                        idade_DER_demonstrativo_corrigida = f"{item.get('idade')} anos de idade"
                        texto_DER_demonstrativo_tempo_total = (f"{rotulo_DER_corrigido} foi apurado o {tempo_DER_demonstrativo_corrigido}, com carência de {carencia_DER_demonstrativo_corrigida} e idade da parte autora de {idade_DER_demonstrativo_corrigida}.")
                        if texto_DER_demonstrativo_tempo_total not in st.session_state["analise_paragrafos"]["demonstrativo_tempo_total"]:
                            st.session_state["analise_paragrafos"]["demonstrativo_tempo_total"].append(texto_DER_demonstrativo_tempo_total)
                        st.markdown(texto_DER_demonstrativo_tempo_total)

                    
                    # 3. BENEFÍCIOS CUMPRIDOS
                    cumpridos = resposta.get("beneficios", {}).get("cumpridos", [])
                    if cumpridos:
                        st.markdown("🟩 Benefícios com requisitos cumpridos")
                        for i, b in enumerate(cumpridos):
                            if b.get("dadosApuracao", {}).get("temDireito"):
                                inicio_vigencia = formatar_data_iso_para_br(b.get("inicioVigencia"))
                                termino_vigencia = (formatar_data_iso_para_br(b.get("terminoVigencia")) if b.get("terminoVigencia") else "data atual")
                                data_apuracao = formatar_data_iso_para_br(b.get("dadosApuracao", {}).get("naApuracao", {}).get("data"))
                                valores = (b.get("dadosApuracao", {}).get("naApuracao", {}).get("valoresApurados", {}))
                                textos_cumpridos = []
                                for chave, detalhe in valores.items():
                                    resultado = detalhe.get("resultado", {})
                                    por_extenso = resultado.get("porExtenso")
                                    if por_extenso:                                
                                        textos_cumpridos.append(f"{por_extenso}")
                                if textos_cumpridos:
                                    lista_requisitos_cumpridos = ""
                                    for id, texto in enumerate(textos_cumpridos, start=1):
                                        lista_requisitos_cumpridos += f" {id} - {texto}."
                                    texto_analise_API_beneficios_cumpridos = f"Em relação à {b.get('descricao')} com fundamento na {b.get('fundamento')} (vigente de {inicio_vigencia} até {termino_vigencia}), foi apurado até a data de {data_apuracao}, o(s) seguinte(s) requisito(s) foi(ram) cumprido(s), resultando no direito à concessão do benefício:{lista_requisitos_cumpridos}"
                                    if texto_analise_API_beneficios_cumpridos not in st.session_state["analise_paragrafos"]["analise_API_beneficios_cumpridos"]:
                                        st.session_state["analise_paragrafos"]["analise_API_beneficios_cumpridos"].append(texto_analise_API_beneficios_cumpridos)
                                    st.markdown(texto_analise_API_beneficios_cumpridos)                        

                    # 4. BENEFÍCIOS NÃO CUMPRIDOS
                    nao_cumpridos = resposta.get("beneficios", {}).get("naoCumpridos", [])
                    if nao_cumpridos:
                        st.markdown("🟥 Benefícios com requisitos não cumpridos")
                        for i, b in enumerate(nao_cumpridos):
                            if b.get("dadosApuracao", {}).get("temDireito") == False:
                                inicio_vigencia = formatar_data_iso_para_br(b.get("inicioVigencia"))
                                termino_vigencia = (formatar_data_iso_para_br(b.get("terminoVigencia")) if b.get("terminoVigencia") else "data atual")
                                data_apuracao = formatar_data_iso_para_br(b.get("dadosApuracao", {}).get("naApuracao", {}).get("data"))
                                valores = (b.get("dadosApuracao", {}).get("naApuracao", {}).get("valoresApurados", {}))
                                textos_nao_cumpridos = []
                                for chave, detalhe in valores.items():
                                    resultado = detalhe.get("resultado", {})
                                    por_extenso = resultado.get("porExtenso")
                                    if por_extenso:
                                        textos_nao_cumpridos.append(f"{por_extenso}")
                                if textos_nao_cumpridos:
                                    lista_requisitos_nao_cumpridos = ""
                                    for id, texto in enumerate(textos_nao_cumpridos, start=1):
                                        lista_requisitos_nao_cumpridos += f" {id} - {texto}."
                                    texto_analise_API_beneficios_nao_cumpridos = f"Em relação à {b.get("descricao")} com fundamento na {b.get("fundamento")} (vigente de {inicio_vigencia} até {termino_vigencia}), foi apurado até a data de {data_apuracao}, o(s) seguinte(s) requisito(s) não foi(ram) cumprido(s), resultando na impossibilidade de concessão do benefício:{lista_requisitos_nao_cumpridos}"                 
                                    if texto_analise_API_beneficios_nao_cumpridos not in st.session_state["analise_paragrafos"]["analise_API_beneficios_nao_cumpridos"]:
                                        st.session_state["analise_paragrafos"]["analise_API_beneficios_nao_cumpridos"].append(texto_analise_API_beneficios_nao_cumpridos)
                                    st.markdown(texto_analise_API_beneficios_nao_cumpridos)


               # EXIBIR OS ARQUIVOS PRODUZIDOS
                # with st.expander("🔍 Ver arquivos produzidos"):
                #     st.write(periodos_para_sentenca) #lista
                #     st.write(paragrafos_sobre_tempo) #dicionario
                #     st.write(paragrafos_sobre_dados_basicos) #dicionario

                # MONTAR REDACAO DA SENTENÇA
                desfecho_dispositivo = [
                                f"Sem condenação em honorários nesta instância.",
                                f"Defiro os benefícios da gratuidade.",
                                f"Em caso de interposição de embargos de declaração, intime-se a parte contrária para contrarrazões no prazo legal, e tornem conclusos para julgamento.",
                                f"Interposto recurso, intime-se a parte contrária para contrarrazões no prazo legal. Após, remetam-se os autos às Turmas Recursais.",
                                f"Com o trânsito em julgado, implantado o benefício, dê-se início ao cumprimento de sentença.",
                                f"Proceda a Secretaria como necessário.",
                                f"Int."
                                ]

                texto_fundamentacao_sentenca = []

                # redação da análise de cada período de trabalho
                for item in periodos_para_sentenca:
                    texto_fundamentacao_sentenca.extend(item.get("texto_final_periodos"))
                # sobre a simulação do tempo de serviço
                texto_fundamentacao_sentenca.append("SOBRE O BENEFÍCIO PLEITEADO:")
                if paragrafos_sobre_dados_basicos.get("redacao"):
                    texto_fundamentacao_sentenca.extend(paragrafos_sobre_dados_basicos.get("redacao"))
                # if paragrafos_sobre_tempo.get("demonstrativo_cada_vinculo"):
                #     texto_fundamentacao_sentenca.append("TEMPO DE TRABALHO EM CADA VÍNCULO:")
                #     texto_fundamentacao_sentenca.extend(paragrafos_sobre_tempo["demonstrativo_cada_vinculo"])
                if paragrafos_sobre_tempo.get("demonstrativo_tempo_total_magisterio"):
                    texto_fundamentacao_sentenca.append("TEMPO DE TRABALHO NO MAGISTÉRIO:")
                    texto_fundamentacao_sentenca.extend(paragrafos_sobre_tempo["demonstrativo_tempo_total_magisterio"])
                if paragrafos_sobre_tempo.get("demonstrativo_tempo_total_PCD"):
                    texto_fundamentacao_sentenca.append("TEMPO DE TRABALHO COM DEFICIÊNCIA:")
                    texto_fundamentacao_sentenca.extend(paragrafos_sobre_tempo["demonstrativo_tempo_total_PCD"]) 
                if paragrafos_sobre_tempo.get("demonstrativo_tempo_total"):
                    texto_fundamentacao_sentenca.append("SÍNTESE DO TEMPO TOTAL DE CONTRIBUIÇÃO:")
                    texto_fundamentacao_sentenca.extend(paragrafos_sobre_tempo["demonstrativo_tempo_total"])                   
                beneficios_cumpridos = paragrafos_sobre_tempo.get("analise_API_beneficios_cumpridos", [])
                beneficios_nao = paragrafos_sobre_tempo.get("analise_API_beneficios_nao_cumpridos", [])
                if beneficios_nao:
                    texto_fundamentacao_sentenca.append("APOSENTADORIAS COM REQUISITOS QUE NÃO FORAM CUMPRIDOS:")
                    texto_fundamentacao_sentenca.extend(beneficios_nao)
                if beneficios_cumpridos:
                    texto_fundamentacao_sentenca.append("APOSENTADORIAS COM REQUISITOS CUMPRIDOS:")
                    texto_fundamentacao_sentenca.extend(beneficios_cumpridos)

                # 2. Inicia lista que armazenará dispositivo
                dispositivo = []

                dispositivo.append(f"Isto posto:")
                for item in periodos_para_sentenca:
                    if item.get("resultado") == "Extinto sem julgamento de mérito - Tema 629 - Falta de início de prova material":
                        texto = item.get("dispositivo")
                        dispositivo.append(texto)

                    if item.get("resultado") == "Procedente":
                        texto = item.get("dispositivo")
                        dispositivo.append(texto)

                    if item.get("resultado") == "Procedente em parte":
                        texto = item.get("dispositivo")
                        dispositivo.append(texto)

                    if item.get("resultado") == "Improcedente":
                        texto = item.get("dispositivo")
                        dispositivo.append(texto)

                if paragrafos_sobre_dados_basicos.get("dispositivo"):
                    dispositivo.extend(paragrafos_sobre_dados_basicos["dispositivo"])
                dispositivo.extend(desfecho_dispositivo)
                

            else:
                st.write("📝 Sentença será de extinção")

                desfecho_dispositivo = [
                                f"Sem condenação em honorários nesta instância.",
                                f"Defiro os benefícios da gratuidade.",
                                f"Em caso de interposição de embargos de declaração, intime-se a parte contrária para contrarrazões no prazo legal, e tornem conclusos para julgamento.",
                                f"Interposto recurso, intime-se a parte contrária para contrarrazões no prazo legal. Após, remetam-se os autos às Turmas Recursais.",
                                f"Com o trânsito em julgado, implantado o benefício, dê-se início ao cumprimento de sentença.",
                                f"Proceda a Secretaria como necessário.",
                                f"Int."
                                ]

                texto_fundamentacao_sentenca = []

                # redação da análise de cada período de trabalho
                for item in periodos_para_sentenca:
                    texto_fundamentacao_sentenca.extend(item.get("texto_final_periodos"))

                # 2. Inicia lista que armazenará dispositivo
                dispositivo = []

                dispositivo.append(f"Isto posto:")
                for item in periodos_para_sentenca:
                    if item.get("resultado") == "Extinto sem julgamento de mérito - Tema 629 - Falta de início de prova material":
                        texto = item.get("dispositivo")
                        dispositivo.append(texto)
                dispositivo.append(f"Prejudicados os demais pedidos sucessivos.")
                dispositivo.extend(desfecho_dispositivo)

                
            # # 8. Exibição final
            # with st.expander("Redação Final da Sentença"):
            #     # st.write(texto_fundamentacao_sentenca)
            #     # st.write(dispositivo)
            #     st.markdown("\n\n".join(texto_fundamentacao_sentenca))
            #     st.markdown("\n\n".join(dispositivo))


            if st.button("Gerar Sentença"):
                doc = Document()
                doc.add_paragraph(f"Processo: {processo_formatado}")
                texto_base(doc, fundamento_questao)
                ft.alinhamento_parag_dispositivo(doc, texto_fundamentacao_sentenca)
                ft.alinhamento_parag_dispositivo(doc, dispositivo)
                ft.salvar_docx_temporario(doc, processo_formatado)              
            
        else:
            st.error(f"Erro {response.status_code}: {response.text}")

if st.button("🔄 Limpar tudo e reiniciar"):
    for key in list(st.session_state.keys()):
        del st.session_state[key]
    st.session_state.clear()
    st.rerun()
