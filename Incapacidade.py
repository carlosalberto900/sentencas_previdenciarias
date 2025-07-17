from datetime import datetime
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import tempfile
import platform
import streamlit as st
import funcoes_texto as ft

def texto_base(doc, fundamento_questao):
    if fundamento_questao == 1:
        for i, paragrafo in enumerate(fundamento_base):
            parag = doc.add_paragraph(paragrafo)
            parag.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY 
            if i in [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 24, 25, 27, 30, 32, 33, 35, 36, 37, 39, 41, 42, 44, 46, 48]:
                parag.paragraph_format.first_line_indent = Cm(2) 
            elif i in [13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 26, 28, 29, 31, 34, 38, 40, 43, 45, 47]:
                parag.paragraph_format.left_indent = Cm(2)
    if fundamento_questao == 2:
        for linha in fundamento_custom.split("\n"):
            if linha.split():
                parag = doc.add_paragraph(linha.strip())
                parag.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                parag.paragraph_format.first_line_indent = Cm(2)

fundamento_questao = st.radio(
    "**Relatório e fundamentação jurídica**\n\n"
    "Toda sentença gerada possui dispensa de relatório e uma fundamentação jurídica básica, que vai até onde se inicia a análise do caso concreto.\n\n"  
    "Você pode usar esta fundamentação, ou fornecer sua **própria fundamentação**, com ou sem relatório (fazendo menção à dispensa, se for o caso.)\n\n"  
    "**Como você deseja prosseguir?**",
    [1, 2],
    format_func=lambda x: "Vou usar a fundamentação padrão deste aplicativo." if x == 1 else "Desejo fornecer a minha fundamentação."
    )
if fundamento_questao == 1:
    fundamento_base = [
                "Vistos.",
                "Trata-se de pedido de concessão/restabelecimento de benefício por incapacidade.",
                "Dispensado o relatório nos termos do art. 38 da Lei n. 9.099/95.",
                "DECIDO.",
                "O feito comporta julgamento imediato.",
                "O Juízo é competente porque o valor da causa está abaixo de 60 salários-mínimos. Não há que se falar em renúncia a valores excedentes, neste caso, justamente porque o valor da causa já foi fixado na alçada do JEF. O tema 1030 do STJ prevê que: “Ao autor que deseje litigar no âmbito de Juizado Especial Federal Cível, é lícito renunciar, de modo expresso e para fins de atribuição de valor à causa, ao montante que exceda os 60 (sessenta) salários mínimos previstos no art. 3º, caput, da Lei 10.259/2001, aí incluídas, sendo o caso, até doze prestações vincendas, nos termos do art. 3º, § 2º, da referida lei, c/c o art. 292, §§ 1º e 2º, do CPC/2015.” Logo, a renúncia só tem sentido nos casos em que o valor da causa, apurado sem qualquer renúncia, pudesse ser superior à alçada, pois nesta hipótese a renúncia daria ensejo à parte litigar sob a competência do Juizado. Não é o caso dos autos, onde o valor da causa apontado já está abaixo da alçada, e a parte ré não aponta erro na sua apuração.",
                "Partes legítimas. Passo ao mérito.",
                "Não há prescrição de fundo de direito, mas apenas das parcelas vencidas, na forma da súmula 85 do STJ. Estão prescritas eventuais parcelas anteriores ao quinquênio que precede a propositura da ação.",
                "Passo a analisar o pedido da parte autora.",
                "A concessão do benefício por incapacidade temporária (auxílio-doença) é devida quando o segurado ficar impossibilitado para o seu trabalho ou para a sua atividade habitual por mais de 15 (quinze) dias consecutivos, respeitada a carência, quando exigida pela lei, conforme determinam, especialmente, as normas dos artigos 25, inciso I e 59 e seguintes da Lei n° 8.213, de 24.07.91, que disciplina o Plano de Benefícios da Previdência Social.",
                "Por sua vez, a concessão do benefício por incapacidade permanente (aposentadoria por invalidez) é devida quando o segurado ficar impossibilitado total e permanentemente, insusceptível de reabilitação, para o exercício de atividade que lhe garanta a subsistência, respeitada a carência, quando exigida pela lei, conforme determinam, especialmente, as normas dos artigos 42 e seguintes da Lei n° 8.213, de 24.07.91, que disciplina o Plano de Benefícios da Previdência Social.",
                "A concessão dos benefícios por incapacidade exige, nos termos dos arts. 42 e 59 da Lei 8.213/91, a presença simultânea dos seguintes requisitos: (a) incapacidade laborativa, (b) qualidade de segurado na época em que iniciada a incapacidade e (c) recolhimento de contribuições mensais em número suficiente para completar a carência legal. Além disso, é necessário que a doença incapacitante não seja preexistente ou, caso seja, que a incapacidade resulte de agravamento da doença, verificado após a filiação ao regime geral de previdência (artigo 42, § 2º, e artigo 59, parágrafo único, da Lei 8.213/91).",
                "A qualidade de segurado mantém enquanto houver recolhimento válido de contribuição previdenciária, até o final do período de graça na forma estipulada no art. 15 da Lei n. 8.213/91:",
                "Art. 15. Mantém a qualidade de segurado, independentemente de contribuições:",
                "I - sem limite de prazo, quem está em gozo de benefício, exceto do auxílio-acidente;",
                "II - até 12 (doze) meses após a cessação das contribuições, o segurado que deixar de exercer atividade remunerada abrangida pela Previdência Social ou estiver suspenso ou licenciado sem remuneração;",
                "III - até 12 (doze) meses após cessar a segregação, o segurado acometido de doença de segregação compulsória;",
                "IV - até 12 (doze) meses após o livramento, o segurado retido ou recluso;",
                "V - até 3 (três) meses após o licenciamento, o segurado incorporado às Forças Armadas para prestar serviço militar;",
                "VI - até 6 (seis) meses após a cessação das contribuições, o segurado facultativo.",
                "§ 1º O prazo do inciso II será prorrogado para até 24 (vinte e quatro) meses se o segurado já tiver pago mais de 120 (cento e vinte) contribuições mensais sem interrupção que acarrete a perda da qualidade de segurado.",
                "§ 2º Os prazos do inciso II ou do § 1º serão acrescidos de 12 (doze) meses para o segurado desempregado, desde que comprovada essa situação pelo registro no órgão próprio do Ministério do Trabalho e da Previdência Social.",
                "§ 3º Durante os prazos deste artigo, o segurado conserva todos os seus direitos perante a Previdência Social.",
                "§ 4º A perda da qualidade de segurado ocorrerá no dia seguinte ao do término do prazo fixado no Plano de Custeio da Seguridade Social para recolhimento da contribuição referente ao mês imediatamente posterior ao do final dos prazos fixados neste artigo e seus parágrafos.",
                "Todas as pessoas que contribuem para a Previdência Social, seja pelo desconto da contribuição no salário, como no caso dos trabalhadores com carteira assinada, seja pelo recolhimento por meio de guia, como fazem os autônomos são consideradas segurados. Quando param de contribuir por um determinado período, perdem a qualidade de segurado e, portanto, deixam de ter direito aos benefícios pagos pela Previdência Social, mesmo que cumpram as outras exigências como, por exemplo, estar incapacitadas para o trabalho no caso do auxílio-doença.",
                "A carência é o número mínimo de contribuições mensais indispensáveis para que o beneficiário faça jus ao benefício, conforme art. 24, da Lei 8.213/91. Para que o segurado possa ser contemplado com um dos benefícios previdenciários deve satisfazer a carência exigida para o benefício pretendido. É a carência, portanto, o instituto que qualifica o segurado como apto ao percebimento de benefícios previdenciários. Para obtenção do benefício previdenciário o segurado deve transpor a barreira da carência, que é exigência típica do seguro privado.",
                "Art. 24. Período de carência é o número mínimo de contribuições mensais indispensáveis para que o beneficiário faça jus ao benefício, consideradas a partir do transcurso do primeiro dia dos meses de suas competências.",
                "Em algumas hipóteses (art. 26, inc. II, da Lei n.º 8.213/91), dispensa-se a carência, ou seja, há uma lista de doenças que dispensam a exigência de carência para auxílio-doença e aposentadoria por Invalidez, conforme disposto no artigo 26, da Lei 8.213/91:",
                "Art. 26. Independe de carência a concessão das seguintes prestações:",
                "II - auxílio-doença e aposentadoria por invalidez nos casos de acidente de qualquer natureza ou causa e de doença profissional ou do trabalho, bem como nos casos de segurado que, após filiar-se ao RGPS, for acometido de alguma das doenças e afecções especificadas em lista elaborada pelos Ministérios da Saúde e da Previdência Social, atualizada a cada 3 (três) anos, de acordo com os critérios de estigma, deformação, mutilação, deficiência ou outro fator que lhe confira especificidade e gravidade que mereçam tratamento particularizado;",
                "A lista consta no art. 151 da lei 8.213/91 e no anexo XLV, da IN 77/2015 e, atualmente, conta com as seguintes enfermidades:",
                "Art. 151 da Lei 8.213/91: Até que seja elaborada a lista de doenças mencionada no inciso II do art. 26, independe de carência a concessão de auxílio-doença e de aposentadoria por invalidez ao segurado que, após filiar-se ao RGPS, for acometido das seguintes doenças: tuberculose ativa, hanseníase, alienação mental, esclerose múltipla, hepatopatia grave, neoplasia maligna, cegueira, paralisia irreversível e incapacitante, cardiopatia grave, doença de Parkinson, espondiloartrose anquilosante, nefropatia grave, estado avançado da doença de Paget (osteíte deformante), síndrome da deficiência imunológica adquirida (aids) ou contaminação por radiação, com base em conclusão da medicina especializada.",
                "Essas doenças, conforme disposto no artigo, dispensam o segurado do regime geral de previdência a cumprir a carência normalmente exigida para a concessão da aposentadoria.",
                "Em caso de perda da qualidade de segurado, para voltar a ter direito aos benefícios, o(a) trabalhador(a) que perdeu a qualidade de segurado(a) terá de contribuir para a Previdência por, pelo menos, com a metade dos períodos previstos nos incisos I, III e IV do caput do art. 25, da Lei 8.213/91 – é a chamada carência – para cada tipo de benefício. No caso do auxílio-doença, por exemplo, a carência exigida é de 12 meses. Mas, para alguém que perdeu a qualidade de segurado(a), as contribuições anteriores só serão consideradas para a concessão do auxílio-doença se, depois de voltar a contribuir, houver, pelo menos, 06 (seis) novas contribuições, conforme art. 27-A, da Lei 8.213/91.",
                "Art. 27-A Na hipótese de perda da qualidade de segurado, para fins da concessão dos benefícios de auxílio-doença, de aposentadoria por invalidez, de salário-maternidade e de auxílio-reclusão, o segurado deverá contar, a partir da data da nova filiação à Previdência Social, com metade dos períodos previstos nos incisos I, III e IV do caput do art. 25 desta Lei.",
                "Para que o segurado possa ser contemplado com um dos benefícios previdenciários deve satisfazer a carência exigida para o benefício pretendido. É a carência, portanto, o instituto que qualifica o(a) segurado(a) como apto(a) ao percebimento de benefícios previdenciários. Para obtenção do benefício previdenciário o(a) segurado(a) deve transpor a barreira da carência, que é exigência típica do seguro privado.",
                "Com relação as mudanças efetuadas na quantidade de carência a partir de 08.07.2016, com a primeira edição e publicação da Medida Provisória nº 739, de 07.07.2016, verifica-se uma variação nos números de contribuição necessárias para reaquisição do direito à contagem dos meses anteriores, conforme as MP´s publicadas: até 07/07/2016 – 04 contribuições; de 08/07/2016 a 04/11/2016 (MP 739/2016) – 12 contribuições; de 05/11/2016 a 05/01/2017 - 04 contribuições; de 06/01/2017 a 26/06/2017 (MP 767/2017) – 12 contribuições; de 27/06/2017 a 17/01/2019 (Lei n. 13.457/2017) – 06 contribuições; de 18/01/2019 a 17/06/2019 (MP 871/2019) – 12 contribuições; a partir de 18/06/2019 (vigente – Lei n. 13.846/2019) – 06 contribuições.",
                "Anoto, também, que o benefício por incapacidade permanente não tem data de cessação predefinida, ao passo que o benefício por incapacidade temporária segue, no que se refere à cessação, o tema 246 da TNU:",
                "Tema 246: I - Quando a decisão judicial adotar a estimativa de prazo de recuperação da capacidade prevista na perícia, o termo inicial é a data da realização do exame, sem prejuízo do disposto no art. 479 do CPC, devendo ser garantido prazo mínimo de 30 dias, desde a implantação, para viabilizar o pedido administrativo de prorrogação. II - quando o ato de concessão (administrativa ou judicial) não indicar o tempo de recuperação da capacidade, o prazo de 120 dias, previsto no § 9º, do art. 60 da Lei 8.213/91, deve ser contado a partir da data da efetiva implantação ou restabelecimento do benefício no sistema de gestão de benefícios da autarquia.",
                "Quanto ao mais, eventual exercício de trabalho remunerado, ao tempo que esteve comprovadamente incapacitado, não impede o segurado de recebe o benefício. O Superior Tribunal de Justiça já definiu sobre a possibilidade de recebimento de benefício por incapacidade do Regime Geral de Previdência Social de caráter substitutivo da renda (auxílio-doença ou aposentadoria por invalidez) concedido judicialmente em período de abrangência concomitante ao que o segurado estava trabalhando e aguardando o deferimento do benefício. A seguinte tese foi firmada (Tema 1013 STJ):",
                "Tema 1013: No período entre o indeferimento administrativo e a efetiva implantação de auxílio-doença ou de aposentadoria por invalidez, mediante decisão judicial, o segurado do RPGS tem direito ao recebimento conjunto das rendas do trabalho exercido, ainda que incompatível com sua incapacidade laboral, e do respectivo benefício previdenciário pago retroativamente.",
                "Por fim, é importante mencionar que neste modelo de benefício, é possível que a perícia constate a existência de incapacidade permanente, mas não total, de forma que o segurado possa vir a exercer outra função para a qual for reabilitado. Na forma do art. 89 da Lei n. 8.213/91: A habilitação e a reabilitação profissional e social deverão proporcionar ao beneficiário incapacitado parcial ou totalmente para o trabalho, e às pessoas portadoras de deficiência, os meios para a (re)educação e de (re)adaptação profissional e social indicados para participar do mercado de trabalho e do contexto em que vive.",
                "Neste panorama, é necessário que se analise a possibilidade de aplicação da súmula 47 da TNU, assim redigida:",
                "Súmula 47 TNU: Uma vez reconhecida a incapacidade parcial para o trabalho, o juiz deve analisar as condições pessoais e sociais do segurado para a concessão de aposentadoria por invalidez.",
                "Também se mostra possível ao Juízo, entendendo o caso, a aplicação da Tema 177 da TNU, restabelecendo/concedendo o benefício por incapacidade temporária (auxílio-doença), e encaminhando o segurado para análise administrativa de elegibilidade à reabilitação profissional:",
                "Tema 177: 1. Constatada a existência de incapacidade parcial e permanente, não sendo o caso de aplicação da Súmula 47 da TNU, a decisão judicial poderá determinar o encaminhamento do segurado para análise administrativa de elegibilidade à reabilitação profissional, sendo inviável a condenação prévia à concessão de aposentadoria por invalidez condicionada ao insucesso da reabilitação; 2. A análise administrativa da elegibilidade à reabilitação profissional deverá adotar como premissa a conclusão da decisão judicial sobre a existência de incapacidade parcial e permanente, ressalvada a possibilidade de constatação de modificação das circunstâncias fáticas após a sentença.",
                "Ainda sobre a possibilidade de reabilitação profissional, importante destacar o tema 272 da TNU:",
                "Tema 272: A circunstância de a recuperação da capacidade depender de intervenção cirúrgica não autoriza, automaticamente, a concessão de aposentadoria por invalidez (aposentadoria por incapacidade permanente), sendo necessário verificar a inviabilidade de reabilitação profissional, consideradas as condições pessoais do segurado, e a sua manifestação inequívoca a respeito da recusa ao procedimento cirúrgico.",
                "Feitas essas premissas, passo a analisar o caso concreto."
                ]
else:
    fundamento_custom = st.text_area(
            "Redija, ou copie e cole, a fundamentação que deseja inserir na sentença. \n\nO texto deve englobar tudo, desde o 'vistos em sentença' até um parágrafo assim redigido: 'Feitas estas considerações, passo a analisar o caso concreto', ou expressão equivalente."
            )

resultado = st.radio("Resultado:", [1, 2], format_func=lambda x: "Procedente" if x == 1 else "Improcedente")

# IMPROCEDÊNCIA==============================================================
if resultado == 2:
    motivo = st.radio("Motivo:", [1, 2],
                    format_func=lambda x: "Ausência de incapacidade" if x == 1 else "Falta qualidade de segurado / carência / incapacidade preexistente")

    if motivo == 1:
        detalhe = st.radio("Por que não existe incapacidade?", [1, 2],
                            format_func=lambda x: "Apenas texto básico da sentença sobre ausencia de incapacidade" if x == 1 else "Além do texto básico, quero redigir maiores esclarecimentos")
        if detalhe == 2:
            ausencia_incapacidade_redigido = st.text_area("Redija o motivo (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final):")
    else:
        sem_QS_carencia = st.text_area("Redija o motivo da falta de qualidade de segurado / carência / incapacidade preexistente (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final):")

    if st.button("Gerar Improcedência"):
        doc = Document()
        doc.add_paragraph(f"Processo: {processo_formatado}")
        texto_base(doc, fundamento_questao)
        if motivo == 1:
            fundamento_improcedente = [
                    (f"Realizada perícia para constatação de incapacidade, em sede judicial, foi verificado que a parte autora não apresenta incapacidade para o trabalho. A prova técnica, portanto, resultou na demonstração de que está ausente o requisito de incapacidade."),
                    (f"O caso não comporta que se produza outro laudo pericial, ou que se exija outros esclarecimentos do perito. O perito é claro em seu laudo. "),
                    (f"A simples constatação da doença não significa, necessariamente, que ela seja incapacitante. A função da perícia judicial é verificar a capacidade, ou não, da parte autora para o trabalho."),
                    (f"Ademais, não se visa na presente demanda a realização de exames médicos para analisar toda a condição de saúde da parte autora, na incessante busca por motivo que leve a concessão de benefício. Trata-se de exame pericial baseado na causa de pedir, apenas."),
                    ]                
            if detalhe == 2:
                for linha in ausencia_incapacidade_redigido.split('\n'):
                    if linha.strip():
                        fundamento_improcedente.append(linha)            
        else:
            fundamento_improcedente = []
            for linha in sem_QS_carencia.split('\n'):
                if linha.strip():
                    fundamento_improcedente.append(linha)
        fundamento_improcedente.extend([
                        (f"Não cumprido um dos requisitos legais, o pedido é improcedente."),
                        (f"Isto posto, com resolução de mérito nos termos do art. 487, I, do CPC, JULGO IMPROCEDENTE o pedido."),
                        (f"Sem condenação em honorários nesta instância."),
                        (f"Defiro os benefícios da gratuidade."),
                        (f"Em caso de interposição de embargos de declaração, intime-se a parte contrária para contrarrazões no prazo legal, e tornem conclusos para julgamento."),
                        (f"Interposto recurso, intime-se a parte contrária para contrarrazões no prazo legal. Após, remetam-se os autos às Turmas Recursais."),
                        (f"Com o trânsito em julgado, arquivem-se oportunamente."),
                        (f"Int."),
                        ])
        ft.alinhamento_parag_dispositivo(doc, fundamento_improcedente)
        ft.salvar_docx_temporario(doc, processo_formatado)

# PROCEDÊNCIA===================================================================
else:
    tipo = st.radio("Benefício a ser concedido:", [1, 2, 3],
                    format_func=lambda x: ["Aposentadoria invalidez", "Auxílio-doença",
                                            "Auxílio-doença + Reabilitação"][x-1])

    incapacidade_tempo = st.radio("A perícia constatou incapacidade:", [1, 2],
                    format_func = lambda x: ["temporária", "permanente"] [x-1])
    if incapacidade_tempo == 1:
        incapacidade_temporaria_ou_permanente = "temporária"
    else:
        incapacidade_temporaria_ou_permanente = "permanente"

    incapacidade_extensao = st.radio("A perícia constatou incapacidade:", [1, 2],
                    format_func = lambda x: ["total", "parcial"] [x-1])
    if incapacidade_extensao == 1:
        incapacidade_total_ou_parcial = "total"
    else:
        incapacidade_total_ou_parcial = "parcial"

    grande_invalidez = st.checkbox("Grande invalidez (+25%)") if tipo == 1 else None
    if grande_invalidez:
        grande_invalidez_redigido = "A perícia constatou necessidade de assistência permanente de outra pessoa, de modo que a parte autora faz jus ao acréscimo de 25% a que se refere o artigo 45 da Lei n. 8.213/91, em relação ao benefício por incapacidade permanente."

    dii = st.text_input("DII (dd/mm/aaaa):")
    motivo_DII = st.radio("Qual o motivo desta DII?", [1,2],
                format_func = lambda x: "DII fixada expressamente na perícia" if x == 1 else "Quero redigir esclarecimentos")
    if motivo_DII == 2:
        motivo_DII_redigido = st.text_area("Esclareça a DII escolhida (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final):")
    else:
        motivo_DII_redigido = "DII fixada expressamente na perícia."

    qualidade = st.text_area("Explique a qualidade de segurado na DII (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final):")
    carencia = st.text_area("Explique a carência na DII (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final):")

    if tipo == 1:
        por_que_concedeu_ap_invalidez = st.radio ("Por que foi concedida aposentadoria por invalidez?", [1, 2],
                                                format_func = lambda x: ["Sendo a incapacidade total e permanente, satisfaz a parte autora o requisito de incapacidade hábil à concessão de benefício por incapacidade permanente.", "Mesmo sendo a incapacidade permanente, mas clinicamente parcial, percebe-se que a parte autora, dada sua situação concreta de instrução e idade, é insusceptível de reabilitação, impedida de prover sua manutenção (súmula 47 TNU), fazendo jus ao benefício por incapacidade permanente."] [x-1])
        if por_que_concedeu_ap_invalidez == 2:
            observacao_sobre_aposentadoria_invalidez = "Mesmo sendo a incapacidade permanente, mas clinicamente parcial, percebe-se que a parte autora, dada sua situação concreta de instrução e idade, é insusceptível de reabilitação, impedida de prover sua manutenção (súmula 47 TNU), fazendo jus ao benefício por incapacidade permanente."
            explicando_sumula_47 = st.text_area("Explique no caso concreto, por qual motivo aplicou-se a súmula 47 da TNU: (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final)")
        else:
            observacao_sobre_aposentadoria_invalidez = "Sendo a incapacidade total e permanente, satisfaz a parte autora o requisito de incapacidade hábil à concessão de benefício por incapacidade permanente."

    if tipo == 2:
        por_que_concedeu_auxilio_doenca = st.radio ("Por que foi concedido auxílio-doença?", [1, 2],
                                        format_func = lambda x: ["Sendo a incapacidade temporária e total, satisfaz a parte autora o requisito de incapacidade hábil à concessão de benefício por incapacidade temporária.", "Mesmo sendo a incapacidade temporária e parcial, verifica-se pelo conjunto probatório que a parte não pode exercer suas atividades habituais, de modo a fazer jus ao benefício por incapacidade temporária. Não se trata de incapacidade total e tampouco permanente, que possa resultar na concessão de outro benefício."] [x-1])
        if por_que_concedeu_auxilio_doenca == 1:
            observacao_sobre_auxilio_doenca = "Sendo a incapacidade temporária e total, satisfaz a parte autora o requisito de incapacidade hábil à concessão de benefício por incapacidade temporária."
        else:
            observacao_sobre_auxilio_doenca = "Mesmo sendo a incapacidade temporária e parcial, verifica-se pelo conjunto probatório que a parte não pode exercer suas atividades habituais, de modo a fazer jus ao benefício por incapacidade temporária. Não se trata de incapacidade total e tampouco permanente, que possa resultar na concessão de outro benefício."

    if tipo == 3:
        por_que_reabilitacao = st.text_area("A concessão de auxílio-doença e encaminhamento para reabilitação advém de incapacidade parcial e permanente, e pelo fato de que não pode ser aplicada a súmula 47 da TNU (que determina a concessão de benefício de aposentadoria por invalidez, diante da idade, condições pessoais, etc. do segurado). \n\nExplique por que não pode ser concedida aposentadoria por invalidez no caso concreto (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final).")

    dib = st.text_input("DIB (dd/mm/aaaa):")
    motivo_DIB = st.radio("DIB fixada na DER?", [1,2,3],
                    format_func = lambda x: "DIB na DER" if x == 1 else "DIB fixada no dia seguinte à cessação de benefício por incapacidade anterior" if x == 2 else "DIB fixada em outra data (necessário esclarecer)")
    if motivo_DIB == 3:
        motivo_DIB_redigido = st.text_area("Esclareça a DIB escolhida (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final):")
    elif motivo_DIB == 2:
        motivo_DIB_redigido = "DIB fixada no dia seguinte à cessação de benefício por incapacidade anterior."
    else:
        motivo_DIB_redigido = "DIB fixada na DER."

    dcb = st.text_input("DCB (dd/mm/aaaa):") if tipo == 2 else None

    motivo_da_DCB = st.radio("Explique a DCB fixada", [1, 2],
                    format_func = lambda x: ["Para fixação da DCB foi adotada a estimativa de prazo de recuperação de capacidade prevista na perícia, com termo inicial na data do exame, na forma do Tema 246 da TNU. O INSS deverá garantir o mínimo de 30 dias de manutenção do benefício, desde a implantação, para viabilizar o pedido administrativo de prorrogação. Compete à parte acompanhar a implantação e prazo para eventual prorrogação.", "Para fixação da DCB foi adotado o prazo de 120 dias a contar da efetiva implantação, na forma do Tema 246 da TNU. O INSS deverá garantir o mínimo de 30 dias de manutenção do benefício, desde a implantação, para viabilizar o pedido administrativo de prorrogação. Compete à parte acompanhar a implantação e prazo para eventual prorrogação."] [x-1]) if tipo == 2 else None
    if motivo_da_DCB == 1:
        motivo_da_DCB_redigido = "Para fixação da DCB foi adotada a estimativa de prazo de recuperação de capacidade prevista na perícia, com termo inicial na data do exame, na forma do Tema 246 da TNU. O INSS deverá garantir o mínimo de 30 dias de manutenção do benefício, desde a implantação, para viabilizar o pedido administrativo de prorrogação. Compete à parte acompanhar a implantação e prazo para eventual prorrogação."
    else:
        motivo_da_DCB_redigido = "Para fixação da DCB foi adotado o prazo de 120 dias a contar da efetiva implantação, na forma do Tema 246 da TNU. O INSS deverá garantir o mínimo de 30 dias de manutenção do benefício, desde a implantação, para viabilizar o pedido administrativo de prorrogação. Compete à parte acompanhar a implantação e prazo para eventual prorrogação."

    procedencia_total_parcial = st.radio("Para fins da redação do dispositivo da sentença, a procedência foi total ou parcial?", [1, 2], format_func=lambda x: "Total" if x == 1 else "Parcial")
    if procedencia_total_parcial == 1:
        resultado_dispositivo = ""
    else:
        resultado_dispositivo = "EM PARTE "
            
    if st.button("Gerar Procedência"):
        data_atual = datetime.now()
        DIP = data_atual.strftime("01/%m/%Y")
        try:
            # Validação de datas
            datetime.strptime(dii, "%d/%m/%Y")
            datetime.strptime(dib, "%d/%m/%Y")
            if dcb: datetime.strptime(dcb, "%d/%m/%Y")
        except ValueError:
            st.error("Formato de data inválido!")
        doc = Document()
        doc.add_paragraph(f"Processo: {processo_formatado}")
        texto_base(doc, fundamento_questao)    
        if tipo == 1:
            beneficio_concedido = "benefício por incapacidade permanente (aposentadoria por invalidez)"
            fundamentacao_tipo1 = [
                    (f"Da perícia médica podemos verificar que o perito concluiu que há incapacidade {incapacidade_total_ou_parcial} e {incapacidade_temporaria_ou_permanente}. Fixo a DII em {dii}. {motivo_DII_redigido}"),
                    (f"Na DII considerada podemos concluir, sobre a qualidade de segurado e carência:"),
                    (f"Qualidade de segurado - {qualidade}"),
                    (f"Carência - {carencia}"),
                    (f"A prova técnica produzida no processo é determinante em casos em que a incapacidade somente pode ser aferida por médico perito, profissional habilitado a fornecer ao Juízo elementos técnicos para formação de sua convicção."),
                    (f"Afasto qualquer necessidade de complementação do laudo, ou necessidade de resposta a novos quesitos. O laudo é claro em sua conclusão, e não há imprecisões que o comprometa ou infirmem suas conclusões. Os peritos deste Juízo são profissionais equidistantes das partes e com habilidades técnicas necessárias para a aferição quanto à existência ou não de incapacidade da parte autora, não há razões para que o laudo médico pericial seja recusado. Ademais, o laudo pericial foi emitido com base no quadro clínico verificado por ocasião da(s) perícia(s) médica(s), através de exames físicos, bem como na história clínica, através dos exames apresentados pela parte autora na data da sua perícia judicial."),
                    (f"{observacao_sobre_aposentadoria_invalidez}")
            ]
            if por_que_concedeu_ap_invalidez == 2:
                for linha in explicando_sumula_47.split('\n'):
                    if linha.strip():
                        fundamentacao_tipo1.append(linha)    
            fundamentacao_tipo1.extend([
                    (f"Fixo a DIB em {dib}. {motivo_DIB_redigido}"),    
                    (f"Isto posto, com resolução de mérito nos termos do art. 487, I, do CPC, JULGO PROCEDENTE {resultado_dispositivo}o pedido para conceder o {beneficio_concedido} nos seguintes parâmetros:"),
                    (f"Benefício: {beneficio_concedido}"),
                    (f"NB: a ser definida pelo INSS"),
                    (f"DII: {dii}"),
                    (f"DIB: {dib}"),
                    (f"RMI e RMA: a serem calculadas pelo INSS." + (f"{grande_invalidez_redigido}" if grande_invalidez else "")),
                    (f"DIP: {DIP}"),
                    (f"Condeno o INSS ao pagamento dos atrasados devidos desde a DIB fixada, até a DIP em {DIP}, atualizados desde cada competência devida e com juros desde a propositura da demanda, pelos índices e percentuais do Manual de Cálculos da Justiça Federal, a ser apurado em cumprimento invertido de sentença."),
                    (f"Fica autorizado o desconto de eventuais valores recebidos a título de benefícios inacumuláveis."),
                    (f"Condeno o INSS ao ressarcimento dos honorários periciais antecipados pela Justiça Federal (art. 82, § 2º, do CPC)."),
                    (f"Considerando que o momento da prolação de sentença é oportuno para distribuir o ônus do tempo do processo, com vistas a salvaguardar a eficácia do princípio constitucional da razoável duração do processo e ao mesmo tempo privilegiar o direito provável em detrimento do improvável, demonstrada a verossimilhança das alegações da parte autora e diante do nítido caráter alimentar da verba pleiteada, nos termos do art. 294 e 300, do CPC ANTECIPA A TUTELA JURISDICIONAL para determinar ao INSS que providencie a implantação do {beneficio_concedido} com data de início de pagamento em {DIP} (DIP)."),
                    (f"O INSS deverá providenciar a implantação do benefício previdenciário ora concedido no prazo legal, sendo a contagem em dias úteis, sendo que constitui ônus das partes informar ao Juízo sobre a efetiva implantação do benefício ou eventual descumprimento do prazo pelo INSS/APSADJ."),
                    (f"Sem condenação em honorários nesta instância."),
                    (f"Defiro os benefícios da gratuidade."),
                    (f"Em caso de interposição de embargos de declaração, intime-se a parte contrária para contrarrazões no prazo legal, e tornem conclusos para julgamento."),
                    (f"Interposto recurso, intime-se a parte contrária para contrarrazões no prazo legal. Após, remetam-se os autos às Turmas Recursais."),
                    (f"Com o trânsito em julgado, implantado o benefício, dê-se início ao cumprimento de sentença."),
                    (f"Proceda a Secretaria como necessário."),
                    (f"Int."),
                    ])
            ft.alinhamento_parag_dispositivo(doc, fundamentacao_tipo1)
            
        if tipo == 2:
            beneficio_concedido = "benefício por incapacidade temporária (auxílio-doença)"
            fundamentacao_tipo2 = [
            (f"Da perícia médica podemos verificar que o perito concluiu que há incapacidade {incapacidade_total_ou_parcial} e {incapacidade_temporaria_ou_permanente}. Fixo a DII em {dii}. {motivo_DII_redigido}"),
            (f"Na DII considerada podemos concluir, sobre a qualidade de segurado e carência:"),
            (f"Qualidade de segurado - {qualidade}"),
            (f"Carência - {carencia}"),
            (f"A prova técnica produzida no processo é determinante em casos em que a incapacidade somente pode ser aferida por médico perito, profissional habilitado a fornecer ao Juízo elementos técnicos para formação de sua convicção."),
            (f"Afasto qualquer necessidade de complementação do laudo, ou necessidade de resposta a novos quesitos. O laudo é claro em sua conclusão, e não há imprecisões que o comprometa ou infirmem suas conclusões. Os peritos deste Juízo são profissionais equidistantes das partes e com habilidades técnicas necessárias para a aferição quanto à existência ou não de incapacidade da parte autora, não há razões para que o laudo médico pericial seja recusado. Ademais, o laudo pericial foi emitido com base no quadro clínico verificado por ocasião da(s) perícia(s) médica(s), através de exames físicos, bem como na história clínica, através dos exames apresentados pela parte autora na data da sua perícia judicial."),
            (f"{observacao_sobre_auxilio_doenca}"),
            (f"Fixo a DIB em {dib}. {motivo_DIB_redigido}"),
            (f"{motivo_da_DCB_redigido}"),
            (f"Isto posto, com resolução de mérito nos termos do art. 487, I, do CPC, JULGO PROCEDENTE {resultado_dispositivo}o pedido para conceder o {beneficio_concedido} nos seguintes parâmetros:"),
            (f"Benefício: {beneficio_concedido}"),
            (f"NB: a ser definida pelo INSS"),
            (f"DII: {dii}"),
            (f"DIB: {dib}"),
            (f"RMI e RMA: a serem calculadas pelo INSS"),
            (f"DCB: {dcb}"),
            (f"DIP: {DIP}"),
            (f"Condeno o INSS ao pagamento dos atrasados devidos desde a DIB fixada, até a DIP em {DIP}, atualizados desde cada competência devida e com juros desde a propositura da demanda, pelos índices e percentuais do Manual de Cálculos da Justiça Federal, a ser apurado em cumprimento invertido de sentença."),
            (f"Fica autorizado o desconto de eventuais valores recebidos a título de benefícios inacumuláveis."),
            (f"Condeno o INSS ao ressarcimento dos honorários periciais antecipados pela Justiça Federal (art. 82, § 2º, do CPC)."),
            (f"Considerando que o momento da prolação de sentença é oportuno para distribuir o ônus do tempo do processo, com vistas a salvaguardar a eficácia do princípio constitucional da razoável duração do processo e ao mesmo tempo privilegiar o direito provável em detrimento do improvável, demonstrada a verossimilhança das alegações da parte autora e diante do nítido caráter alimentar da verba pleiteada, nos termos do art. 294 e 300, do CPC ANTECIPA A TUTELA JURISDICIONAL para determinar ao INSS que providencie a implantação do {beneficio_concedido} com data de início de pagamento em {DIP} (DIP)."),
            (f"O INSS deverá providenciar a implantação do benefício previdenciário ora concedido no prazo legal, sendo a contagem em dias úteis, sendo que constitui ônus das partes informar ao Juízo sobre a efetiva implantação do benefício ou eventual descumprimento do prazo pelo INSS/APSADJ."),
            (f"O INSS deverá garantir o mínimo de 30 dias de manutenção do benefício, desde a implantação, para viabilizar o pedido administrativo de prorrogação, mesmo nas hipóteses em que a DCB fixada na sentença seja anterior à data de sua prolatação. Compete à parte acompanhar a implantação e prazo para eventual prorrogação, não havendo intimação por este Juízo."),
            (f"Sem condenação em honorários nesta instância."),
            (f"Defiro os benefícios da gratuidade."),
            (f"Em caso de interposição de embargos de declaração, intime-se a parte contrária para contrarrazões no prazo legal, e tornem conclusos para julgamento."),
            (f"Interposto recurso, intime-se a parte contrária para contrarrazões no prazo legal. Após, remetam-se os autos às Turmas Recursais."),
            (f"Com o trânsito em julgado, implantado o benefício, dê-se início ao cumprimento de sentença."),
            (f"Proceda a Secretaria como necessário."),
            (f"Int."),
            ]
            ft.alinhamento_parag_dispositivo(doc, fundamentacao_tipo2)
                
        if tipo == 3:
            beneficio_concedido = "benefício por incapacidade temporária (auxilio doença), com encaminhamento ao serviço de reabilitação profissional"
            fundamentacao_tipo3 = [
            (f"Da perícia médica podemos verificar que o perito concluiu que há incapacidade {incapacidade_total_ou_parcial} e {incapacidade_temporaria_ou_permanente}. Fixo a DII em {dii}. {motivo_DII_redigido}"),
            (f"Na DII considerada podemos concluir, sobre a qualidade de segurado e carência:"),
            (f"Qualidade de segurado - {qualidade}"),
            (f"Carência - {carencia}"),
            (f"A prova técnica produzida no processo é determinante em casos em que a incapacidade somente pode ser aferida por médico perito, profissional habilitado a fornecer ao Juízo elementos técnicos para formação de sua convicção."),
            (f"Afasto qualquer necessidade de complementação do laudo, ou necessidade de resposta a novos quesitos. O laudo é claro em sua conclusão, e não há imprecisões que o comprometa ou infirmem suas conclusões. Os peritos deste Juízo são profissionais equidistantes das partes e com habilidades técnicas necessárias para a aferição quanto à existência ou não de incapacidade da parte autora, não há razões para que o laudo médico pericial seja recusado. Ademais, o laudo pericial foi emitido com base no quadro clínico verificado por ocasião da(s) perícia(s) médica(s), através de exames físicos, bem como na história clínica, através dos exames apresentados pela parte autora na data da sua perícia judicial."),
            (f"Diante do cenário do caso concreto, restando comprovada que a incapacidade laboral da parte autora o benefício que deve ser concedido é o {beneficio_concedido}"),
            (f"O tema 177 da TNU é claro determinar o encaminhamento do segurado incapaz parcialmente para a reabilitação, quando não é o caso de concessão de benefício por incapacidade permanente. É este o caso dos autos."),
            ]
            for linha in por_que_reabilitacao.split("\n"):
                if linha.strip():
                    fundamentacao_tipo3.append(linha)
            fundamentacao_tipo3.extend([    
            (f"Fixo a DIB em {dib}. {motivo_DIB_redigido}"),
            (f"Isto posto, com resolução de mérito nos termos do art. 487, I, do CPC, JULGO PROCEDENTE {resultado_dispositivo}o pedido para conceder o {beneficio_concedido} nos seguintes parâmetros:"),
            (f"Benefício: {beneficio_concedido}"),
            (f"NB: a ser definida pelo INSS"),
            (f"DII: {dii}"),
            (f"DIB: {dib}"),
            (f"RMI e RMA: a serem calculadas pelo INSS"),
            (f"DCB: o benefício deve ser mantido enquanto não for decidido sobre a elegibilidade da parte segurada ao serviço de reabilitação, e, se o caso, enquanto durar a própria reabilitação"),
            (f"DIP: {DIP}"),
            (f"Condeno o INSS ao pagamento dos atrasados devidos desde a DIB fixada, até a DIP em {DIP}, atualizados desde cada competência devida e com juros desde a propositura da demanda, pelos índices e percentuais do Manual de Cálculos da Justiça Federal, a ser apurado em cumprimento invertido de sentença."),
            (f"Fica autorizado o desconto de eventuais valores recebidos a título de benefícios inacumuláveis."),
            (f"Condeno o INSS ao ressarcimento dos honorários periciais antecipados pela Justiça Federal (art. 82, § 2º, do CPC)."),
            (f"Considerando que o momento da prolação de sentença é oportuno para distribuir o ônus do tempo do processo, com vistas a salvaguardar a eficácia do princípio constitucional da razoável duração do processo e ao mesmo tempo privilegiar o direito provável em detrimento do improvável, demonstrada a verossimilhança das alegações da parte autora e diante do nítido caráter alimentar da verba pleiteada, nos termos do art. 294 e 300, do CPC ANTECIPA A TUTELA JURISDICIONAL para determinar ao INSS que providencie a implantação do {beneficio_concedido} com data de início de pagamento em {DIP} (DIP)."),
            (f"O INSS deverá providenciar a implantação do benefício previdenciário ora concedido no prazo legal, sendo a contagem em dias úteis, sendo que constitui ônus das partes informar ao Juízo sobre a efetiva implantação do benefício ou eventual descumprimento do prazo pelo INSS/APSADJ."),
            (f"Sem condenação em honorários nesta instância."),
            (f"Defiro os benefícios da gratuidade."),
            (f"Em caso de interposição de embargos de declaração, intime-se a parte contrária para contrarrazões no prazo legal, e tornem conclusos para julgamento."),
            (f"Interposto recurso, intime-se a parte contrária para contrarrazões no prazo legal. Após, remetam-se os autos às Turmas Recursais."),
            (f"Com o trânsito em julgado, implantado o benefício, dê-se início ao cumprimento de sentença."),
            (f"Proceda a Secretaria como necessário."),
            (f"Int."),
            ])
            ft.alinhamento_parag_dispositivo(doc, fundamentacao_tipo3)
            
        ft.salvar_docx_temporario(doc, processo_formatado)
