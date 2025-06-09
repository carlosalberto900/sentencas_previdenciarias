from datetime import datetime, date
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import tempfile
import platform
import streamlit as st
import funcoes_texto as ft

def texto_base(doc, fundamento_questao):
    if fundamento_questao == 1:
        for i, paragrafo in enumerate(fundamento_base):
            parag = doc.add_paragraph(paragrafo)
            parag.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY 
            if i in [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 25, 26, 27, 28, 31, 33, 59]:
                parag.paragraph_format.first_line_indent = Cm(2) 
            elif i in [13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24, 29, 30, 32, 34, 35, 36, 37, 38 ,39, 40 ,41, 42, 43, 44, 45, 46, 47 ,48 ,49, 50 ,51, 52, 53, 54, 55, 56, 57, 58]:
                parag.paragraph_format.left_indent = Cm(2)
    if fundamento_questao == 2:
        for linha in fundamento_custom.split("\n"):
            if linha.split():
                parag = doc.add_paragraph(linha.strip())
                parag.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                parag.paragraph_format.first_line_indent = Cm(2)

def tema_629(doc):
    tema_629_STJ = [
(f"Ausente início de prova material, inviável o acolhimento do pleito. Em que pese, em regra, a ausência de prova implique na improcedência do feito, a jurisprudência do STJ consolidou-se no sentido de que, nas demandas previdenciárias, a ausência de prova de tempo de contribuição deve resultar na extinção do feito sem resolução de mérito, por ausência de pressuposto processual, diante do caráter social que a lide envolve. Trata-se do Resp 1352721/SP, Rel. Ministro NAPOLEÃO NUNES MAIA FILHO, julgado em regime de recursos repetitivos (tema 629):"), 
(f"DIREITO PREVIDENCIÁRIO. RECURSO ESPECIAL REPRESENTATIVO DA CONTROVÉRSIA. ART. 543-C DO CPC. RESOLUÇÃO No. 8/STJ. APOSENTADORIA POR IDADE RURAL. AUSÊNCIA DE PROVA MATERIAL APTA A COMPROVAR O EXERCÍCIO DA ATIVIDADE RURAL. CARÊNCIA DE PRESSUPOSTO DE CONSTITUIÇÃO E DESENVOLVIMENTO VÁLIDO DO PROCESSO. EXTINÇÃO DO FEITO SEM JULGAMENTO DO MÉRITO, DE MODO QUE A AÇÃO PODE SER REPROPOSTA, DISPONDO A PARTE DOS ELEMENTOS NECESSÁRIOS PARA COMPROVAR O SEU DIREITO. RECURSO ESPECIAL DO INSS DESPROVIDO."), 
(f"1. Tradicionalmente, o Direito Previdenciário se vale da processualística civil para regular os seus procedimentos, entretanto, não se deve perder de vista as peculiaridades das demandas previdenciárias, que justificam a flexibilização da rígida metodologia civilista, levando-se em conta os cânones constitucionais atinentes à Seguridade Social, que tem como base o contexto social adverso em que se inserem os que buscam judicialmente os benefícios previdenciários."),
(f"2. As normas previdenciárias devem ser interpretadas de modo a favorecer os valores morais da Constituição Federal/1988, que prima pela proteção do Trabalhador Segurado da Previdência Social, motivo pelo qual os pleitos previdenciários devem ser julgados no sentido de amparar a parte hipossuficiente e que, por esse motivo, possui proteção legal que lhe garante a flexibilização dos rígidos institutos processuais. Assim, deve-se procurar encontrar na hermenêutica previdenciária a solução que mais se aproxime do caráter social da Carta Magna, a fim de que as normas processuais não venham a obstar a concretude do direito fundamental à prestação previdenciária a que faz jus o segurado."), 
(f"3. Assim como ocorre no Direito Sancionador, em que se afastam as regras da processualística civil em razão do especial garantismo conferido por suas normas ao indivíduo, deve-se dar prioridade ao princípio da busca da verdade real, diante do interesse social que envolve essas demandas."), 
(f"4. A concessão de benefício devido ao trabalhador rural configura direito subjetivo individual garantido constitucionalmente, tendo a CF/88 dado primazia à função social do RGPS ao erigir como direito fundamental de segunda geração o acesso à Previdência do Regime Geral; sendo certo que o trabalhador rural, durante o período de transição, encontra-se constitucionalmente dispensado do recolhimento das contribuições, visando à universalidade da cobertura previdenciária e a inclusão de contingentes desassistidos por meio de distribuição de renda pela via da assistência social."), 
(f"5. A ausência de conteúdo probatório eficaz a instruir a inicial, conforme determina o art. 283 do CPC, implica a carência de pressuposto de constituição e desenvolvimento válido do processo, impondo a sua extinção sem o julgamento do mérito (art. 267, IV do CPC) e a consequente possibilidade de o autor intentar novamente a ação (art. 268 do CPC), caso reúna os elementos necessários à tal iniciativa."),
(f"6. Recurso Especial do INSS desprovido."), 
(f"(REsp 1352721/SP, Rel. Ministro NAPOLEÃO NUNES MAIA FILHO, CORTE ESPECIAL, julgado em 16/12/2015, DJe 28/04/2016)"), 
(f"Isto posto, sem resolução de mérito nos termos do art. 485, IV do CPC, JULGO EXTINTO O FEITO."),
(f"Sem condenação em honorários nesta instância."),
(f"Defiro os benefícios da gratuidade."),
(f"Em caso de interposição de embargos de declaração, intime-se a parte contrária para contrarrazões no prazo legal, e tornem conclusos para julgamento."),
(f"Interposto recurso, intime-se a parte contrária para contrarrazões no prazo legal. Após, remetam-se os autos às Turmas Recursais."),
(f"Com o trânsito em julgado, arquivem-se oportunamente."),
(f"Int.") 
    ]
    for i, paragrafo in enumerate(tema_629_STJ):
        parag = doc.add_paragraph(paragrafo)
        parag.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY 
        if i in [0, 9, 10, 11, 12, 13, 14, 15]:
            parag.paragraph_format.first_line_indent = Cm(2) 
        elif i in [1, 2, 3, 4, 5, 6, 7, 8]:
            parag.paragraph_format.left_indent = Cm(2)

def improcedencia(doc):
    improcedencia = [
                    (f"Não cumprido um dos requisitos legais, o pedido é improcedente."),
                    (f"Isto posto, com resolução de mérito nos termos do art. 487, I, do CPC, JULGO IMPROCEDENTE o pedido."),
                    (f"Sem condenação em honorários nesta instância."),
                    (f"Defiro os benefícios da gratuidade."),
                    (f"Em caso de interposição de embargos de declaração, intime-se a parte contrária para contrarrazões no prazo legal, e tornem conclusos para julgamento."),
                    (f"Interposto recurso, intime-se a parte contrária para contrarrazões no prazo legal. Após, remetam-se os autos às Turmas Recursais."),
                    (f"Com o trânsito em julgado, arquivem-se oportunamente."),
                    (f"Int.")
                    ]
    for linha in improcedencia:
        if linha.strip():
            parag = doc.add_paragraph(linha.strip())
            parag.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            parag.paragraph_format.first_line_indent = Cm(2)

def prazo_pensao_conjuge_companheira(data_do_obito_convertida):
    data_vigencia = date(2015, 1, 3) #vigencia da Lei 13.146
    data_atualizacao = date(2021, 1, 1) #vigencia da portaria que aumentou um ano
    tempo_casamento_uniao = ""
    carencia_instituidor = ""
    tempo = ""    
    if data_do_obito_convertida < data_vigencia:
        tempo = "de forma vitalícia"
    if data_vigencia <= data_do_obito_convertida < data_atualizacao or data_atualizacao <= data_do_obito_convertida:
        idade_autor = int(st.number_input("Qual a idade da parte autora na data do óbito? (Digite apenas números):", min_value=0, max_value=150, step=1, key="idade_autor_pensao"))
        carencia_instituidor_opcoes = st.radio("O instituidor tem quantas contribuições vertidas em vida?", [1,2], format_func=lambda x: "menos de 18 contribuições" if x == 1 else "18 contribuições ou mais", index=1, key="carencia_instituidor_pensao")
        if carencia_instituidor_opcoes == 1:
            carencia_instituidor = f"menos de 18 contribuições"
        else:
            carencia_instituidor = f"18 (dezoito) contribuições, ou mais,"
        tempo_casamento_uniao_opcoes = st.radio("Qual o tempo do casamento/união estável da parte autora?", [1,2], format_func=lambda x: "menos de 02 anos" if x == 1 else "02 anos ou mais", index=1, key="tempo_casamento_uniao_pensao")
        if tempo_casamento_uniao_opcoes == 1:
            tempo_casamento_uniao = f"menos de dois anos"
        else:
            tempo_casamento_uniao = f"dois anos, ou mais,"
        if carencia_instituidor_opcoes == 1 or tempo_casamento_uniao_opcoes == 1:
            tempo = "por 4 meses"
        if carencia_instituidor_opcoes == 2 and tempo_casamento_uniao_opcoes == 2 and data_vigencia <= data_do_obito_convertida < data_atualizacao:
            if idade_autor < 21:
                tempo = "por 3 anos"
            if 21 <= idade_autor <= 26:
                tempo = "por 6 anos"
            if 27 <= idade_autor <= 29:
                tempo = "por 10 anos"
            if 30 <= idade_autor <= 40:
                tempo = "por 15 anos"
            if 41 <= idade_autor <= 43:
                tempo = "por 20 anos"
            if idade_autor >= 44:
                tempo = "de forma vitalícia"
        if carencia_instituidor_opcoes == 2 and tempo_casamento_uniao_opcoes == 2 and data_atualizacao <= data_do_obito_convertida:
            if idade_autor < 22:
                tempo = "por 3 anos"
            if 22 <= idade_autor <= 27:
                tempo = "por 6 anos"
            if 28 <= idade_autor <= 30:
                tempo = "por 10 anos"
            if 31 <= idade_autor <= 41:
                tempo = "por 15 anos"
            if 42 <= idade_autor <= 44:
                tempo = "por 20 anos"
            if idade_autor >= 45:
                tempo = "de forma vitalícia"
return tempo_casamento_uniao, carencia_instituidor, tempo

fundamento_questao = st.radio(
"**Relatório e fundamentação jurídica**\n\n"
"Toda sentença gerada possui dispensa de relatório e uma fundamentação jurídica básica, que vai até onde se inicia a análise do caso concreto.\n\n"  
"Você pode usar esta fundamentação, ou fornecer sua **própria fundamentação**, com ou sem relatório (fazendo menção à dispensa, se for o caso.)\n\n"  
"**Como você deseja prosseguir?**",
[1, 2],
format_func=lambda x: "Vou usar a fundamentação padrão deste aplicativo." if x == 1 else "Desejo fornecer a minha fundamentação."
)
if fundamento_questao == 1:
    fundamento_base = [
            (f"Vistos."),
            (f"Trata-se de pedido de pensão por morte."),
            (f"Dispensado o relatório nos termos do art. 38 da Lei n. 9.099/95."),
            (f"DECIDO."),
            (f"O feito comporta julgamento imediato."),
            (f"O Juízo é competente, pois o valor da causa é inferior a 60 salários mínimos. Não há que se falar em renúncia ao valor excedente a esta alçada."),
            (f"O tema 1030 do STJ prevê que: “Ao autor que deseje litigar no âmbito de Juizado Especial Federal Cível, é lícito renunciar, de modo expresso e para fins de atribuição de valor à causa, ao montante que exceda os 60 (sessenta) salários mínimos previstos no art. 3º, caput, da Lei 10.259/2001, aí incluídas, sendo o caso, até doze prestações vincendas, nos termos do art. 3º, § 2º, da referida lei, c/c o art. 292, §§ 1º e 2º, do CPC/2015.” Logo, a renúncia só tem sentido nos casos em que o valor da causa, apurado sem qualquer renúncia, seria superior a alçada, pois nesta hipótese a renúncia daria ensejo à parte litigar sob a competência do Juizado. Não é o caso dos autos, onde o valor da causa apontado já está abaixo da alçada, e a parte ré não aponta erro na sua apuração."),
            (f"Cumprido o estipulado no Tema 350 do STF. Partes legítimas e bem representadas. Passo ao mérito."),
            (f"Prejudicialmente, analiso a prescrição. Não há prescrição de fundo de direito, mas apenas das parcelas vencidas, na forma da súmula 85 do STJ. Estão prescritas eventuais parcelas anteriores ao quinquênio que precede a propositura da ação."),
            (f"Em relação ao mérito propriamente dito, o art. 74 da Lei n. 8.213/91 aduz que a pensão por morte é devida ao conjunto dos dependentes do segurado que falecer, aposentado ou não. Como requisito a lei exige a qualidade de segurado do falecido. "),
            (f"A qualidade de segurado mantém enquanto houver recolhimento válido de contribuição previdenciária, até o final do período de graça na forma estipulada no art. 15 da Lei n. 8.213/91."),
            (f"Prevê a súmula 416 do STJ que: “É devida a pensão por morte aos dependentes do segurado que, apesar de ter perdido essa qualidade, preencheu os requisitos legais para a obtenção de aposentadoria até a data do seu óbito”. A interpretação corrente é que é devida a pensão por morte caso comprovado que o falecido teria direito adquirido à obtenção de qualquer benefício que lhe garantisse o período de graça a que se refere o art. 15, I da Lei n. 8.213/91."),
            (f"Beneficiários da pensão por morte são os dependentes, segundo dispõe o artigo 16 da Lei n. 8.213/91:"),
            (f"Art. 16. São beneficiários do Regime Geral de Previdência Social, na condição de dependentes do segurado:"),
            (f"I - o cônjuge, a companheira, o companheiro e o filho não emancipado, de qualquer condição, menor de 21 (vinte e um) anos ou inválido ou que tenha deficiência intelectual ou mental ou deficiência grave; (Redação dada pela Lei nº 13.146, de 2015)"),
            (f"II - os pais;"),
            (f"III - o irmão não emancipado, de qualquer condição, menor de 21 (vinte e um) anos ou inválido ou que tenha deficiência intelectual ou mental ou deficiência grave; (Redação dada pela Lei nº 13.146, de 2015)"),
            (f"IV - (Revogada pela Lei nº 9.032, de 1995)"),
            (f"§ 1º A existência de dependente de qualquer das classes deste artigo exclui do direito às prestações os das classes seguintes."),
            (f"§ 2º O enteado, o menor sob tutela e o menor sob guarda judicial equiparam-se a filho, mediante declaração do segurado e desde que não possuam condições suficientes para o próprio sustento e educação. (Redação dada pela Lei nº 15.108, de 2025)"),
            (f"§ 3º Considera-se companheira ou companheiro a pessoa que, sem ser casada, mantém união estável com o segurado ou com a segurada, de acordo com o § 3º do art. 226 da Constituição Federal."),
            (f"§ 4º A dependência econômica das pessoas indicadas no inciso I é presumida e a das demais deve ser comprovada."),
            (f"§ 5º As provas de união estável e de dependência econômica exigem início de prova material contemporânea dos fatos, produzido em período não superior a 24 (vinte e quatro) meses anterior à data do óbito ou do recolhimento à prisão do segurado, não admitida a prova exclusivamente testemunhal, exceto na ocorrência de motivo de força maior ou caso fortuito, conforme disposto no regulamento. (Incluído pela Lei nº 13.846, de 2019)"),
            (f"§ 6º Na hipótese da alínea c do inciso V do § 2º do art. 77 desta Lei, a par da exigência do § 5º deste artigo, deverá ser apresentado, ainda, início de prova material que comprove união estável por pelo menos 2 (dois) anos antes do óbito do segurado. (Incluído pela Lei nº 13.846, de 2019)"),
            (f"§ 7º Será excluído definitivamente da condição de dependente quem tiver sido condenado criminalmente por sentença com trânsito em julgado, como autor, coautor ou partícipe de homicídio doloso, ou de tentativa desse crime, cometido contra a pessoa do segurado, ressalvados os absolutamente incapazes e os inimputáveis. (Incluído pela Lei nº 13.846, de 2019)"),
            (f"Além destes, é importante mencionar que o art. 76, § 2º da Lei n. 8.213/91 prevê que: “o cônjuge divorciado ou separado judicialmente ou de fato que recebia pensão de alimentos concorrerá em igualdade de condições com os dependentes referidos no inciso I do art. 16 desta Lei.”"),
            (f"Para ter direito, basta ao ex-cônjuge comprovar dependência econômica em relação ao segurado falecido, desde que anterior ao falecimento, ainda que posterior à separação. Neste sentido o tema 45 da TNU:"),
            (f"É devida pensão por morte ao ex-cônjuge que não percebe alimentos, desde que comprovada dependência econômica superveniente à separação, demonstrada em momento anterior ao óbito."),
            (f"Observe-se que se trata de ex-cônjuge, e não da manutenção de casamento e outra união, que não poderia ser reconhecida como “união estável” nesta hipótese, mas mero concubinato, para o qual o Supremo Tribunal Federal não reconhece efeitos previdenciários:"),
            (f"EMENTA Direito Previdenciário e Constitucional. Recurso extraordinário. Sistemática da repercussão geral. Tema nº 526. Pensão por morte. Rateio entre a concubina e a viúva. Convivência simultânea. Concubinato e Casamento. Impossibilidade. Recurso extraordinário provido. 1. Assentou-se no acórdão recorrido que, comprovada a convivência e a dependência econômica, faz jus a concubina à quota parte de pensão deixada por ex-combatente, em concorrência com a viúva, a contar do pedido efetivado na seara administrativa. Tal orientação, contudo, contraria a tese fixada pelo Supremo Tribunal Federal no julgamento do processo paradigma do Tema nº 529 sob a sistemática da repercussão geral, in verbis: “A preexistência de casamento ou de união estável de um dos conviventes, ressalvada a exceção do artigo 1723, § 1º, do Código Civil, impede o reconhecimento de novo vínculo referente ao mesmo período, inclusive para fins previdenciários, em virtude da consagração do dever de fidelidade e da monogamia pelo ordenamento jurídico-constitucional brasileiro”. 2. Antes do advento da Constituição de 1988, havia o emprego indistinto da expressão concubinato para qualquer relação não estabelecida sob as formalidades da lei, daí porque se falava em concubinato puro (hoje união estável) e concubinato impuro (relações duradoras com impedimento ao casamento). Erigida a união estável, pelo texto constitucional (art. 226, § 3º, da CF), ao status de entidade familiar e tendo o Código Civil traçado sua distinção em face do concubinato (art. 1.723, § 1º, c/c art. 1.521, VI e art. 1.727 do CC), os termos passaram a disciplinar situações diversas, o que não pode ser desconsiderado pelo intérprete da Constituição. 3. O art. 1.521 do Código Civil – que trata dos impedimentos para casar -, por força da legislação (art. 1.723, § 1º), também se aplica à união estável, sob claro reconhecimento de que a ela, como entidade familiar, também se assegura proteção à unicidade do vínculo. A espécie de vínculo que se interpõe a outro juridicamente estabelecido (seja casamento ou união estável) a legislação nomina concubinato (art. 1.727 do CC). Assim, a pessoa casada não pode ter reconhecida uma união estável concomitante, por força do art. 1.723, § 1º, c/c o art. 1.521, VI, do Código Civil. 4. Considerando que não é possível reconhecer, nos termos da lei civil (art. 1.723, § 1º, c/c art. 1.521, VI e art. 1.727 do Código Civil Brasileiro), a concomitância de casamento e união estável (salvo na hipótese do § 1º, art. 1.723, do CC/02), impende concluir que o concubinato – união entre pessoas impedidas de casar - não gera efeitos previdenciários. 5. A exegese constitucional mais consentânea ao telos implícito no microssistema jurídico que rege a família, entendida como base da sociedade (art. 226, caput, da CF), orienta-se pelos princípios da exclusividade e da boa-fé, bem como pelos deveres de lealdade e fidelidade que visam a assegurar maior estabilidade e segurança às relações familiares. 5. Foi fixada a seguinte tese de repercussão geral: “É incompatível com a Constituição Federal o reconhecimento de direitos previdenciários (pensão por morte) à pessoa que manteve, durante longo período e com aparência familiar, união com outra casada, porquanto o concubinato não se equipara, para fins de proteção estatal, às uniões afetivas resultantes do casamento e da união estável”. 6. Recurso extraordinário a que se dá provimento."),
            (f"(RE 883168, Relator(a): DIAS TOFFOLI, Tribunal Pleno, julgado em 03-08-2021, PROCESSO ELETRÔNICO REPERCUSSÃO GERAL - MÉRITO DJe-200  DIVULG 06-10-2021  PUBLIC 07-10-2021)"),
            (f"No que toca ao menor sob guarda, mesmo antes da Lei n. 15.108/2025, a jurisprudência do STJ já o considerava dependente para fins previdenciários:"),
            (f"Tema 732 do STJ: O menor sob guarda tem direito à concessão do benefício de pensão por morte do seu mantenedor, comprovada sua dependência econômica, nos termos do art. 33, § 3º do Estatuto da Criança e do Adolescente, ainda que o óbito do instituidor da pensão seja posterior à vigência da Medida Provisória 1.523/96, reeditada e convertida na Lei 9.528/97. Funda-se essa conclusão na qualidade de lei especial do Estatuto da Criança e do Adolescente (8.069/90), frente à legislação previdenciária."),
            (f"Por fim, verifica-se que a pensão por morte pode ser vitalícia ou temporária, na forma do art. 77 da Lei n. 8.213/91:"),
            (f"Art. 77. A pensão por morte, havendo mais de um pensionista, será rateada entre todos em parte iguais. (Redação dada pela Lei nº 9.032, de 1995)"),
            (f"§ 1º Reverterá em favor dos demais a parte daquele cujo direito à pensão cessar. (Redação dada pela Lei nº 9.032, de 1995)"),
            (f"§ 2º O direito à percepção da cota individual cessará: (Redação dada pela Lei nº 13.846, de 2019)"),
            (f"I - pela morte do pensionista; (Incluído pela Lei nº 9.032, de 1995)"),
            (f"II - para o filho, a pessoa a ele equiparada ou o irmão, de ambos os sexos, ao completar vinte e um anos de idade, salvo se for inválido ou tiver deficiência intelectual ou mental ou deficiência grave; (Redação dada pela Lei nº 13.183, de 2015)"),
            (f"III - para filho ou irmão inválido, pela cessação da invalidez; (Redação dada pela Lei nº 13.135, de 2015)"),
            (f"IV - para filho ou irmão que tenha deficiência intelectual ou mental ou deficiência grave, pelo afastamento da deficiência, nos termos do regulamento; (Redação dada pela Lei nº 13.135, de 2015)"),
            (f"V - para cônjuge ou companheiro: (Incluído pela Lei nº 13.135, de 2015)"),
            (f"a) se inválido ou com deficiência, pela cessação da invalidez ou pelo afastamento da deficiência, respeitados os períodos mínimos decorrentes da aplicação das alíneas “b” e “c”; (Incluído pela Lei nº 13.135, de 2015)"),
            (f"b) em 4 (quatro) meses, se o óbito ocorrer sem que o segurado tenha vertido 18 (dezoito) contribuições mensais ou se o casamento ou a união estável tiverem sido iniciados em menos de 2 (dois) anos antes do óbito do segurado; (Incluído pela Lei nº 13.135, de 2015)"),
            (f"c) transcorridos os seguintes períodos, estabelecidos de acordo com a idade do beneficiário na data de óbito do segurado, se o óbito ocorrer depois de vertidas 18 (dezoito) contribuições mensais e pelo menos 2 (dois) anos após o início do casamento ou da união estável: (Incluído pela Lei nº 13.135, de 2015)"),
            (f"1) 3 (três) anos, com menos de 21 (vinte e um) anos de idade; (Incluído pela Lei nº 13.135, de 2015)"),
            (f"2) 6 (seis) anos, entre 21 (vinte e um) e 26 (vinte e seis) anos de idade; (Incluído pela Lei nº 13.135, de 2015)"),
            (f"3) 10 (dez) anos, entre 27 (vinte e sete) e 29 (vinte e nove) anos de idade; (Incluído pela Lei nº 13.135, de 2015)"),
            (f"4) 15 (quinze) anos, entre 30 (trinta) e 40 (quarenta) anos de idade; (Incluído pela Lei nº 13.135, de 2015)"),
            (f"5) 20 (vinte) anos, entre 41 (quarenta e um) e 43 (quarenta e três) anos de idade; (Incluído pela Lei nº 13.135, de 2015)"),
            (f"6) vitalícia, com 44 (quarenta e quatro) ou mais anos de idade. (Incluído pela Lei nº 13.135, de 2015)"),
            (f"VI - pela perda do direito, na forma do § 1º do art. 74 desta Lei. (Incluído pela Lei nº 13.846, de 2019)"),
            (f"§ 2o-A.  Serão aplicados, conforme o caso, a regra contida na alínea “a” ou os prazos previstos na alínea “c”, ambas do inciso V do § 2o, se o óbito do segurado decorrer de acidente de qualquer natureza ou de doença profissional ou do trabalho, independentemente do recolhimento de 18 (dezoito) contribuições mensais ou da comprovação de 2 (dois) anos de casamento ou de união estável. (Incluído pela Lei nº 13.135, de 2015)"),
            (f"§ 2o-B.  Após o transcurso de pelo menos 3 (três) anos e desde que nesse período se verifique o incremento mínimo de um ano inteiro na média nacional única, para ambos os sexos, correspondente à expectativa de sobrevida da população brasileira ao nascer, poderão ser fixadas, em números inteiros, novas idades para os fins previstos na alínea “c” do inciso V do § 2o, em ato do Ministro de Estado da Previdência Social, limitado o acréscimo na comparação com as idades anteriores ao referido incremento. (Incluído pela Lei nº 13.135, de 2015)"),
            (f"§ 3º Com a extinção da parte do último pensionista a pensão extinguir-se-á.  (Incluído pela Lei nº 9.032, de 1995)"),
            (f"§ 4o  (Revogado)."),    
            (f"§ 5o  O tempo de contribuição a Regime Próprio de Previdência Social (RPPS) será considerado na contagem das 18 (dezoito) contribuições mensais de que tratam as alíneas “b” e “c” do inciso V do § 2o. (Incluído pela Lei nº 13.135, de 2015)"),
            (f"§ 6º O exercício de atividade remunerada, inclusive na condição de microempreendedor individual, não impede a concessão ou manutenção da parte individual da pensão do dependente com deficiência intelectual ou mental ou com deficiência grave. (Incluído pela Lei nº 13.183, de 2015)"),
            (f"§ 7º Se houver fundados indícios de autoria, coautoria ou participação de dependente, ressalvados os absolutamente incapazes e os inimputáveis, em homicídio, ou em tentativa desse crime, cometido contra a pessoa do segurado, será possível a suspensão provisória de sua parte no benefício de pensão por morte, mediante processo administrativo próprio, respeitados a ampla defesa e o contraditório, e serão devidas, em caso de absolvição, todas as parcelas corrigidas desde a data da suspensão, bem como a reativação imediata do benefício. (Incluído pela Lei nº 13.846, de 2019)"),
            (f"Feitas estas premissas, passo ao caso concreto."),
            ]
else:
    fundamento_custom = st.text_area(
        "Redija, ou copie e cole, a fundamentação que deseja inserir na sentença. \nO texto deve englobar tudo, desde o 'vistos em sentença' até um parágrafo assim redigido: 'Feitas estas considerações, passo a analisar o caso concreto', ou expressão equivalente."
        )
st.write("Forneça as seguintes informações, sobre o processo: ")

col1, col2 = st.columns(2)
with col1:
    instituidor = st.text_input("Qual o nome completo do segurado, falecido, instituidor?")
with col2:
    data_do_obito = st.text_input("Qual a data do óbito? Digite no formato DD/MM/AAAA")
    try:
        data_do_obito_convertida = datetime.strptime(data_do_obito, "%d/%m/%Y").date()
        st.success(f"Data válida: {data_do_obito_convertida.strftime('%d/%m/%Y')}")
    except ValueError:
        st.error("Data inválida. Use o formato DD/MM/AAAA.")
        st.stop()
resultado = st.radio("O pedido é procedente ou improcedente (ou aplicação Tema 629 STJ)?", [1, 2], format_func=lambda x: "Procedente" if x == 1 else "Improcedente (ou aplicação de Tema 629 STJ)")
#IMPROCEDENTE
if resultado == 2:
    motivo_improcedencia = st.radio("Qual o motivo da improcedência?", [1, 2], format_func=lambda x: "Falta de qualidade de segurado do(a) instituidor(a) falecido" if x==1 else "A parte autora não comprova ser dependente do falecido")
    #FALTA DE QUALIDADE DE SEGURADO DO INSTITUIDOR
    if motivo_improcedencia == 1:
        hipotese = st.radio("Qual hipótese melhor se enquadra no caso?", [1, 2, 3, 4, 5, 6], format_func=lambda x: 
                                "A parte autora apresentou sentença trabalhista meramente homologatória de acordo, ou derivada de revelia, que não pode ser reconhecida como início de prova material. Além da sentença, pode ter apresentado, ou não, outros documentos que não são início de prova material." if x == 1 else
                                "A parte autora não apresentou nenhum início de prova material, ou apresentou documentos (exceto sentença trabalhista) que não podem ser considerados início de prova material." if x == 2 else
                                "As testemunhas ouvidas em Juízo não comprovam o exercício de labor pelo falecido, como segurado empregado, avulso, doméstico ou especial." if x ==3 else
                                "O trabalho do segurado era na condição de contribuinte individual, sendo vedado o reconhecimento de qualidade de segurado 'post mortem' ao segurado falecido como contribuinte individual, com recolhimento extemporâneo (ou sem recolhimento)." if x == 4 else
                                "A parte autora não comprova que o segurado falecido tinha direito adquirido a um benefício ao tempo do óbito." if x == 5 else
                                "Outro motivo sobre falta de qualidade de segurado do falecido (a ser redigido)"
                                )
        if hipotese in [1, 2, 3, 4]:
            if hipotese in [2, 3, 4]:
                segurado_instituidor_opcoes = st.radio("Em que categoria enquadra-se o segurado instituidor?", [1, 2, 3, 4, 5, 6], format_func=lambda x: "empregado(a)" if x == 1 else "trabalhador(a) avulso" if x == 2 else "empregado(a) doméstico(a)" if x == 3 else "contribuinte individual" if x == 4 else "segurado especial" if x == 5 else "segurado facultativo")
                if segurado_instituidor_opcoes == 1:
                    segurado_instituidor = "empregado(a)"
                elif segurado_instituidor_opcoes == 2:
                    segurado_instituidor = "trabalhador(a) avulso"
                elif segurado_instituidor_opcoes == 3:
                    segurado_instituidor = "empregado(a) doméstico(a)"
                elif segurado_instituidor_opcoes == 4:
                    segurado_instituidor = "contribuinte individual"
                elif segurado_instituidor_opcoes == 5:
                    segurado_instituidor = "segurado especial"
                elif segurado_instituidor_opcoes == 6:
                    segurado_instituidor = "segurado facultativo"
            data_inicial = st.text_input("Qual a **data inicial do vínculo** que a parte autora pede reconhecimento? (responda em formato DD/MM/AAAA): ")
            data_final = st.text_input("Qual a **data final** do vínculo que a parte autora pede reconhecimento? (responda em formato DD/MM/AAAA): ")
            vinculo_controvertido = st.text_input("Qual o nome do empregador ou contratante, em relação a este vínculo controvertido? Não havendo, informar o tipo de segurado (contribuinte individual, segurado especial ou segurado facultativo) ")
        #EXTINÇÃO TEMA 629 SENTENÇA TRABALHISTA     
        if hipotese == 1:
            segurado_instituidor = (f"empregado(a)")
            ha_mais_documentos = st.radio("Além da sentença trabalhista, há outros documentos que a parte autora trouxe para serem valorados como início de prova material, que não configuram início de prova material?", [1, 2], format_func=lambda x: "Sim" if x == 1 else "Não")
            if ha_mais_documentos == 1:
                outros_documentos = st.text_area("Redija o motivo pelo qual os documentos apresentados pela parte não podem ser considerados início de prova material (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final): ")
            else:
                pass
            if st.button("Gerar Sentença"):
                doc = Document()
                doc.add_paragraph(f"Processo: {processo_formatado}")
                sentenca_homologatoria = [
                        (f"Vistos."),
                        (f"Trata-se de pedido de pensão por morte."),
                        (f"Dispensado o relatório nos termos do art. 38 da Lei n. 9.099/95."),
                        (f"DECIDO."),
                        (f"O feito comporta julgamento imediato."),
                        (f"A parte autora pede o reconhecimento do tempo de serviço/contribuição do(a) segurado(a) falecido(a), Sr(a). {instituidor}, entre {data_inicial} e {data_final} referente a (origem) {vinculo_controvertido}, como {segurado_instituidor}, para, com isso, comprovar a qualidade de segurado do(a) falecido(a) instituidor(a) ao tempo do óbito."),
                        (f"A pretensão do autor, para acolhida, deve vir acompanhada de início de prova material, a rigor do que dispõe o artigo 55, § 3º da Lei n. 8.213/91."), 
                        (f"No caso concreto, a prova material da atividade laboral que pretende ver declarada para fins previdenciários é a sentença trabalhista. A sentença trabalhista é considerada início de prova material, mas não toda e qualquer sentença. Somente a sentença de mérito baseada em prova produzida na Justiça do Trabalho. A sentença meramente homologatória de acordo, assim como a que deriva de confissão ficta (revelia) não servem como início de prova material. Demais disso, também não podem ser opostas diretamente contra o INSS, que não participou da lide e a seus termos não está vinculado."),
                        (f"""Mesmo o registro em CTPS, neste caso, não é autônomo, derivando da sentença trabalhista. A sólida jurisprudência do STJ é no sentido de que "a sentença trabalhista pode ser considerada como início de prova material, desde que prolatada com base em elementos probatórios capazes de demonstrar o exercício da atividade laborativa, durante o período que se pretende ter reconhecido na ação previdenciária." """),
                        (f"Nesse sentido:"),   
                        (f"PREVIDENCIÁRIO. PROCESSUAL CIVIL. INOVAÇÃO RECURSAL. SENTENÇA TRABALHISTA MERAMENTE HOMOLOGATÓRIA DE ACORDO. IMPRESTABILIDADE DE UTILIZAÇÃO COMO INÍCIO DE PROVA MATERIAL. JURISPRUDÊNCIA CONSOLIDADA DO STJ."),   
                        (f"1. A jurisprudência desta Corte está firmada no sentido de que a sentença trabalhista pode ser considerada como início de prova material, desde que prolatada com base em elementos probatórios capazes de demonstrar o exercício da atividade laborativa, durante o período que se pretende ter reconhecido na ação previdenciária."),   
                        (f"2. Na espécie, ao que se tem dos autos, a sentença trabalhista está fundada apenas nos depoimentos das partes, motivo pelo qual não se revela possível a sua consideração como início de prova material para fins de reconhecimento da qualidade de segurado do instituidor do benefício e, por conseguinte, como direito da parte autora à pensão por morte."),  
                        (f"3. Agravo interno a que se nega provimento."), 
                        (f"(AgInt no AREsp 1405520/SP, Rel. Ministro SÉRGIO KUKINA, PRIMEIRA TURMA, julgado em 07/11/2019, DJe 12/11/2019)"),
                        (f" "),
                        (f"PREVIDENCIÁRIO. PENSÃO POR MORTE. QUALIDADE DE SEGURADO. SENTENÇA TRABALHISTA HOMOLOGATÓRIA DE ACORDO. UTILIZAÇÃO. IMPOSSIBILIDADE."),   
                        (f"1. A jurisprudência do Superior Tribunal de Justiça está firmada no sentido de que a sentença trabalhista pode ser considerada como início de prova material, desde que prolatada com base em elementos probatórios capazes de demonstrar o exercício da atividade laborativa, durante o período que se pretende ter reconhecido na ação previdenciária. Precedentes: AgInt no AREsp 529.963/RS, Rel. Ministro Benedito Gonçalves, Primeira Turma, DJe 28.2.2019; REsp 1.758.094/RJ, Rel. Ministro Herman Benjamim, Segunda Turma, DJe 17.12.2018; e AgInt no AREsp 688.117/SP, Rel. Ministro Sérgio Kukina, Primeira Turma, DJe 11.12.2017."),
                        (f"""2. O Tribunal a quo reconheceu a qualidade de segurado do instituidor da pensão, com base na "sentença homologatória de acordo realizado em sede de Reclamação Trabalhista (fl. 110), em que foi reconhecida a relação de emprego entre o de cujus e a empresa DIVIPISO COMÉRCIO DE DIVISÓRIAS FORROS E PISOS LTDA-ME., no período de 03/05/2004 a 17/11/2005, na função de montador" (fl. 278, e-STJ) 3. Na espécie, ao que se tem dos autos, a sentença judicial trabalhista só homologou os termos de acordo entre as partes, para o reconhecimento de vínculo laboral do trabalhador já falecido, sem nenhuma incursão em matéria probatória."""),   
                        (f"4. Assim, inexistindo, quer naqueles autos da Justiça Especializada, quer nos da Justiça Federal, a produção de prova documental ou mesmo testemunhal, para se reconhecer o período de tempo em que o falecido teria trabalhado para a empresa firmatária do acordo, a sentença homologatória trabalhista é insuficiente, no caso, para embasar a pensão por morte aos dependentes do segurado."),
                        (f"5. Recurso Especial provido."),  
                        (f"(REsp 1760216/SP, Rel. Ministro HERMAN BENJAMIN, SEGUNDA TURMA, julgado em 26/03/2019, DJe 23/04/2019)"), 
                        (f" "),
                        (f"""PREVIDENCIÁRIO. APOSENTADORIA POR IDADE URBANA. ART. 25 E 48 DA LEI Nº 8.213/91. REQUISITOS LEGAIS NÃO PREENCHIDOS. AVERBAÇÃO DE TEMPO DE SERVIÇO. SENTENÇA TRABALHISTA. REVELIA. AUSÊNCIA DE COMPROVAÇÃO DO EFETIVO TEMPO DE SERVIÇO. 1. O benefício de aposentadoria por idade urbana exige o cumprimento de dois requisitos: a) idade mínima, de 65 anos, se homem, ou 60 anos, se mulher; e b) período de carência (art. 48, "caput", da Lei nº 8.213/91). 2. O Colendo Superior Tribunal de Justiça consolidou o entendimento de que a sentença trabalhista pode ser considerada como início de prova material para a determinação de tempo de serviço, desde que tenha sido fundada em outros elementos de prova que evidenciem o labor exercido e os períodos alegados pelo trabalhador. Todavia, no presente caso, a procedência da reclamação trabalhista decorreu da revelia da reclamada, não havendo notícia da produção de nenhuma prova documental ou testemunhal que ampare o reconhecimento do labor alegadamente desempenhado pela parte autora. 3. Condenação da parte autora ao pagamento das custas processuais e honorários advocatícios, fixados em 10% (dez por cento) do valor da causa, nos termos do artigo 85, § 2º, do CPC/2015, observando-se, na execução, o disposto no artigo 98, § 3º, do CPC/2015. 4. Remessa necessária e apelação do INSS providas para julgar improcedente o pedido."""), 
                        (f"(ApReeNec 00239448920174039999, DESEMBARGADOR FEDERAL NELSON PORFIRIO, TRF3 - DÉCIMA TURMA, e-DJF3 Judicial 1 DATA:23/05/2018 ..FONTE_REPUBLICACAO:.)"), 
                        (f" ")
                    ]
                for i, paragrafo in enumerate(sentenca_homologatoria):
                    parag = doc.add_paragraph(paragrafo)
                    parag.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY 
                    if i in [0, 1, 2, 3, 4, 20]:
                        parag.paragraph_format.first_line_indent = Cm(2) 
                    elif i in [5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19,]:
                        parag.paragraph_format.left_indent = Cm(2)            
                if ha_mais_documentos == 1:
                    lista = [
                        (f"Observo que a parte autora, além da sentença, trouxe outros elementos que não podem ser aceitos como início de prova material.")
                        ]
                    for linha in lista.split("\n"):
                        if linha.split():
                            lista.append(linha)
                    ft.alinhamento_parag_dispositivo(doc, lista)
                tema_629(doc)
                ft.salvar_docx_temporario(doc, processo_formatado)
        #NÃO APRESENTOU INÍCIO DE PROVA MATERIAL, OU TROUXE ELEMENTOS QUE NÃO PODEM SER CONSIDERADOS
        if hipotese == 2:
            ha_inicio_de_prova_material = st.radio("A parte autora trouxe algum documento que quer ver apreciado como início de prova material?", [1,2], format_func=lambda x: "Sim" if x == 1 else "Não")
            if ha_inicio_de_prova_material == 1:
                documentos_trazidos = st.text_area("Redija quais documentos foram apresentados e o motivo pelo qual tais documentos não podem ser considerados início de prova material (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final): ")
            else: 
                documentos_trazidos = (f"A parte autora não trouxe qualquer documento para ser apreciado como início de prova material.")
            if st.button("Gerar Sentença"):
                doc = Document()
                doc.add_paragraph(f"Processo: {processo_formatado}")
                fundamento_improcedencia = [
                        (f"Vistos."),
                        (f"Trata-se de pedido de pensão por morte."),
                        (f"Dispensado o relatório nos termos do art. 38 da Lei n. 9.099/95."),
                        (f"DECIDO."),
                        (f"O feito comporta julgamento imediato."),
                        (f"A parte autora pede o reconhecimento do tempo de serviço/contribuição do(a) segurado(a) falecido(a), Sr(a). {instituidor}, entre {data_inicial} e {data_final} referente a (origem) {vinculo_controvertido}, para, com isso, comprovar a qualidade de segurado do(a) falecido(a) instituidor(a) ao tempo do óbito."),
                        (f"A pretensão da parte autora, para acolhida, deve vir acompanhada de início de prova material, a rigor do que dispõe o artigo 55, § 3º da Lei n. 8.213/91."),
                        ]
                if ha_inicio_de_prova_material == 1:
                    fundamento_improcedencia.append(f"Os documentos apresentados pela parte autora não podem ser considerados início de prova material suficiente.")
                    for linha in documentos_trazidos.split("\n"):
                        if linha.split():
                            fundamento_improcedencia.append(linha)
                else:
                    fundamento_improcedencia.append({documentos_trazidos})
                ft.alinhamento_parag_dispositivo(doc, fundamento_improcedencia)
                tema_629(doc)
                ft.salvar_docx_temporario(doc, processo_formatado)                   
        #TESTEMUNHAS OUVIDAS EM JUÍZO NÃO COMPROVAM QUE O FALECIDO ERA SEGURADO
        if hipotese == 3:
            depoimentos = st.text_area("Redija o(s) depoimento(s) da prova testemunhal. O que for redigido será inserida como parágrafo na sentença (iniciar com letra maiúscula e colocar ponto final): ")
            conclusao = st.text_area("Redija a conclusão explicando o motivo da prova testemunhal não comprovar que o falecido era segurado. O que for redigido será inserida como parágrafo na sentença (iniciar com letra maiúscula e colocar ponto final): ") 
            if st.button("Gerar Sentença"):
                doc = Document()
                doc.add_paragraph(f"Processo: {processo_formatado}")
                texto_base(doc, fundamento_questao)
                fundamento_improcedencia = [
                        (f"A parte autora pede o reconhecimento do tempo de serviço/contribuição do(a) segurado(a) falecido(a), Sr(a). {instituidor}, entre {data_inicial} e {data_final} referente a (origem) {vinculo_controvertido}, para, com isso, comprovar a qualidade de segurado do(a) falecido(a) instituidor(a) ao tempo do óbito."),
                        (f"A prova testemunhal produzida em Juízo não comprova o alegado.")
                    ]
                for linha in depoimentos.split("\n"):
                    if linha.split():
                            fundamento_improcedencia.append(linha)
                for linha in conclusao.split("\n"):
                    if linha.split():
                            fundamento_improcedencia.append(linha)            
                fundamento_improcedencia.extend([(f"Mostra-se ausente, portanto, a qualidade de segurado do falecido ao tempo do óbito.")])
                ft.alinhamento_parag_dispositivo(doc, fundamento_improcedencia)
                improcedencia(doc)
                ft.salvar_docx_temporario(doc, processo_formatado)               
        #NÃO RECONHECIMENTO DE QUALIDADE DE SEGURADO POST MORTEM
        if hipotese == 4:
            depoimentos = st.text_area("Redija o(s) depoimento(s) da prova testemunhal. O que for redigido será inserida como parágrafo na sentença (iniciar com letra maiúscula e colocar ponto final): ")
            if st.button("Gerar Sentença"):
                doc = Document()
                doc.add_paragraph(f"Processo: {processo_formatado}")
                texto_base(doc, fundamento_questao)
                fundamento_improcedencia = [
                        (f"A parte autora pede o reconhecimento do tempo de serviço/contribuição do(a) segurado(a) falecido(a), Sr(a). {instituidor}, entre {data_inicial} e {data_final} referente a (origem) {vinculo_controvertido}."),
                        (f"A prova testemunhal produzida comprova que o trabalho do autor efetivamente se deu na condição de contribuinte individual, sem vínculo empregatício.")
                    ]
                for linha in depoimentos.split("\n"):
                    if linha.split():
                            fundamento_improcedencia.append(linha)
                fundamento_improcedencia.extend([
                        (f"O Superior Tribunal de Justiça entende pela impossibilidade de recolhimento, pelos dependentes, para fins de concessão do benefício de pensão por morte, de contribuições vertidas após o óbito do(a) instituidor(a), no caso de contribuinte individual, nas hipóteses em que a contribuição não é devida por seu contratante (Lei n. 10.666/03)."),
                        (f"Por isso, no caso concreto, não pode o trabalho do(a) instituidor(a) ser considerado para fins de manutenção da qualidade de segurado, para concessão de pensão por morte.")
                        ])
                ft.alinhamento_parag_dispositivo(doc, fundamento_improcedencia)
                improcedencia(doc)
                ft.salvar_docx_temporario(doc, processo_formatado)                
        #IMPROCEDENCIA FALTA DE DIREITO ADQUIRIDO DO INSTITUIDOR A UM BENEFÍCIO NO ÓBITO
        if hipotese == 5:
            direito_aquirido = st.text_input("Qual benefício a parte autora alega que o instituidor teria direito adquirido, quando faleceu?")
            sem_direito_adquirido = st.text_area("Explique por que o instituidor não possuía direito adquirito ao benefício informado, na data do óbito (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final): ")
            if st.button("Gerar Sentença"):
                doc = Document()
                doc.add_paragraph(f"Processo: {processo_formatado}")
                texto_base(doc, fundamento_questao)
                fundamento_improcedencia = [
                    (f"A parte autora alega que o(a) instituidor(a), Sr(a). {instituidor}, possuía direito adquirido ao recebimento de {direito_aquirido}, ao tempo do óbito."),
                ]
                for linha in sem_direito_adquirido.split("\n"):
                    if linha.split():
                            fundamento_improcedencia.append(linha)
                fundamento_improcedencia.extend(f"Mostra-se ausente, portanto, a qualidade de segurado do falecido ao tempo do óbito.")
                ft.alinhamento_parag_dispositivo(doc, fundamento_improcedencia)
                improcedencia(doc)
                ft.salvar_docx_temporario(doc, processo_formatado)             
        #OUTROS MOTIVOS PARA NÃO HAVER QUALIDADE DE SEGURADO DO FALECIDO
        if hipotese == 6:
            motivo_outro = st.text_area("Explique por que o falecido não detinha qualidade de segurado? O que for redigido será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final): ")
            if st.button("Gerar Sentença"):
                doc = Document()
                doc.add_paragraph(f"Processo: {processo_formatado}")
                texto_base(doc, fundamento_questao)
                fundamento_improcedencia = [
                        (f"A parte autora pede não prova que o(a) falecido(a) instituidor(a), Sr(a). {instituidor}, detinha qualidade de segurado ao tempo do óbito.")
                    ]
                for linha in motivo_outro.split("\n"):
                    if linha.split():
                            fundamento_improcedencia.append(linha)
                ft.alinhamento_parag_dispositivo(doc, fundamento_improcedencia)
                improcedencia(doc)
                ft.salvar_docx_temporario(doc, processo_formatado)            
    # A PARTE AUTORA NÃO É DEPENDENTE
    if motivo_improcedencia == 2:
        hipotese = st.radio("Qual hipótese melhor se enquadra no caso?", [1, 2, 3, 4, 5, 6], format_func=lambda x: 
                                "A parte autora alega ser companheiro(a) do(a) falecido(a), mas não apresenta início de prova material da união estável, ou os documentos que trouxe não podem ser considerados início de prova material." if x == 1 else
                                "A parte autora alega ser companheiro(a) do(a) falecido(a), mas a prova testemunhal não comprova união estável." if x == 2 else
                                "A parte autora alega ser companheiro(a) do(a) falecido(a), mas era concubino(a), pois o falecido(a) era casado(a)." if x ==3 else
                                "A parte autora é pai/mãe/enteado/tutelado/menor sob guarda/ex-cônjuge/irmão(ã) do falecido, mas não apresenta início de prova material da dependência econômica, ou os documentos que trouxe não pode ser considerados início de prova material." if x == 4 else
                                "A parte autora é pai/mãe/enteado/tutelado/menor sob guarda/ex-cônjuge/irmão(ã) do falecido, mas a prova testemunhal não comprova dependência econômica." if x == 5 else
                                "Outro motivo sobre a parte autora não ser dependente do falecido (a ser redigido)"
                                )
        #NÃO HÁ INÍCIO DE PROVA MATERIAL DA UNIÃO ESTÁVEL
        if hipotese == 1:
            ha_inicio_de_prova_material = st.radio("A parte autora trouxe algum documento que quer ver apreciado como início de prova material?", [1,2], format_func=lambda x: "Sim" if x == 1 else "Não")
            if ha_inicio_de_prova_material == 1:
                documentos_trazidos = st.text_area("Redija quais documentos foram aprentados e o motivo pelo qual tais documentos não podem ser considerados início de prova material (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final): ")
            else: 
                documentos_trazidos = (f"A parte autora não trouxe qualquer documento para ser apreciado como início de prova material.")
            if st.button("Gerar Sentença"):
                doc = Document()
                doc.add_paragraph(f"Processo: {processo_formatado}")
                fundamento_improcedencia = [
                        (f"Vistos."),
                        (f"Trata-se de pedido de pensão por morte."),
                        (f"Dispensado o relatório nos termos do art. 38 da Lei n. 9.099/95."),
                        (f"DECIDO."),
                        (f"O feito comporta julgamento imediato."),
                        (f"A parte autora alega que viveu em união estável com (a) segurado(a) falecido(a), Sr(a). {instituidor}."),
                        (f"""A Lei n. 8.213/91, em seu artigo 16, § 5º dispõe que: "As provas de união estável e de dependência econômica exigem início de prova material contemporânea dos fatos, produzido em período não superior a 24 (vinte e quatro) meses anterior à data do óbito ou do recolhimento à prisão do segurado, não admitida a prova exclusivamente testemunhal, exceto na ocorrência de motivo de força maior ou caso fortuito, conforme disposto no regulamento." """),
                        (f"A pretensão da parte autora, portanto, para acolhida, deve vir acompanhada de início de prova material."),
                        ]
                if ha_inicio_de_prova_material == 1:
                    fundamento_improcedencia.append(f"Os documentos apresentados pela parte autora não podem ser considerados início de prova material suficiente.")
                    for linha in documentos_trazidos.split("\n"):
                        if linha.split():
                            fundamento_improcedencia.append(linha)
                else:
                    fundamento_improcedencia.append({documentos_trazidos})
                ft.alinhamento_parag_dispositivo(doc, fundamento_improcedencia)
                tema_629(doc)
                ft.salvar_docx_temporario(doc, processo_formatado)  
        #A PROVA TESTEMUNHAL NÃO COMPROVA UNIÃO ESTÁVEL
        if hipotese == 2:
            testemunhas = st.text_area("Redija o(s) depoimento(s) da prova testemunhal. O que for redigido será inserida como parágrafo na sentença (iniciar com letra maiúscula e colocar ponto final): ")
            if st.button("Gerar Sentença"):
                doc = Document()
                doc.add_paragraph(f"Processo: {processo_formatado}")
                texto_base(doc, fundamento_questao)
                fundamento_improcedencia = [
                        (f"A parte autora alega que viveu em união estável com o(a) segurado(a) falecido(a), Sr(a). {instituidor}."),
                        (f"A prova testemunhal produzida em Juízo não comprova o alegado.")
                    ]
                for linha in testemunhas.split("\n"):
                    if linha.split():
                            fundamento_improcedencia.append(linha)
                fundamento_improcedencia.extend([(f"Por não restar comprovada a união estável, a parte autora não pode ser considerada dependente do falecido.")])
                ft.alinhamento_parag_dispositivo(doc, fundamento_improcedencia)
                improcedencia(doc)
                ft.salvar_docx_temporario(doc, processo_formatado)               
        # CONCUBINATO
        if hipotese == 3:
            testemunhas = st.text_area("Redija o(s) depoimento(s) da prova testemunhal. O que for redigido será inserida como parágrafo na sentença (iniciar com letra maiúscula e colocar ponto final): ")
            if st.button("Gerar Sentença"):
                doc = Document()
                doc.add_paragraph(f"Processo: {processo_formatado}")
                texto_base(doc, fundamento_questao)
                fundamento_improcedencia = [
                        (f"A parte autora alega que viveu em união estável com o(a) segurado(a) falecido(a), Sr(a). {instituidor}."),
                        (f"A prova testemunhal produzida em Juízo aponta para a existência de concubinato.")
                    ]
                for linha in testemunhas.split("\n"):
                    if linha.split():
                            fundamento_improcedencia.append(linha)
                fundamento_improcedencia.extend([(f"Por não restar comprovada a união estável, mas sim concubinato, a parte autora não pode ser considerada dependente do falecido.")])
                ft.alinhamento_parag_dispositivo(doc, fundamento_improcedencia)
                improcedencia(doc)
                ft.salvar_docx_temporario(doc, processo_formatado)                
        #NÃO HÁ INÍCIO DE PROVA MATERIAL DA DEPENDÊNCIA ECONÔMICA
        if hipotese in [4, 5]:
            dependente_opcoes = st.radio(
                    "O que a parte autora é, em relação ao instituidor?", 
                    [1, 2, 3, 4, 5, 6, 7], 
                    format_func=lambda x:                      
                    "ex-cônjuge" if x == 1 else
                    "enteado(a)" if x == 2 else
                    "tutelado(a)" if x == 3 else 
                    "menor sob guarda" if x == 4 else
                    "pai" if x == 5 else
                    "mãe" if x == 6 else
                    "irmão(ã)"
                    )
            if dependente_opcoes == 1:
                dependente_economico = "ex-cônjuge"
            elif dependente_opcoes == 2:
                dependente_economico = "enteado(a)"
            elif dependente_opcoes == 3:
                dependente_economico = "tutelado(a)"
            elif dependente_opcoes == 4:
                dependente_economico = "menor sob guarda"
            elif dependente_opcoes == 5:
                dependente_economico = "pai"
            elif dependente_opcoes == 6:
                dependente_economico = "mãe"
            elif dependente_opcoes == 7:
                dependente_economico = "irmão(ã)"
        if hipotese == 4:
            ha_inicio_de_prova_material = st.radio("A parte autora trouxe algum documento que quer ver apreciado como início de prova material?", [1,2], format_func=lambda x: "Sim" if x == 1 else "Não")
            if ha_inicio_de_prova_material == 1:
                documentos_trazidos = st.text_area("Redija quais documentos foram apresentados e o motivo pelo qual tais documentos apresentados não podem ser considerados início de prova material (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final): ")
            else: 
                documentos_trazidos = (f"A parte autora não trouxe qualquer documento para ser apreciado como início de prova material.")
            if st.button("Gerar Sentença"):
                doc = Document()
                doc.add_paragraph(f"Processo: {processo_formatado}")
                fundamento_improcedencia = [
                        (f"Vistos."),
                        (f"Trata-se de pedido de pensão por morte."),
                        (f"Dispensado o relatório nos termos do art. 38 da Lei n. 9.099/95."),
                        (f"DECIDO."),
                        (f"O feito comporta julgamento imediato."),
                        (f"A parte autora é {dependente_economico} do(a) segurado(a) falecido(a), Sr(a). {instituidor}, e, portanto, deve comprovar efetiva dependência econômica em relação a ele."),
                        (f"""A Lei n. 8.213/91, em seu artigo 16, § 5º dispõe que: "As provas de união estável e de dependência econômica exigem início de prova material contemporânea dos fatos, produzido em período não superior a 24 (vinte e quatro) meses anterior à data do óbito ou do recolhimento à prisão do segurado, não admitida a prova exclusivamente testemunhal, exceto na ocorrência de motivo de força maior ou caso fortuito, conforme disposto no regulamento." """),
                        (f"A pretensão da parte autora, portanto, para acolhida, deve vir acompanhada de início de prova material."),
                        ]
                if ha_inicio_de_prova_material == 1:
                    fundamento_improcedencia.append(f"Os documentos apresentados pela parte autora não podem ser considerados início de prova material suficiente.")
                    for linha in documentos_trazidos.split("\n"):
                        if linha.split():
                            fundamento_improcedencia.append(linha)
                else:
                    fundamento_improcedencia.append({documentos_trazidos})
                ft.alinhamento_parag_dispositivo(doc, fundamento_improcedencia)
                tema_629(doc)
                ft.salvar_docx_temporario(doc, processo_formatado)   
        #A PROVA TESTEMUNHAL NÃO COMPROVA DEPENDÊNCIA ECONÔMICA
        if hipotese == 5:
            dependente_economico()
            testemunhas = st.text_area("Redija o(s) depoimento(s) da prova testemunhal. O que for redigido será inserida como parágrafo na sentença (iniciar com letra maiúscula e colocar ponto final): ")
            if st.button("Gerar Sentença"):
                doc = Document()
                doc.add_paragraph(f"Processo: {processo_formatado}")
                texto_base(doc, fundamento_questao)
                fundamento_improcedencia = [
                        (f"A parte autora é {dependente_economico} do(a) segurado(a) falecido(a), Sr(a). {instituidor}, e, portanto, deve comprovar efetiva dependência econômica em relação a ele."),
                        (f"A prova testemunhal produzida em Juízo não comprova o alegado.")
                    ]
                for linha in testemunhas.split("\n"):
                    if linha.split():
                        fundamento_improcedencia.append(linha)
                fundamento_improcedencia.extend([(f"Por não restar comprovada a dependência econômica, a parte autora não pode ser considerada dependente do falecido.")])
                ft.alinhamento_parag_dispositivo(doc, fundamento_improcedencia)
                improcedencia(doc)
                ft.salvar_docx_temporario(doc, processo_formatado)                  
        #OUTROS MOTIVOS PELOS QUAIS A PARTE AUTORA NÃO É DEPENDENTE
        if hipotese == 6:
            motivo_outro = st.text_area("Explique por que a parte autora não é dependente do falecido. O que for redigido será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final): ")
            if st.button("Gerar Sentença"):
                doc = Document()
                doc.add_paragraph(f"Processo: {processo_formatado}")
                texto_base(doc, fundamento_questao)
                fundamento_improcedencia = [
                        (f"A parte autora não prova que era dependente, nos termos da Lei n. 8.213/91, do(a) falecido(a) instituidor(a), Sr(a). {instituidor}.")
                    ]
                for linha in motivo_outro.split("\n"):
                    if linha.split():
                            fundamento_improcedencia.append(linha)
                ft.alinhamento_parag_dispositivo(doc, fundamento_improcedencia)
                improcedencia(doc)
                ft.salvar_docx_temporario(doc, processo_formatado)               
#PROCEDENTE
if resultado == 1:
    caso_exige_prova_material_qs_instituidor = st.radio("A qualidade de segurado do instituidor depende da comprovação em Juízo de tempo de serviço, que exija início de prova material nos termos do art. 55, § 3º da Lei n. 8.213/91?", [1, 2], format_func=lambda x: "Sim" if x == 1 else "Não", index=1)
    if caso_exige_prova_material_qs_instituidor == 1:
        prova_material_qs_instituidor = st.text_area("O que a parte autora apresentou como início de prova material? O que for redigido será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final):", key="prova_material_qs_instituidor")
        depoimentos_qs_instituidor = st.text_area("Redija o(s) depoimento(s) da prova testemunhal. O que for redigido será inserida como parágrafo na sentença (iniciar com letra maiúscula e colocar ponto final): ")
        qualidade_segurado = st.text_area("Os depoimentos comprovam a qualidade de segurado do instituidor, mas é preciso especificar o que levou a esta conclusão. O que for redigido será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final): ")
    else:
        qualidade_segurado = st.text_area("Por que o segurado instituidor mantinha a qualidade de segurado no óbito? O que for redigido será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final): ")
    dependente_opcoes = st.radio(
                    "O que a parte autora é, em relação ao instituidor?", 
                    [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15], 
                    format_func=lambda x:
                    "cônjuge" if x == 1 else
                    "companheiro(a)" if x == 2 else
                    "filho(a) menor de 21 anos de idade" if x == 3 else
                    "filho(a) inválido, deficiente intelectual ou mental, ou deficiente grave" if x == 4 else               
                    "ex-marido" if x == 5 else
                    "ex-esposa" if x == 6 else
                    "enteado(a) menor de 21 anos de idade" if x == 7 else
                    "tutelado(a) menor de 21 anos de idade" if x == 8 else 
                    "enteado(a) inválido, deficiente intelectual ou mental, ou deficiente grave" if x == 9 else
                    "tutelado(a) inválido, deficiente intelectual ou mental, ou deficiente grave" if x == 10 else 
                    "menor sob guarda" if x == 11 else
                    "pai" if x == 12 else
                    "mãe" if x == 13 else
                    "irmão(ã) menor de 21 anos de idade" if x == 14 else
                    "irmão(ã) inválido, deficiente intelectual ou mental, ou deficiente grave",
                    index=1 
                    )
    if dependente_opcoes == 1:
        dependente = "cônjuge"
    elif dependente_opcoes == 2:
        dependente = "companheiro(a)"
    elif dependente_opcoes == 3:
        dependente = "filho(a) menor de 21 anos de idade"
    elif dependente_opcoes == 4:
        dependente = "filho(a) inválido, deficiente intelectual ou mental, ou deficiente grave"
    elif dependente_opcoes == 5:
        dependente = "ex-marido"
    elif dependente_opcoes == 6:
        dependente = "ex-esposa"
    elif dependente_opcoes == 7:
        dependente = "enteado(a) menor de 21 anos de idade"
    elif dependente_opcoes == 8:
        dependente = "tutelado(a) menor de 21 anos de idade"
    elif dependente_opcoes == 9:
        dependente = "enteado(a) inválido, deficiente intelectual ou mental, ou deficiente grave"
    elif dependente_opcoes == 10:
        dependente = "tutelado(a) inválido, deficiente intelectual ou mental, ou deficiente grave"
    elif dependente_opcoes == 11:
        dependente = "menor sob guarda"
    elif dependente_opcoes == 12:
        dependente = "pai"
    elif dependente_opcoes == 13:
        dependente = "mãe"
    elif dependente_opcoes == 14:
        dependente = "irmão(ã) menor de 21 anos de idade"
    elif dependente_opcoes == 15:
        dependente = "irmão(ã) inválido, deficiente intelectual ou mental, ou deficiente grave"

    if dependente_opcoes == 1:
        data_casamento = st.text_input("Qual a data do casamento? Escreva no formato DD/MM/AAAA")
    if dependente_opcoes in [3, 7, 8, 11, 14]:
        data_nascimento = st.text_input("Qual a data de nascimento do(a) dependente? Escreva no formato DD/MM/AAAA")
        data_aniversario = st.text_input("Qual a data em que o(a) dependente completará 21 anos de idade? Escreva no formato DD/MM/AAAA")
        pensao_prazo = (f"até completar 21 anos de idade em {data_aniversario}")
    if dependente_opcoes in [4, 9, 10, 15]:
        pericia = st.text_area("O que a perícia médica constatou sobre invalidez, deficiência intelectual ou mental ou deficiência grave? O que for redigido será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final): ")
        pensao_prazo = "enquanto permanecer inválido, deficiente intelectual ou mental, ou deficiente grave"    
    if dependente_opcoes in [5, 6]:
        ha_sentenca_alimentos = st.radio("A parte trouxe sentença que comprove a fixação de pensão alimentícia, ou fará prova de dependência econômica por testemunha?", [1, 2], format_func=lambda x: "Trouxe sentença" if x == 1 else "Ouvirá testemunhas")    
        if ha_sentenca_alimentos == 1:
            alimentos = st.radio("No caso, a parte trouxe sentença de comprove fixação de pensão alimentícia temporária, ou não há prazo final fixado para a pensão alimentícia?", [1, 2], format_func=lambda x: "Temporária" if x == 1 else "Sem prazo")
            if alimentos == 1:
                prazo_alimentos = st.text_input("Qual o prazo de duração da pensão alimentícia temporária? (Escreva no formato dd/mm/aaaa):")
                tempo_casamento_uniao, carencia_instituidor, tempo = prazo_pensao_conjuge_companheira(data_do_obito_convertida)
                if not tempo == "de forma vitalícia":        
                    dependente_incapaz = st.radio("O dependente é inválido ou deficiente?", [1, 2], format_func=lambda x: "Sim" if x == 1 else "Não", index=1, key="dependente_incapaz_pensao")
                    if dependente_incapaz == 1:
                        pericia = st.text_area("O que a perícia médica constatou? O que for redigido será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final): ")
                        pensao_prazo = (f"{tempo}"+", no mínimo, e enquanto durar a invalidez/deficiência da parte autora")
                    if dependente_incapaz == 2:
                        pensao_prazo = tempo
                if tempo == "de forma vitalícia":
                    pensao_prazo = "de forma vitalícia"    
            if alimentos == 2:
                tempo_casamento_uniao, carencia_instituidor, tempo = prazo_pensao_conjuge_companheira(data_do_obito_convertida)
        if ha_sentenca_alimentos == 2:
            lei13846 = date(2019, 6, 18)
            if data_do_obito_convertida < lei13846:
                inicio_prova_material_dependente = "O óbito ocorreu antes da vigência da Lei n. 13.846/2019, não havendo previsão legal para necessidade de apresentação de início de prova material"
            if lei13846 <= data_do_obito_convertida:
                inicio_prova_material_dependente = st.text_area("Qual(is) elemento(s) a parte autora trouxe como início de prova material? O que for redigido será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final): ")
            depoimentos = st.text_area("Redija o(s) depoimento(s) da prova testemunhal. O que for redigido será inserida como parágrafo na sentença (iniciar com letra maiúscula e colocar ponto final): ")
            conclusao = st.text_area("Os depoimentos comprovam a dependência econômica, mas é preciso especificar o que levou a esta conclusão. O que for redigido será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final): ")
            tempo_casamento_uniao, carencia_instituidor, tempo = prazo_pensao_conjuge_companheira(data_do_obito_convertida)
            if not tempo == "de forma vitalícia":        
                dependente_incapaz = st.radio("O dependente é inválido ou deficiente?", [1, 2], format_func=lambda x: "Sim" if x == 1 else "Não", index=1, key="dependente_incapaz_pensao")
                if dependente_incapaz == 1:
                    pericia = st.text_area("O que a perícia médica constatou? O que for redigido será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final): ")
                    pensao_prazo = (f"{tempo}"+", no mínimo, e enquanto durar a invalidez/deficiência da parte autora")
                if dependente_incapaz == 2:
                    pensao_prazo = tempo
                if tempo == "de forma vitalícia":
                    pensao_prazo = "de forma vitalícia"                
    if dependente_opcoes in [2, 7, 8, 9, 10, 11, 12, 13, 14, 15]:
        lei13846 = date(2019, 6, 18)
        if data_do_obito_convertida < lei13846:
            inicio_prova_material_dependente = "O óbito ocorreu antes da vigência da Lei n. 13.846/2019, não havendo previsão legal para necessidade de apresentação de início de prova material"
        if lei13846 <= data_do_obito_convertida:
            inicio_prova_material_dependente = st.text_area("Qual(is) elemento(s) a parte autora trouxe como início de prova material? O que for redigido será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final): ")
        depoimentos = st.text_area("Redija o(s) depoimento(s) da prova testemunhal. O que for redigido será inserida como parágrafo na sentença (iniciar com letra maiúscula e colocar ponto final): ")
        if dependente_opcoes == 2:
            conclusao = st.text_area("Os depoimentos comprovam a união estável, mas é preciso esclarecer desde quando, e, especificar o que levou a esta conclusão. O que for redigido será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final): ")
        if dependente_opcoes in [7, 8, 9, 10, 11, 12, 13, 14, 15]:
            conclusao = st.text_area("Os depoimentos comprovam a dependência econômica, mas é preciso especificar o que levou a esta conclusão. O que for redigido será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final): ")
    if dependente_opcoes in [1, 2]:
        tempo_casamento_uniao, carencia_instituidor, tempo = prazo_pensao_conjuge_companheira(data_do_obito_convertida)
        if not tempo == "de forma vitalícia":        
            dependente_incapaz = st.radio("O dependente é inválido ou deficiente?", [1, 2], format_func=lambda x: "Sim" if x == 1 else "Não", index=1, key="dependente_incapaz_pensao")
            if dependente_incapaz == 1:
                pericia = st.text_area("O que a perícia médica constatou? O que for redigido será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final): ")
                pensao_prazo = (f"{tempo}"+", no mínimo, e enquanto durar a invalidez/deficiência da parte autora")
            if dependente_incapaz == 2:
                pensao_prazo = tempo
        if tempo == "de forma vitalícia":
            pensao_prazo = "de forma vitalícia"                
    dib = st.text_input("DIB (dd/mm/aaaa):")
    motivo_DIB = st.radio("DIB fixada na DER?", [1,2,3],
                    format_func = lambda x: "DIB na DER" if x == 1 else "DIB no óbito" if x == 2 else "DIB fixada em outra data (necessário esclarecer)")
    if motivo_DIB == 3:
        motivo_DIB_redigido = st.text_area("Esclareça a DIB escolhida (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final):")
    elif motivo_DIB == 2:
        motivo_DIB_redigido = "DIB fixada no óbito (data da entrada do requerimento no prazo legal, considerando a data do óbito)."
    else:
        motivo_DIB_redigido = "DIB fixada na DER (data da entrada do requerimento acima do prazo legal, considerando a data do óbito)."
    procedencia_total_parcial = st.radio("Para fins da redação do dispositivo da sentença, a procedência foi total ou parcial?", [1, 2], format_func=lambda x: "Total" if x == 1 else "Parcial")
    if procedencia_total_parcial == 1:
        resultado_dispositivo = ""
    else:
        resultado_dispositivo = "EM PARTE "

    if st.button("Gerar Sentença"):
        data_atual = datetime.now()
        dip = data_atual.strftime("01/%m/%Y")        
        doc = Document()
        doc.add_paragraph(f"Processo: {processo_formatado}")
        texto_base(doc, fundamento_questao)
        fundamento_procedencia = [
            (f"A parte autora comprova a qualidade de segurado do(a) instituidor(a), Sr(a). {instituidor}, ao tempo do óbito em {data_do_obito}, conforme o disposto no artigo 15, inciso I, da Lei n. 8.213/91."),
            ]
        if caso_exige_prova_material_qs_instituidor == 2:
            for linha in qualidade_segurado.split("\n"):
                if linha.split():
                    fundamento_procedencia.append(linha)
        if caso_exige_prova_material_qs_instituidor == 1:
            fundamento_procedencia.append([(f"Foi apresentado início de prova material, que comprova a qualidade de segurado do(a) instituidor(a), conforme o disposto no artigo 55, § 3º, da Lei n. 8.213/91.")])
            for linha in prova_material_qs_instituidor.split("\n"):
                if linha.split():
                    fundamento_procedencia.append(linha)
            fundamento_procedencia.append([(f"Foram ouvidas as seguintes testemunhas:")])
            for linha in depoimentos_qs_instituidor.split("\n"):
                if linha.split():
                    fundamento_procedencia.append(linha)
            fundamento_procedencia.append([(f"{qualidade_segurado}")])           
        fundamento_procedencia.append([(f"Quanto à condição de dependente, a parte autora comprova ser {dependente} do(a) instituidor(a).")])
        if dependente_opcoes == 1:
            fundamento_procedencia.append(f"Foi apresentado certidão casamento, realizado em {data_casamento}.")
        if dependente_opcoes == 3:
            fundamento_procedencia.append(f"Foi apresentado certidão de nascimento / documento de identificação civil, que comprova a filiação, com nascimento em {data_nascimento}.")
        if dependente_opcoes == 4:
            fundamento_procedencia.append(f"Foi realizada perícia médica que comprova a condição de inválido, deficiente intelectual ou mental, ou deficiente grave, da parte autora.")
            for linha in pericia.split("\n"):
                if linha.split():
                    fundamento_procedencia.append(linha)
        if dependente_opcoes in [2, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15]:
            if data_do_obito_convertida < lei13846:
                fundamento_procedencia.append(f"inicio_prova_material_dependente")
            if lei13846 <= data_do_obito_convertida:
                fundamento_procedencia.append(f"Foi apresentado início de prova material, que comprova a condição de dependente da parte autora, conforme o disposto no artigo 16, § 5º, da Lei n. 8.213/91.")  
                for linha in inicio_prova_material_dependente.split("\n"):
                    if linha.split():
                        fundamento_procedencia.append(linha)
            fundamento_procedencia.append(f"Foram ouvidas as seguintes testemunhas:")
            for linha in depoimentos.split("\n"):
                if linha.split():
                    fundamento_procedencia.append(linha)
            fundamento_procedencia.append(f"{conclusao}")
            if dependente_incapaz == 1:
                fundamento_procedencia.append(f"Foi realizada perícia médica que comprova a condição de inválido, deficiente intelectual ou mental, ou deficiente grave, conforme o disposto no artigo 16, § 2º, da Lei n. 8.213/91.")
                for linha in pericia.split("\n"):
                    if linha.split():
                        fundamento_procedencia.append(linha)
        if dependente_opcoes == 1 and data_do_obito_convertida < date(2015, 1, 3):
            fundamento_procedencia.append(f"Considerando que o obito ocorreu antes da vigência da Lei n. 13.135/2015, fruto da conversão da MP 664/2014, a pensão por morte é vitalícia")  
        if dependente_opcoes == 1 and data_do_obito_convertido => date(2015, 1, 3):
            fundamento_procedencia.append(f"Considerando que o tempo de casamento da parte autora (data de casamento: {data_casamento}) do(a) instituidor(a) , Sr(a). {instituidor}, que em vida possuía {carencia_instituidor} vertidas, faz jus ao benefício pleiteado de pensão por morte, {pensao_prazo}.")
        if dependente_opcoes == 2 and data_do_obito_convertida < date(2015, 1, 3):
            fundamento_procedencia.append(f"Considerando que o obito ocorreu antes da vigência da Lei n. 13.135/2015, fruto da conversão da MP 664/2014, a pensão por morte é vitalícia")  
        if dependente_opcoes == 2:
            fundamento_procedencia.append(f"Considerando que a parte autora vive em união estável por {tempo_casamento_uniao} com o(a) instituidor(a), Sr(a). {instituidor}, que em vida possuía {carencia_instituidor} vertidas, faz jus ao benefício pleiteado de pensão por morte, {pensao_prazo}.")
        if dependente_opcoes in [3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15]:
            fundamento_procedencia.append(f"Assim, a parte autora faz jus ao benefício pleiteado de pensão por morte {pensao_prazo}.")
        fundamento_procedencia.extend([
            (f"A DIB será fixada em {dib}. {motivo_DIB_redigido}"),
            (f"Isto posto, com resolução de mérito nos termos do artigo 487, inciso I, do Código de Processo Civil, JULGO PROCEDENTE {resultado_dispositivo}o pedido para condenar o INSS a conceder o benefício de pensão por morte à parte autora, tendo como instituidor(a) o(a) Sr(a). {instituidor}, com DIB em {dib}, {pensao_prazo}. RMI e RMA a serem calculadas pelo INSS."),
            (f"Condeno o INSS ao pagamento dos atrasados devidos desde a DIB fixada, até a DIP em {dip}, atualizados desde cada competência devida e com juros desde a propositura da demanda, pelos índices e percentuais do Manual de Cálculos da Justiça Federal, a ser apurado em cumprimento invertido de sentença."),
            (f"Fica autorizado o desconto de eventuais valores recebidos a título de benefícios inacumuláveis."),
            (f"Condeno o INSS ao ressarcimento de eventuais honorários periciais antecipados pela Justiça Federal nesta lide (art. 82, § 2º, do CPC)."),
            (f"Considerando que o momento da prolação de sentença é oportuno para distribuir o ônus do tempo do processo, com vistas a salvaguardar a eficácia do princípio constitucional da razoável duração do processo e ao mesmo tempo privilegiar o direito provável em detrimento do improvável, demonstrada a verossimilhança das alegações da parte autora e diante do nítido caráter alimentar da verba pleiteada, nos termos do art. 294 e 300, do CPC ANTECIPA A TUTELA JURISDICIONAL para determinar ao INSS que providencie a implantação da pensão por morte na forma concedida, com data de início de pagamento em {dip} (DIP)."),
            (f"O INSS deverá providenciar a implantação do benefício previdenciário ora concedido no prazo legal, sendo a contagem em dias úteis, sendo que constitui ônus das partes informar ao Juízo sobre a efetiva implantação do benefício ou eventual descumprimento do prazo pelo INSS/APSADJ."),
            (f"Sem condenação em honorários nesta instância."),
            (f"Defiro os benefícios da gratuidade."),
            (f"Em caso de interposição de embargos de declaração, intime-se a parte contrária para contrarrazões no prazo legal, e tornem conclusos para julgamento."),
            (f"Interposto recurso, intime-se a parte contrária para contrarrazões no prazo legal. Após, remetam-se os autos às Turmas Recursais."),
            (f"Com o trânsito em julgado, implantado o benefício, dê-se início ao cumprimento de sentença."),
            (f"Proceda a Secretaria como necessário."),
            (f"Int.")
            ]) 
        ft.alinhamento_parag_dispositivo(doc, fundamento_procedencia)
        ft.salvar_docx_temporario(doc, processo_formatado)              
