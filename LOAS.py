from datetime import datetime
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import tempfile
import platform
import streamlit as st
import funcoes_texto as ft

def texto_base(doc, fundamento_questao):
    if fundamento_questao == 1:
        for i, paragrafo in enumerate(fundamento_base):
            parag = doc.add_paragraph(paragrafo)
            parag.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY 
            if i in [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 14, 18, 21, 24, 27, 31]:
                parag.paragraph_format.first_line_indent = Cm(2) 
            elif i in [11, 12, 13, 15, 16, 17, 19, 20, 22, 23, 25, 26, 28, 29, 30]:
                parag.paragraph_format.left_indent = Cm(2)
    if fundamento_questao == 2:
        for linha in fundamento_custom.split("\n"):
            if linha.split():
                parag = doc.add_paragraph(linha.strip())
                parag.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                parag.paragraph_format.first_line_indent = Cm(2)

fundamento_questao = st.radio(
"**Relatório e fundamentação jurídica**\n\n"
"Toda sentença gerada possui dispensa de relatório e uma fundamentação jurídica básica, que vai até onde se inicia a análise do caso concreto.\n\n"  
"Você pode usar esta fundamentação, ou fornecer sua **própria fundamentação**, com ou sem relatório (fazendo menção à dispensa, se for o caso.)\n\n"  
"**Como você deseja prosseguir?**",
[1, 2],
format_func=lambda x: "Vou usar a fundamentação padrão deste aplicativo." if x == 1 else "Desejo fornecer a minha fundamentação."
)
if fundamento_questao == 1:
    fundamento_base = [
            (f"Vistos."),
            (f"Trata-se de pedido de benefício de prestação continuada."),
            (f"Dispensado o relatório nos termos do art. 38 da Lei n. 9.099/95."),
            (f"DECIDO."),
            (f"O feito comporta julgamento imediato."),
            (f"O valor da causa está abaixo de 60 salários-mínimos, motivo pelo qual o Juízo é competente, e não há que se falar em renúncia. O tema 1030 do STJ prevê que: “Ao autor que deseje litigar no âmbito de Juizado Especial Federal Cível, é lícito renunciar, de modo expresso e para fins de atribuição de valor à causa, ao montante que exceda os 60 (sessenta) salários mínimos previstos no art. 3º, caput, da Lei 10.259/2001, aí incluídas, sendo o caso, até doze prestações vincendas, nos termos do art. 3º, § 2º, da referida lei, c/c o art. 292, §§ 1º e 2º, do CPC/2015.” Logo, a renúncia só tem sentido nos casos em que o valor da causa, apurado sem qualquer renúncia, pudesse ser superior à alçada, pois nesta hipótese a renúncia daria ensejo à parte litigar sob a competência do Juizado. Não é o caso dos autos, onde o valor da causa apontado já está abaixo da alçada, e a parte ré não aponta erro na sua apuração."),
            (f"Partes legítimas e bem representadas."),
            (f"Passo ao mérito."),
            (f"Quanto a prescrição, nos termos da súmula 85 do STJ estão prescritas as parcelas eventualmente devidas além dos cinco anos que antecedem a propositura da demanda."),
            (f"Passo ao mérito propriamente dito."),
            (f"Dispõe o art. 203, V, da Constituição Federal:"),
            (f"Art. 203. A assistência social será prestada a quem dela necessitar, independentemente de contribuição à seguridade social, e tem por objetivos:"),
            (f"(...)"),
            (f"V - a garantia de um salário mínimo de benefício mensal à pessoa portadora de deficiência e ao idoso que comprovem não possuir meios de prover à própria manutenção ou de tê-la provida por sua família, conforme dispuser a lei."),
            (f"Por sua vez, a Lei n. 8.742/93 assim disciplina a matéria:"),
            (f"Art. 20.  O benefício de prestação continuada é a garantia de um salário-mínimo mensal à pessoa com deficiência e ao idoso com 65 (sessenta e cinco) anos ou mais que comprovem não possuir meios de prover a própria manutenção nem de tê-la provida por sua família."),
            (f"(...)"),
            (f"§ 3º Observados os demais critérios de elegibilidade definidos nesta Lei, terão direito ao benefício financeiro de que trata o caput deste artigo a pessoa com deficiência ou a pessoa idosa com renda familiar mensal per capita igual ou inferior a 1/4 (um quarto) do salário-mínimo. (Redação dada pela Lei nº 14.176, de 2021)"),
            (f"Ocorre que o Supremo Tribunal Federal, no julgamento do Recurso Extraordinário n. 567.985 assim decidiu:"),
            (f"Benefício assistencial de prestação continuada ao idoso e ao deficiente. Art. 203, V, da Constituição. A Lei de Organização da Assistência Social (LOAS), ao regulamentar o art. 203, V, da Constituição da República, estabeleceu os critérios para que o benefício mensal de um salário mínimo seja concedido aos portadores de deficiência e aos idosos que comprovem não possuir meios de prover a própria manutenção ou de tê-la provida por sua família. 2. Art. 20, § 3º, da Lei 8.742/1993 e a declaração de constitucionalidade da norma pelo Supremo Tribunal Federal na ADI 1.232. Dispõe o art. 20, § 3º, da Lei 8.742/93 que “considera-se incapaz de prover a manutenção da pessoa portadora de deficiência ou idosa a família cuja renda mensal per capita seja inferior a 1/4 (um quarto) do salário mínimo”. O requisito financeiro estabelecido pela lei teve sua constitucionalidade contestada, ao fundamento de que permitiria que situações de patente miserabilidade social fossem consideradas fora do alcance do benefício assistencial previsto constitucionalmente. Ao apreciar a Ação Direta de Inconstitucionalidade 1.232-1/DF, o Supremo Tribunal Federal declarou a constitucionalidade do art. 20, § 3º, da LOAS. 3. Decisões judiciais contrárias aos critérios objetivos preestabelecidos e Processo de inconstitucionalização dos critérios definidos pela Lei 8.742/1993. A decisão do Supremo Tribunal Federal, entretanto, não pôs termo à controvérsia quanto à aplicação em concreto do critério da renda familiar per capita estabelecido pela LOAS. Como a lei permaneceu inalterada, elaboraram-se maneiras de se contornar o critério objetivo e único estipulado pela LOAS e de se avaliar o real estado de miserabilidade social das famílias com entes idosos ou deficientes. Paralelamente, foram editadas leis que estabeleceram critérios mais elásticos para a concessão de outros benefícios assistenciais, tais como: a Lei 10.836/2004, que criou o Bolsa Família; a Lei 10.689/2003, que instituiu o Programa Nacional de Acesso à Alimentação; a Lei 10.219/01, que criou o Bolsa Escola; a Lei 9.533/97, que autoriza o Poder Executivo a conceder apoio financeiro a Municípios que instituírem programas de garantia de renda mínima associados a ações socioeducativas. O Supremo Tribunal Federal, em decisões monocráticas, passou a rever anteriores posicionamentos acerca da intransponibilidade do critérios objetivos. Verificou-se a ocorrência do processo de inconstitucionalização decorrente de notórias mudanças fáticas (políticas, econômicas e sociais) e jurídicas (sucessivas modificações legislativas dos patamares econômicos utilizados como critérios de concessão de outros benefícios assistenciais por parte do Estado brasileiro). 4. Declaração de inconstitucionalidade parcial, sem pronúncia de nulidade, do art. 20, § 3º, da Lei 8.742/1993. 5. Recurso extraordinário a que se nega provimento."),
            (f"(RE - RECURSO EXTRAORDINÁRIO 567.985, REL. MIN. MARCO AURÉLIO, STF.)"),
            (f"No julgamento do Recurso Extraordinário n. 580.963 assim decidiu:"),
            (f"Benefício assistencial de prestação continuada ao idoso e ao deficiente. Art. 203, V, da Constituição. A Lei de Organização da Assistência Social (LOAS), ao regulamentar o art. 203, V, da Constituição da República, estabeleceu os critérios para que o benefício mensal de um salário mínimo seja concedido aos portadores de deficiência e aos idosos que comprovem não possuir meios de prover a própria manutenção ou de tê-la provida por sua família. 2. Art. 20, § 3º, da Lei 8.742/1993 e a declaração de constitucionalidade da norma pelo Supremo Tribunal Federal na ADI 1.232. Dispõe o art. 20, § 3º, da Lei 8.742/93 que: “considera-se incapaz de prover a manutenção da pessoa portadora de deficiência ou idosa a família cuja renda mensal per capita seja inferior a 1/4 (um quarto) do salário mínimo”. O requisito financeiro estabelecido pela Lei teve sua constitucionalidade contestada, ao fundamento de que permitiria que situações de patente miserabilidade social fossem consideradas fora do alcance do benefício assistencial previsto constitucionalmente. Ao apreciar a Ação Direta de Inconstitucionalidade 1.232-1/DF, o Supremo Tribunal Federal declarou a constitucionalidade do art. 20, § 3º, da LOAS. 3. Decisões judiciais contrárias aos critérios objetivos preestabelecidos e processo de inconstitucionalização dos critérios definidos pela Lei 8.742/1993. A decisão do Supremo Tribunal Federal, entretanto, não pôs termo à controvérsia quanto à aplicação em concreto do critério da renda familiar per capita estabelecido pela LOAS. Como a Lei permaneceu inalterada, elaboraram-se maneiras de contornar o critério objetivo e único estipulado pela LOAS e de avaliar o real estado de miserabilidade social das famílias com entes idosos ou deficientes. Paralelamente, foram editadas leis que estabeleceram critérios mais elásticos para concessão de outros benefícios assistenciais, tais como: a Lei 10.836/2004, que criou o Bolsa Família; a Lei 10.689/2003, que instituiu o Programa Nacional de Acesso à Alimentação; a Lei 10.219/01, que criou o Bolsa Escola; a Lei 9.533/97, que autoriza o Poder Executivo a conceder apoio financeiro a municípios que instituírem programas de garantia de renda mínima associados a ações socioeducativas. O Supremo Tribunal Federal, em decisões monocráticas, passou a rever anteriores posicionamentos acerca da intransponibilidade dos critérios objetivos. Verificou-se a ocorrência do processo de inconstitucionalização decorrente de notórias mudanças fáticas (políticas, econômicas e sociais) e jurídicas (sucessivas modificações legislativas dos patamares econômicos utilizados como critérios de concessão de outros benefícios assistenciais por parte do Estado brasileiro). 4. A inconstitucionalidade por omissão parcial do art. 34, parágrafo único, da Lei 10.741/2003. O Estatuto do Idoso dispõe, no art. 34, parágrafo único, que o benefício assistencial já concedido a qualquer membro da família não será computado para fins do cálculo da renda familiar per capita a que se refere a LOAS. Não exclusão dos benefícios assistenciais recebidos por deficientes e de previdenciários, no valor de até um salário mínimo, percebido por idosos. Inexistência de justificativa plausível para discriminação dos portadores de deficiência em relação aos idosos, bem como dos idosos beneficiários da assistência social em relação aos idosos titulares de benefícios previdenciários no valor de até um salário mínimo. Omissão parcial inconstitucional. 5. Declaração de inconstitucionalidade parcial, sem pronúncia de nulidade, do art. 34, parágrafo único, da Lei 10.741/2003. 6. Recurso extraordinário a que se nega provimento."),
            (f"(RE - RECURSO EXTRAORDINÁRIO 580.963, REL. MIN. GILMAR MENDES, STF.)"),
            (f"A partir destes julgamentos, a conclusão a que se chega é que o critério de miserabilidade baseado na renda per capita, dentro dos limites legais, não é absoluto. É possível a prova por outros meios, sendo importante, para tanto, a conclusão da perícia socioeconômica. Neste sentido:"),
            (f"PREVIDENCIÁRIO. BENEFÍCIO ASSISTENCIAL. ART. 20, § 3º, DA LEI 8.742/93. INCONSTITUCIONALIDADE. TERMO INICIAL. DATA DA CITAÇÃO. 1. O limite legal estabelecido no art. 20, § 3º, da Lei 8.742/93 não é critério absoluto, de modo que a necessidade/miserabilidade do postulante pode ser comprovada de outras maneiras. 2. O STF, no recente julgamento dos REs 567.985 e 580.963, assentou a inconstitucionalidade do art. 20, § 3º, da Lei 8.742/93. 3. O termo inicial do benefício assistencial é a data da citação da autarquia previdenciária. Precedentes. 4. Agravo regimental não provido."),
            (f"(AGRESP - AGRAVO REGIMENTAL NO RECURSO ESPECIAL - 1341655 2012.01.62185-5, CASTRO MEIRA, STJ - SEGUNDA TURMA, DJE DATA:16/08/2013)"),
            (f"Por fim, em que pese a decisão do Supremo Tribunal Federal, é certo que houve inovação legislativa posterior, com a inclusão do § 14º no art. 20 da Lei n. 8.742/93, que assim dispõe:"),
            (f"Art. 20..."),
            (f"(...)"),
            (f"§ 14. O benefício de prestação continuada ou o benefício previdenciário no valor de até 1 (um) salário-mínimo concedido a idoso acima de 65 (sessenta e cinco) anos de idade ou pessoa com deficiência não será computado, para fins de concessão do benefício de prestação continuada a outro idoso ou pessoa com deficiência da mesma família, no cálculo da renda a que se refere o § 3º deste artigo. (Incluído pela Lei nº 13.982, de 2020)"),
            (f"Assim é que, ao Juízo, é lícito analisar o cumprimento dos requisitos para concessão do benefício, atentando-se para o que dispõe a lei, mas sem se descurar do quanto apontado no laudo socioeconômico."),
            ]
else:
    fundamento_custom = st.text_area(
        "Redija, ou copie e cole, a fundamentação que deseja inserir na sentença. \nO texto deve englobar tudo, desde o 'vistos em sentença' até um parágrafo assim redigido: 'Feitas estas considerações, passo a analisar o caso concreto', ou expressão equivalente."
        )

# Input do resultado (procedente ou improcedente)
resultado = st.radio("O pedido é procedente (no todo ou em parte) ou improcedente?", [1, 2], format_func=lambda x: "Procedente" if x == 1 else "Improcedente")

if resultado == 2:
    motivo_improcedencia = st.radio("Qual o motivo da improcedência?", [1, 2, 3], format_func=lambda x: "Não tem idade mínima" if x == 1 else "Não apresenta deficiência" if x == 2 else "Não cumpriu o requisito da miserabilidade")

    if motivo_improcedencia == 1:
        idade_insuficiente = st.number_input("Qual a idade do requerente na DER? (Digite apenas números):", min_value=0, max_value=150, step=1)
        if st.button("Gerar Sentença"):
            doc = Document()
            doc.add_paragraph(f"Processo: {processo_formatado}")
            texto_base(doc, fundamento_questao)
            fundamento_improcedencia1 = [
                (f"No caso dos autos trata-se de pedido de benefício de prestação continuada (LOAS) - Idoso. Verifica-se que na DER a parte autora possuía {idade_insuficiente} anos de idade, e o benefício exige 65 anos de idade. Não cumprido um dos requisitos legais, o pedido é improcedente."),
                (f"Isto posto, com resolução de mérito nos termos do art. 487, I, do CPC, JULGO IMPROCEDENTE o pedido."),
                (f"Sem condenação em honorários nesta instância."),
                (f"Defiro os benefícios da gratuidade."),
                (f"Em caso de interposição de embargos de declaração, intime-se a parte contrária para contrarrazões no prazo legal, e tornem conclusos para julgamento."),
                (f"Interposto recurso, intime-se a parte contrária para contrarrazões no prazo legal. Após, remetam-se os autos às Turmas Recursais."),
                (f"Com o trânsito em julgado, arquivem-se oportunamente."),
                (f"Int."),
                ]
            ft.alinhamento_parag_dispositivo(doc, fundamento_improcedencia1)
            ft.salvar_docx_temporario(doc, processo_formatado)

    elif motivo_improcedencia == 2:
        sem_deficiencia = st.radio("Por que não existe deficiência?", [1, 2, 3], format_func=lambda x: "Estado mórbido não impede participação social" if x == 1 else "Impedimento de longo prazo não comprovado" if x == 2 else "Outro motivo")
        if sem_deficiencia == 3:
            sem_deficiencia_redigido = st.text_area("Redija o motivo pelo qual o requerente não tem deficiência (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final):")
        if st.button("Gerar Sentença"):
            doc = Document()
            doc.add_paragraph(f"Processo: {processo_formatado}")
            texto_base(doc, fundamento_questao)
            fundamento_improcedencia2 = []
            if sem_deficiencia == 1:
                fundamento_improcedencia2.append(f"No caso dos autos trata-se de pedido de benefício de prestação continuada (LOAS) - Deficiente. Em que pese a perícia tenha constatado que a parte autora é acometida de estado mórbido que a aflige, não se constatou deficiência que a impeça de participar da vida social em igualdade de condições.")
            elif sem_deficiencia == 2:
                fundamento_improcedencia2.append(f"No caso dos autos trata-se de pedido de benefício de prestação continuada (LOAS) - Deficiente. A lei 8.742/93 considera deficiente somente quem possui impedimento de longo prazo, assim entendido aquele superior a 2 anos, ainda que em prognóstico. A perícia constatou que a parte autora não tem impedimento de longo prazo.")
            elif sem_deficiencia == 3:
                fundamento_improcedencia2.append(f"A parte autora não comprova a existência de deficiência no caso concreto.")
                for linha in sem_deficiencia_redigido.split("\n"):
                    if linha.strip():
                        fundamento_improcedencia2.append(linha)
            fundamento_improcedencia2.extend([
                (f"O caso não comporta que se produza outro laudo pericial, ou que se exija outros esclarecimentos do perito. O perito é claro em seu laudo."),
                (f"Não cumprido um dos requisitos legais, o pedido é improcedente."),
                (f"Isto posto, com resolução de mérito nos termos do art. 487, I, do CPC, JULGO IMPROCEDENTE o pedido."),
                (f"Sem condenação em honorários nesta instância."),
                (f"Defiro os benefícios da gratuidade."),
                (f"Em caso de interposição de embargos de declaração, intime-se a parte contrária para contrarrazões no prazo legal, e tornem conclusos para julgamento."),
                (f"Interposto recurso, intime-se a parte contrária para contrarrazões no prazo legal. Após, remetam-se os autos às Turmas Recursais."),
                (f"Com o trânsito em julgado, arquivem-se oportunamente."),
                (f"Int."),
                ])
            ft.alinhamento_parag_dispositivo(doc, fundamento_improcedencia2)
            ft.salvar_docx_temporario(doc, processo_formatado)

    elif motivo_improcedencia == 3:
        sem_miserabilidade = st.radio("Por que não existe miserabilidade?", [1, 2], format_func=lambda x: "Renda per capita familiar supera o limite legal" if x == 1 else "Outro motivo")
        if sem_miserabilidade == 2:
            sem_miserabilidade_redigido = st.text_area("Redija o motivo pelo qual o requerente não cumpre o requisito da miserabilidade (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final):")
        if st.button("Gerar Sentença"):
            doc = Document()
            doc.add_paragraph(f"Processo: {processo_formatado}")
            texto_base(doc, fundamento_questao)
            fundamento_improcedencia3 = []
            if sem_miserabilidade == 1:
                fundamento_improcedencia3.append(f"A perícia social constatou que a renda per capita familiar supera o limite legal que 1/4 do salário-mínimo, e a situação concreta apresentada no laudo demonstra que, apesar das dificuldades enfrentadas, a parte autora possui o necessário para sua manutenção.")
            elif sem_miserabilidade == 2:
                fundamento_improcedencia3.append(f"A parte autora não comprova a existência de miserabilidade no caso concreto. A situação concreta apresentada no laudo demonstra que, apesar das dificuldades enfrentadas, a parte autora possui o necessário para sua manutenção.")
                for linha in sem_miserabilidade_redigido.split("\n"):
                    if linha.split():
                        fundamento_improcedencia3.append(linha)                
            fundamento_improcedencia3.extend([
                (f"O caso não comporta que se produza outro laudo pericial, ou que se exija outros esclarecimentos do perito. O perito é claro em seu laudo."),
                (f"Não cumprido um dos requisitos legais, o pedido é improcedente."),
                (f"Isto posto, com resolução de mérito nos termos do art. 487, I, do CPC, JULGO IMPROCEDENTE o pedido."),
                (f"Sem condenação em honorários nesta instância."),
                (f"Defiro os benefícios da gratuidade."),
                (f"Em caso de interposição de embargos de declaração, intime-se a parte contrária para contrarrazões no prazo legal, e tornem conclusos para julgamento."),
                (f"Interposto recurso, intime-se a parte contrária para contrarrazões no prazo legal. Após, remetam-se os autos às Turmas Recursais."),
                (f"Com o trânsito em julgado, arquivem-se oportunamente."),
                (f"Int."),
                ])
            ft.alinhamento_parag_dispositivo(doc, fundamento_improcedencia3)
            ft.salvar_docx_temporario(doc, processo_formatado)

elif resultado == 1:
    tipo_de_loas = st.radio("Trata-se de LOAS Idoso ou LOAS Deficiente?", [1, 2], format_func=lambda x: "Idoso" if x == 1 else "Deficiente")

    if tipo_de_loas == 1:
        idade_idoso = st.number_input("Qual a idade do requerente na DER? (Digite apenas números):", min_value=0, max_value=150, step=1)
        miserabilidade_presente = st.text_area("Por que a parte autora cumpre o requisito de miserabilidade?")
        DIB = st.text_input("Qual a DIB do benefício concedido? (Digite no formato dd/mm/aaaa):")
        DIB_na_DER = st.radio("A DIB foi fixada na DER?", [1, 2], format_func=lambda x: "Sim" if x == 1 else "Não")
        if DIB_na_DER == 2:
            motivo_DIB = st.text_area("Explique por que a DIB não foi fixada na DER (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final):")
        else:
            motivo_DIB = "A DIB deve ser fixada na DER do benefício junto ao INSS."
        procedencia_total_parcial = st.radio("Para fins da redação do dispositivo da sentença, a procedência foi total ou parcial?", [1, 2], format_func=lambda x: "Total" if x == 1 else "Parcial")
        if procedencia_total_parcial == 1:
            resultado_dispositivo = ""
        else:
            resultado_dispositivo = "EM PARTE "

        if st.button("Gerar Sentença"):
            data_atual = datetime.now()
            DIP = data_atual.strftime("01/%m/%Y")

            doc = Document()
            doc.add_paragraph(f"Processo: {processo_formatado}")
            texto_base(doc, fundamento_questao)
            fundamento_procedencia1 = [
                (f"No presente caso, trata-se de pedido de benefício de prestação continuada - LOAS - Idoso."),
                (f"A parte autora possuía {idade_idoso} anos de idade no requerimento."),
                (f"Para comprovação da situação econômica foi realizada perícia socioeconômica, onde se vê que o requisito de miserabilidade foi cumprido.")
                ]
            for linha in miserabilidade_presente.split("\n"):
                if linha.split():
                    fundamento_procedencia1.append(linha)
            fundamento_procedencia1.extend([
                (f"Tendo em vista este quadro, e o posicionamento jurisprudencial, entendo que está comprovada a miserabilidade a que se refere a Constituição Federal para garantir ao autor o benefício pleiteado."),
                (f"Quanto à DIB, fixo em {DIB}. {motivo_DIB}"),
                (f"Isto posto, com resolução de mérito nos termos do art. 487, I, do CPC, JULGO PROCEDENTE {resultado_dispositivo}o pedido para condenar o réu a conceder a parte autora o benefício de prestação continuada – LOAS  Idoso, desde {DIB}, no valor de um salário mínimo vigente ao tempo."),
                (f"Condeno o INSS ao pagamento dos atrasados devidos desde a DIB fixada, até a DIP em {DIP}, atualizados desde cada competência devida e com juros desde a propositura da demanda, pelos índices e percentuais do Manual de Cálculos da Justiça Federal, a ser apurado em cumprimento invertido de sentença."),
                (f"Sem condenação em honorários nesta instância."),
                (f"Defiro os benefícios da gratuidade."),
                (f"Diante da certeza do direito, e do fundado receio de dano de difícil reparação ao autor, que comprovadamente necessita desta verba de natureza alimentar, CONCEDO A ANTECIPAÇÃO DE TUTELA para determinar a implantação do benefício no prazo de até 60 dias, com DIP em {DIP}."),
                (f"Em caso de interposição de embargos de declaração, intime-se a parte contrária para contrarrazões no prazo legal, e tornem conclusos para julgamento."),
                (f"Interposto recurso, intime-se a parte contrária para contrarrazões no prazo legal. Após, remetam-se os autos às Turmas Recursais."),
                (f"Com o trânsito em julgado, implantado o benefício, dê-se início ao cumprimento de sentença."),
                (f"Proceda a Secretaria como necessário."),
                (f"Int."),
                ])
            ft.alinhamento_parag_dispositivo(doc, fundamento_procedencia1)
            ft.salvar_docx_temporario(doc, processo_formatado)

    elif tipo_de_loas == 2:
        deficiencia = st.text_area("Por que a parte autora pode ser considerada deficiente? (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final)")
        miserabilidade_presente = st.text_area("Por que a parte autora cumpre o requisito de miserabilidade? (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final)")
        DIB1 = st.text_input("Qual a DIB do benefício concedido? (Digite no formato dd/mm/aaaa):")
        DIB_na_DER1 = st.radio("A DIB foi fixada na DER?", [1, 2], format_func=lambda x: "Sim" if x == 1 else "Não")
        if DIB_na_DER1 == 2:
            motivo_DIB1 = st.text_area("Explique por que a DIB não foi fixada na DER (será inserido como parágrafo na sentença - iniciar com letra maiúscula e colocar ponto final):")
        else:
            motivo_DIB1 = "A DIB deve ser fixada na DER do benefício junto ao INSS."
        procedencia_total_parcial = st.radio("Para fins da redação do dispositivo da sentença, a procedência foi total ou parcial?", [1, 2], format_func=lambda x: "Total" if x == 1 else "Parcial")
        if procedencia_total_parcial == 1:
            resultado_dispositivo = ""
        else:
            resultado_dispositivo = "EM PARTE "

        if st.button("Gerar Sentença"):
            data_atual = datetime.now()
            DIP1 = data_atual.strftime("01/%m/%Y")

            doc = Document()
            doc.add_paragraph(f"Processo: {processo_formatado}")
            texto_base(doc, fundamento_questao)
            fundamento_procedencia2 = [
                (f"No presente caso, trata-se de pedido de benefício de prestação continuada - LOAS - Deficiente."),
                (f"A parte autora enquadra-se como deficiente nos termos da lei."),
                ]

            for linha in deficiencia.split("\n"):
                if linha.split():
                    fundamento_procedencia2.append(linha)

            fundamento_procedencia2.append(f"Para comprovação da situação econômica foi realizada perícia socioeconômica, onde se vê que o requisito de miserabilidade foi cumprido.")
            
            for linha in miserabilidade_presente.split("\n"):
                if linha.split():
                    fundamento_procedencia2.append(linha)

            fundamento_procedencia2.extend([
                (f"Tendo em vista este quadro, e o posicionamento jurisprudencial, entendo que está comprovada a miserabilidade a que se refere a Constituição Federal para garantir ao autor o benefício pleiteado."),
                (f"Quanto à DIB, fixo em {DIB1}. {motivo_DIB1}"),
                (f"Isto posto, com resolução de mérito nos termos do art. 487, I, do CPC, JULGO PROCEDENTE {resultado_dispositivo}o pedido para condenar o réu a conceder a parte autora o benefício de prestação continuada – LOAS  Deficiente, desde {DIB1}, no valor de um salário mínimo vigente ao tempo."),
                (f"Condeno o INSS ao pagamento dos atrasados devidos desde a DIB fixada, até a DIP em {DIP1}, atualizados desde cada competência devida e com juros desde a propositura da demanda, pelos índices e percentuais do Manual de Cálculos da Justiça Federal, a ser apurado em cumprimento invertido de sentença."),
                (f"Sem condenação em honorários nesta instância."),
                (f"Defiro os benefícios da gratuidade."),
                (f"Diante da certeza do direito, e do fundado receio de dano de difícil reparação ao autor, que comprovadamente necessita desta verba de natureza alimentar, CONCEDO A ANTECIPAÇÃO DE TUTELA para determinar a implantação do benefício no prazo de até 60 dias, com DIP em {DIP1}."),
                (f"Em caso de interposição de embargos de declaração, intime-se a parte contrária para contrarrazões no prazo legal, e tornem conclusos para julgamento."),
                (f"Interposto recurso, intime-se a parte contrária para contrarrazões no prazo legal. Após, remetam-se os autos às Turmas Recursais."),
                (f"Com o trânsito em julgado, implantado o benefício, dê-se início ao cumprimento de sentença."),
                (f"Proceda a Secretaria como necessário."),
                (f"Int."),
                ])
            ft.alinhamento_parag_dispositivo(doc, fundamento_procedencia2)
            ft.salvar_docx_temporario(doc, processo_formatado)
